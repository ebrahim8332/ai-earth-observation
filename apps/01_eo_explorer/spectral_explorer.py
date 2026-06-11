"""
spectral_explorer.py — Core logic for the Spectral Explorer tab.

Responsibilities:
  - Search Planetary Computer for scenes matching location, date, satellite, cloud cover
  - Render any RGB band combination using the PC rendering API
  - Build a contact sheet (all presets as thumbnails)
  - Build a scene timeline chart
  - Generate AI explanations for band combinations

The black image fix from the Day 2 notebook is baked in:
  - Individual bands (not the visual composite asset)
  - rescale per satellite from satellite_catalog
  - gamma correction
  - Coverage scoring to find the best scene
  - STAC query filtered to avoid tile boundary issues
"""

import requests
import numpy as np
import plotly.graph_objects as go
import pystac_client
import planetary_computer
from PIL import Image
from io import BytesIO
from datetime import datetime

import config
import satellite_catalog
import ai_assistant

# Planetary Computer STAC API endpoint
PC_STAC_URL = "https://planetarycomputer.microsoft.com/api/stac/v1"

# Rendering API endpoint — renders the full item tile (100 km × 100 km for Sentinel-2).
# Geographic clipping is done locally after fetching using _clip_to_bbox().
PC_RENDER_URL = "https://planetarycomputer.microsoft.com/api/data/v1/item/preview"


# ---------------------------------------------------------------------------
# STAC search
# ---------------------------------------------------------------------------

def _pc_search_with_retry(search, retries=3, delay=5):
    """
    Call list(search.items()) with up to `retries` attempts.
    Planetary Computer occasionally returns a timeout on the free tier.
    Waiting `delay` seconds between attempts clears most transient errors.
    Raises the last exception if all attempts fail.
    """
    import time
    last_exc = None
    for attempt in range(1, retries + 1):
        try:
            return list(search.items())
        except Exception as e:
            last_exc = e
            if attempt < retries:
                time.sleep(delay)
    raise last_exc


def get_catalog():
    """
    Open and return a signed connection to the Planetary Computer STAC catalog.
    Cached in session state by the caller to avoid reconnecting on every interaction.
    """
    return pystac_client.Client.open(PC_STAC_URL, modifier=planetary_computer.sign_inplace)


def search_scenes(catalog, collection: str, bbox: list, date_range: str,
                  max_cloud: int, cloud_field: str | None) -> list:
    """
    Search the STAC catalog for scenes matching the given parameters.
    Returns a list of STAC items sorted by cloud cover (ascending).
    If cloud_field is None (SAR), no cloud filter is applied.
    """
    query = {}
    if cloud_field:
        query[cloud_field] = {"lt": max_cloud}

    search = catalog.search(
        collections=[collection],
        bbox=bbox,
        datetime=date_range,
        query=query if query else None,
    )

    items = _pc_search_with_retry(search)

    # Sort by cloud cover if available, else by date
    if cloud_field:
        items.sort(key=lambda x: x.properties.get(cloud_field, 100))
    else:
        items.sort(key=lambda x: x.datetime, reverse=True)

    return items


def score_scene_coverage(item, r_band: str, g_band: str, b_band: str,
                          satellite_key: str, width: int = 200) -> float:
    """
    Fetch a small test render and count non-zero pixels.
    Returns coverage as a percentage (0-100).
    Used to rank scenes by actual spatial coverage, not just cloud metadata.
    """
    arr = render_combination(item, r_band, g_band, b_band, satellite_key, width=width)
    if arr is None:
        return 0.0
    valid = (arr.max(axis=2) > 5).sum()
    total = arr.shape[0] * arr.shape[1]
    return round(100.0 * valid / total, 1)


def find_best_scene(items: list, r_band: str, g_band: str, b_band: str,
                    satellite_key: str, max_to_check: int = 6) -> tuple:
    """
    Score the top N scenes by coverage and return the best one.
    Returns (item, coverage_pct) or (None, 0) if nothing has valid data.
    """
    best_item    = None
    best_pct     = 0.0
    best_results = []

    for item in items[:max_to_check]:
        pct = score_scene_coverage(item, r_band, g_band, b_band, satellite_key)
        date  = item.datetime.strftime("%Y-%m-%d")
        cloud = item.properties.get("eo:cloud_cover", 0)
        best_results.append((date, cloud, pct, item))
        if pct > best_pct:
            best_pct  = pct
            best_item = item

    return best_item, best_pct, best_results


# ---------------------------------------------------------------------------
# Rendering
# ---------------------------------------------------------------------------

def render_combination(item, r_band: str, g_band: str, b_band: str,
                        satellite_key: str, width: int = 600,
                        bbox: list = None) -> np.ndarray | None:
    """
    Call the Planetary Computer rendering API to produce an RGB image from
    any three bands. Returns a numpy uint8 array or None on failure.

    The black image fix is applied here:
    - Individual named bands (not the visual composite)
    - Per-satellite rescale values from satellite_catalog
    - Gamma correction
    - Crop to valid (non-black) region

    bbox: optional [min_lon, min_lat, max_lon, max_lat] clip window.
    When provided, the bbox is passed directly to the PC rendering API
    which renders only the requested geographic window. This avoids all
    tile-edge and nodata-border issues from local pixel clipping.
    Falls back to local clip if the API does not accept the bbox param.
    """
    sat     = satellite_catalog.SATELLITES[satellite_key]
    rescale = sat["rescale"]
    gamma   = sat["color_formula"]

    # Base parameters — shared by both bbox and full-tile paths.
    base_params = [
        ("collection", sat["collection"]),
        ("item",       item.id),
        ("assets",     r_band),
        ("assets",     g_band),
        ("assets",     b_band),
        ("asset_bidx", f"{r_band}|1"),
        ("asset_bidx", f"{g_band}|1"),
        ("asset_bidx", f"{b_band}|1"),
        ("rescale",    rescale),
        ("rescale",    rescale),
        ("rescale",    rescale),
    ]
    if gamma:
        base_params.append(("color_formula", gamma))

    try:
        if bbox:
            # --- Server-side bbox crop (preferred path) ---
            # Pass the bbox directly to the PC/titiler rendering API.
            # The server renders only the requested geographic window,
            # which avoids tile-edge geometry and nodata-border issues.
            #
            # Adjust height to match the real-world aspect ratio of the
            # bbox, correcting for longitude compression at higher latitudes.
            lon_span = bbox[2] - bbox[0]
            lat_span = bbox[3] - bbox[1]
            lat_mid  = (bbox[1] + bbox[3]) / 2.0
            lon_km   = lon_span * 111.0 * float(np.cos(np.radians(lat_mid)))
            lat_km   = lat_span * 111.0
            ratio    = lat_km / lon_km if lon_km > 0 else 1.0
            ratio    = min(max(ratio, 0.33), 3.0)   # cap to reasonable range
            h        = int(width * ratio)

            api_params = base_params + [
                ("width",  str(width)),
                ("height", str(h)),
                ("bbox",   f"{bbox[0]},{bbox[1]},{bbox[2]},{bbox[3]}"),
            ]

            resp = requests.get(PC_RENDER_URL, params=api_params, timeout=60)
            if resp.status_code == 200:
                arr = np.array(Image.open(BytesIO(resp.content)).convert("RGB"))
                # Light clean-up: remove any residual nodata pixels at the edges.
                return _crop_to_valid(arr)

            # API didn't accept bbox — fall back to local clip below.
            params = base_params + [("width", str(width)), ("height", str(width))]
            resp = requests.get(PC_RENDER_URL, params=params, timeout=60)
            if resp.status_code != 200:
                return None

            arr     = np.array(Image.open(BytesIO(resp.content)).convert("RGB"))
            clipped = _clip_to_bbox(arr, item.bbox, bbox)
            cropped = _crop_to_valid(clipped, skip_ratio_guard=True)

            clip_pixels = clipped.shape[0] * clipped.shape[1]
            crop_pixels = cropped.shape[0] * cropped.shape[1]
            ch, cw      = cropped.shape[:2]
            crop_ratio  = max(ch / cw, cw / ch) if ch > 0 and cw > 0 else 99

            good_clip = (
                clip_pixels > 0
                and (crop_pixels / clip_pixels) >= 0.40
                and crop_ratio <= 3.0
            )
            return cropped if good_clip else _crop_to_valid(arr)

        else:
            # --- Full-tile path (no bbox) ---
            params = base_params + [("width", str(width)), ("height", str(width))]
            resp = requests.get(PC_RENDER_URL, params=params, timeout=60)
            if resp.status_code != 200:
                return None
            arr = np.array(Image.open(BytesIO(resp.content)).convert("RGB"))
            return _crop_to_valid(arr)

    except Exception:
        return None


def render_ndvi(item, satellite_key: str, width: int = 600,
                bbox: list = None) -> np.ndarray | None:
    """
    Render an NDVI map using the PC titiler expression endpoint.
    NDVI = (NIR - Red) / (NIR + Red).
    Red-yellow-green colormap: green = healthy vegetation, red = bare/stressed.
    bbox: optional [min_lon, min_lat, max_lon, max_lat] clip window.
    """
    sat = satellite_catalog.SATELLITES[satellite_key]
    if sat.get("sar"):
        return None

    if satellite_key == "Sentinel-2 L2A":
        expression = "(B08-B04)/(B08+B04)"
    elif satellite_key == "Landsat 8/9":
        expression = "(nir08-red)/(nir08+red)"
    else:
        return None

    params = [
        ("collection",    sat["collection"]),
        ("item",          item.id),
        ("expression",    expression),
        ("colormap_name", "rdylgn"),
        ("rescale",       "-1,1"),
        ("width",         str(width)),
        ("height",        str(width)),
    ]
    try:
        resp = requests.get(PC_RENDER_URL, params=params, timeout=60)
        if resp.status_code != 200:
            return None
        img = Image.open(BytesIO(resp.content)).convert("RGB")
        arr = np.array(img)
        if bbox:
            arr = _clip_to_bbox(arr, item.bbox, bbox)
        else:
            arr = _crop_to_valid(arr)
        return arr
    except Exception:
        return None


def render_ndwi(item, satellite_key: str, width: int = 600,
                bbox: list = None) -> np.ndarray | None:
    """
    Render an NDWI map (water body index).
    NDWI = (Green - NIR) / (Green + NIR).
    Blue colormap: bright blue = open water.
    bbox: optional [min_lon, min_lat, max_lon, max_lat] clip window.
    """
    sat = satellite_catalog.SATELLITES[satellite_key]
    if sat.get("sar"):
        return None

    if satellite_key == "Sentinel-2 L2A":
        expression = "(B03-B08)/(B03+B08)"
    elif satellite_key == "Landsat 8/9":
        expression = "(green-nir08)/(green+nir08)"
    else:
        return None

    params = [
        ("collection",    sat["collection"]),
        ("item",          item.id),
        ("expression",    expression),
        ("colormap_name", "blues"),
        ("rescale",       "-1,1"),
        ("width",         str(width)),
        ("height",        str(width)),
    ]
    try:
        resp = requests.get(PC_RENDER_URL, params=params, timeout=60)
        if resp.status_code != 200:
            return None
        img = Image.open(BytesIO(resp.content)).convert("RGB")
        arr = np.array(img)
        if bbox:
            arr = _clip_to_bbox(arr, item.bbox, bbox)
        else:
            arr = _crop_to_valid(arr)
        return arr
    except Exception:
        return None


def render_ndmi(item, satellite_key: str, width: int = 600,
                bbox: list = None) -> np.ndarray | None:
    """
    Render an NDMI map (vegetation moisture index).
    NDMI = (NIR - SWIR1) / (NIR + SWIR1).
    Blue colormap: bright blue = high moisture, pale = dry or stressed.
    """
    sat = satellite_catalog.SATELLITES[satellite_key]
    if sat.get("sar"):
        return None

    if satellite_key == "Sentinel-2 L2A":
        expression = "(B08-B11)/(B08+B11)"
    elif satellite_key == "Landsat 8/9":
        expression = "(nir08-swir16)/(nir08+swir16)"
    else:
        return None

    params = [
        ("collection",    sat["collection"]),
        ("item",          item.id),
        ("expression",    expression),
        ("colormap_name", "blues"),
        ("rescale",       "-0.5,0.5"),
        ("width",         str(width)),
        ("height",        str(width)),
    ]
    try:
        resp = requests.get(PC_RENDER_URL, params=params, timeout=60)
        if resp.status_code != 200:
            return None
        img = Image.open(BytesIO(resp.content)).convert("RGB")
        arr = np.array(img)
        if bbox:
            arr = _clip_to_bbox(arr, item.bbox, bbox)
        else:
            arr = _crop_to_valid(arr)
        return arr
    except Exception:
        return None


def render_nbr(item, satellite_key: str, width: int = 600,
               bbox: list = None) -> np.ndarray | None:
    """
    Render an NBR map (Normalized Burn Ratio).
    NBR = (NIR - SWIR2) / (NIR + SWIR2).
    Green-yellow-red colormap: green = healthy vegetation, red = burned or bare.
    """
    sat = satellite_catalog.SATELLITES[satellite_key]
    if sat.get("sar"):
        return None

    if satellite_key == "Sentinel-2 L2A":
        expression = "(B08-B12)/(B08+B12)"
    elif satellite_key == "Landsat 8/9":
        expression = "(nir08-swir22)/(nir08+swir22)"
    else:
        return None

    params = [
        ("collection",    sat["collection"]),
        ("item",          item.id),
        ("expression",    expression),
        ("colormap_name", "rdylgn"),
        ("rescale",       "-1,1"),
        ("width",         str(width)),
        ("height",        str(width)),
    ]
    try:
        resp = requests.get(PC_RENDER_URL, params=params, timeout=60)
        if resp.status_code != 200:
            return None
        img = Image.open(BytesIO(resp.content)).convert("RGB")
        arr = np.array(img)
        if bbox:
            arr = _clip_to_bbox(arr, item.bbox, bbox)
        else:
            arr = _crop_to_valid(arr)
        return arr
    except Exception:
        return None


def render_savi(item, satellite_key: str, width: int = 600,
                bbox: list = None) -> np.ndarray | None:
    """
    Render a SAVI map (Soil-Adjusted Vegetation Index).
    SAVI = 1.5 * (NIR - Red) / (NIR + Red + 0.5).
    Same green-yellow-red colormap as NDVI but corrects for bright bare soil.
    Better than NDVI in arid and semi-arid areas with sparse vegetation.
    """
    sat = satellite_catalog.SATELLITES[satellite_key]
    if sat.get("sar"):
        return None

    if satellite_key == "Sentinel-2 L2A":
        expression = "1.5*(B08-B04)/(B08+B04+0.5)"
    elif satellite_key == "Landsat 8/9":
        expression = "1.5*(nir08-red)/(nir08+red+0.5)"
    else:
        return None

    params = [
        ("collection",    sat["collection"]),
        ("item",          item.id),
        ("expression",    expression),
        ("colormap_name", "rdylgn"),
        ("rescale",       "0,0.8"),
        ("width",         str(width)),
        ("height",        str(width)),
    ]
    try:
        resp = requests.get(PC_RENDER_URL, params=params, timeout=60)
        if resp.status_code != 200:
            return None
        img = Image.open(BytesIO(resp.content)).convert("RGB")
        arr = np.array(img)
        if bbox:
            arr = _clip_to_bbox(arr, item.bbox, bbox)
        else:
            arr = _crop_to_valid(arr)
        return arr
    except Exception:
        return None


def render_evi(item, satellite_key: str, width: int = 600,
               bbox: list = None) -> np.ndarray | None:
    """
    Render an EVI map (Enhanced Vegetation Index).
    EVI = 2.5 * (NIR - Red) / (NIR + 6*Red - 7.5*Blue + 1).
    Better than NDVI in dense canopy areas where NDVI saturates.
    Green-yellow-red colormap: green = dense healthy vegetation.
    """
    sat = satellite_catalog.SATELLITES[satellite_key]
    if sat.get("sar"):
        return None

    if satellite_key == "Sentinel-2 L2A":
        expression = "2.5*(B08-B04)/(B08+6*B04-7.5*B02+1)"
    elif satellite_key == "Landsat 8/9":
        expression = "2.5*(nir08-red)/(nir08+6*red-7.5*blue+1)"
    else:
        return None

    params = [
        ("collection",    sat["collection"]),
        ("item",          item.id),
        ("expression",    expression),
        ("colormap_name", "rdylgn"),
        ("rescale",       "-0.2,1"),
        ("width",         str(width)),
        ("height",        str(width)),
    ]
    try:
        resp = requests.get(PC_RENDER_URL, params=params, timeout=60)
        if resp.status_code != 200:
            return None
        img = Image.open(BytesIO(resp.content)).convert("RGB")
        arr = np.array(img)
        if bbox:
            arr = _clip_to_bbox(arr, item.bbox, bbox)
        else:
            arr = _crop_to_valid(arr)
        return arr
    except Exception:
        return None


def render_bsi(item, satellite_key: str, width: int = 600,
               bbox: list = None) -> np.ndarray | None:
    """
    Render a BSI map (Bare Soil Index).
    BSI = ((SWIR1 + Red) - (NIR + Blue)) / ((SWIR1 + Red) + (NIR + Blue)).
    Orange-red colormap: bright orange-red = exposed bare soil.
    """
    sat = satellite_catalog.SATELLITES[satellite_key]
    if sat.get("sar"):
        return None

    if satellite_key == "Sentinel-2 L2A":
        expression = "((B11+B04)-(B08+B02))/((B11+B04)+(B08+B02))"
    elif satellite_key == "Landsat 8/9":
        expression = "((swir16+red)-(nir08+blue))/((swir16+red)+(nir08+blue))"
    else:
        return None

    params = [
        ("collection",    sat["collection"]),
        ("item",          item.id),
        ("expression",    expression),
        ("colormap_name", "ylorrd"),
        ("rescale",       "-0.5,0.5"),
        ("width",         str(width)),
        ("height",        str(width)),
    ]
    try:
        resp = requests.get(PC_RENDER_URL, params=params, timeout=60)
        if resp.status_code != 200:
            return None
        img = Image.open(BytesIO(resp.content)).convert("RGB")
        arr = np.array(img)
        if bbox:
            arr = _clip_to_bbox(arr, item.bbox, bbox)
        else:
            arr = _crop_to_valid(arr)
        return arr
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Contact sheet (all presets as thumbnails)
# ---------------------------------------------------------------------------

def render_contact_sheet(item, satellite_key: str,
                          include_ndvi: bool = True,
                          include_ndwi: bool = True,
                          include_indices: bool = True,
                          thumb_size: int = 300,
                          bbox: list = None) -> list:
    """
    Render every named preset for the selected satellite as a small thumbnail.
    Also renders NDVI and NDWI if requested and supported.

    Returns a list of dicts: [{label, note, array}, ...]
    Arrays are numpy uint8. None means the render failed.
    bbox: optional [min_lon, min_lat, max_lon, max_lat] clip window.
    """
    sat     = satellite_catalog.SATELLITES[satellite_key]
    bands   = sat["bands"]
    presets = sat["presets"]
    results = []

    for label, preset in presets.items():
        r, g, b = preset["r"], preset["g"], preset["b"]
        arr = render_combination(
            item, r, g, b,
            satellite_key,
            width=thumb_size,
            bbox=bbox,
        )
        # Build short channel label: "R → B11 (SWIR 1)  ·  G → B08 (NIR)  ·  B → B04 (Red)"
        def band_label(code):
            info = bands.get(code, {})
            name = info.get("name", code)
            return f"{code} ({name})"

        results.append({
            "label":    label,
            "note":     preset["note"],
            "array":    arr,
            "type":     "preset",
            "channels": f"R → {band_label(r)}  ·  G → {band_label(g)}  ·  B → {band_label(b)}",
        })

    if include_ndvi and not sat.get("sar"):
        arr = render_ndvi(item, satellite_key, width=thumb_size, bbox=bbox)
        if satellite_key == "Sentinel-2 L2A":
            expr = "(B08 − B04) / (B08 + B04)"
        else:
            expr = "(NIR − Red) / (NIR + Red)"
        results.append({
            "label":    "NDVI",
            "note":     "Vegetation health index. Green = healthy. Red = bare or stressed.",
            "array":    arr,
            "type":     "index",
            "channels": f"Expression: {expr}",
        })

    if include_ndwi and not sat.get("sar"):
        arr = render_ndwi(item, satellite_key, width=thumb_size, bbox=bbox)
        if satellite_key == "Sentinel-2 L2A":
            expr = "(B03 − B08) / (B03 + B08)"
        else:
            expr = "(Green − NIR) / (Green + NIR)"
        results.append({
            "label":    "NDWI",
            "note":     "Water body index. Bright blue = open water.",
            "array":    arr,
            "type":     "index",
            "channels": f"Expression: {expr}",
        })

    if include_indices and not sat.get("sar"):

        # NDMI — vegetation moisture
        arr = render_ndmi(item, satellite_key, width=thumb_size, bbox=bbox)
        if satellite_key == "Sentinel-2 L2A":
            expr = "(B08 − B11) / (B08 + B11)"
        else:
            expr = "(NIR − SWIR1) / (NIR + SWIR1)"
        results.append({
            "label":    "NDMI",
            "note":     "Vegetation moisture. Bright blue = high water content. Pale = dry or stressed.",
            "array":    arr,
            "type":     "index",
            "channels": f"Expression: {expr}",
        })

        # NBR — burn severity
        arr = render_nbr(item, satellite_key, width=thumb_size, bbox=bbox)
        if satellite_key == "Sentinel-2 L2A":
            expr = "(B08 − B12) / (B08 + B12)"
        else:
            expr = "(NIR − SWIR2) / (NIR + SWIR2)"
        results.append({
            "label":    "NBR",
            "note":     "Burn ratio. Green = healthy vegetation. Red = burned or severely damaged.",
            "array":    arr,
            "type":     "index",
            "channels": f"Expression: {expr}",
        })

        # SAVI — soil-adjusted vegetation
        arr = render_savi(item, satellite_key, width=thumb_size, bbox=bbox)
        if satellite_key == "Sentinel-2 L2A":
            expr = "1.5 × (B08 − B04) / (B08 + B04 + 0.5)"
        else:
            expr = "1.5 × (NIR − Red) / (NIR + Red + 0.5)"
        results.append({
            "label":    "SAVI",
            "note":     "Soil-adjusted vegetation. Better than NDVI in arid areas with sparse cover.",
            "array":    arr,
            "type":     "index",
            "channels": f"Expression: {expr}",
        })

        # EVI — enhanced vegetation
        arr = render_evi(item, satellite_key, width=thumb_size, bbox=bbox)
        if satellite_key == "Sentinel-2 L2A":
            expr = "2.5 × (B08 − B04) / (B08 + 6×B04 − 7.5×B02 + 1)"
        else:
            expr = "2.5 × (NIR − Red) / (NIR + 6×Red − 7.5×Blue + 1)"
        results.append({
            "label":    "EVI",
            "note":     "Enhanced vegetation. Better than NDVI in dense canopy where NDVI saturates.",
            "array":    arr,
            "type":     "index",
            "channels": f"Expression: {expr}",
        })

        # BSI — bare soil
        arr = render_bsi(item, satellite_key, width=thumb_size, bbox=bbox)
        if satellite_key == "Sentinel-2 L2A":
            expr = "((B11 + B04) − (B08 + B02)) / ((B11 + B04) + (B08 + B02))"
        else:
            expr = "((SWIR1 + Red) − (NIR + Blue)) / ((SWIR1 + Red) + (NIR + Blue))"
        results.append({
            "label":    "BSI",
            "note":     "Bare soil index. Bright orange-red = exposed bare soil or cleared land.",
            "array":    arr,
            "type":     "index",
            "channels": f"Expression: {expr}",
        })

    return results


# ---------------------------------------------------------------------------
# Index statistics (approximate — derived from greyscale render pixel values)
# ---------------------------------------------------------------------------

# Rescale ranges used when rendering each index — maps pixel 0-255 back to real value
_INDEX_RESCALE = {
    "NDVI": (-1.0,  1.0),
    "NDWI": (-1.0,  1.0),
    "NDMI": (-0.5,  0.5),
    "NBR":  (-1.0,  1.0),
    "SAVI": ( 0.0,  0.8),
    "EVI":  (-0.2,  1.0),
    "BSI":  (-0.5,  0.5),
}

_INDEX_EXPRESSIONS = {
    "Sentinel-2 L2A": {
        "NDVI": "(B08-B04)/(B08+B04)",
        "NDWI": "(B03-B08)/(B03+B08)",
        "NDMI": "(B08-B11)/(B08+B11)",
        "NBR":  "(B08-B12)/(B08+B12)",
        "SAVI": "1.5*(B08-B04)/(B08+B04+0.5)",
        "EVI":  "2.5*(B08-B04)/(B08+6*B04-7.5*B02+1)",
        "BSI":  "((B11+B04)-(B08+B02))/((B11+B04)+(B08+B02))",
    },
    "Landsat 8/9": {
        "NDVI": "(nir08-red)/(nir08+red)",
        "NDWI": "(green-nir08)/(green+nir08)",
        "NDMI": "(nir08-swir16)/(nir08+swir16)",
        "NBR":  "(nir08-swir22)/(nir08+swir22)",
        "SAVI": "1.5*(nir08-red)/(nir08+red+0.5)",
        "EVI":  "2.5*(nir08-red)/(nir08+6*red-7.5*blue+1)",
        "BSI":  "((swir16+red)-(nir08+blue))/((swir16+red)+(nir08+blue))",
    },
}

_BAND_EXPRESSIONS = {
    "Sentinel-2 L2A": ["B02", "B03", "B04", "B08", "B11", "B12"],
    "Landsat 8/9":    ["blue", "green", "red", "nir08", "swir16", "swir22"],
}

_BAND_WAVELENGTHS = {
    "Sentinel-2 L2A": {
        "B02": 490, "B03": 560, "B04": 665,
        "B08": 842, "B11": 1610, "B12": 2190,
    },
    "Landsat 8/9": {
        "blue": 485, "green": 560, "red": 660,
        "nir08": 865, "swir16": 1610, "swir22": 2200,
    },
}


def compute_index_stats(item, satellite_key: str, bbox: list = None) -> dict:
    """Fetch each spectral index as a greyscale image and compute min/mean/max.

    Pixel values (0-255) are mapped back to real index values using the known
    rescale range for each index. Returns a dict keyed by index name.
    Each value is {'min': float, 'mean': float, 'max': float} or None on failure.
    """
    sat = satellite_catalog.SATELLITES[satellite_key]
    if sat.get("sar"):
        return {}

    expressions = _INDEX_EXPRESSIONS.get(satellite_key, {})
    results = {}

    for idx_name, expr in expressions.items():
        lo, hi = _INDEX_RESCALE[idx_name]
        params = [
            ("collection",    sat["collection"]),
            ("item",          item.id),
            ("expression",    expr),
            ("colormap_name", "greys"),
            ("rescale",       f"{lo},{hi}"),
            ("width",         "200"),
            ("height",        "200"),
        ]
        if bbox:
            params.append(("bbox", f"{bbox[0]},{bbox[1]},{bbox[2]},{bbox[3]}"))
        try:
            resp = requests.get(PC_RENDER_URL, params=params, timeout=30)
            if resp.status_code != 200:
                results[idx_name] = None
                continue
            arr = np.array(Image.open(BytesIO(resp.content)).convert("L"), dtype=np.float32)
            # Remove nodata (pure black or pure white borders)
            valid = (arr > 5) & (arr < 250)
            if not valid.any():
                results[idx_name] = None
                continue
            vals = arr[valid] / 255.0 * (hi - lo) + lo
            results[idx_name] = {
                "min":  float(np.percentile(vals, 5)),
                "mean": float(np.mean(vals)),
                "max":  float(np.percentile(vals, 95)),
            }
        except Exception:
            results[idx_name] = None

    return results


def compute_spectral_signature(item, satellite_key: str, bbox: list = None) -> dict:
    """Fetch each optical band as greyscale and return mean reflectance per band.

    Returns a dict: band_name -> {'wavelength': int, 'mean_reflectance': float}
    Reflectance is approximate (derived from rendered pixel values).
    """
    sat = satellite_catalog.SATELLITES[satellite_key]
    if sat.get("sar"):
        return {}

    band_list = _BAND_EXPRESSIONS.get(satellite_key, [])
    wavelengths = _BAND_WAVELENGTHS.get(satellite_key, {})
    rescale = sat["rescale"]  # e.g. "0,3000" or "7000,22000"
    lo_r, hi_r = [float(x) for x in rescale.split(",")]
    results = {}

    for band in band_list:
        params = [
            ("collection",    sat["collection"]),
            ("item",          item.id),
            ("assets",        band),
            ("asset_bidx",    f"{band}|1"),
            ("rescale",       rescale),
            ("colormap_name", "greys"),
            ("width",         "150"),
            ("height",        "150"),
        ]
        if bbox:
            params.append(("bbox", f"{bbox[0]},{bbox[1]},{bbox[2]},{bbox[3]}"))
        try:
            resp = requests.get(PC_RENDER_URL, params=params, timeout=30)
            if resp.status_code != 200:
                continue
            arr = np.array(Image.open(BytesIO(resp.content)).convert("L"), dtype=np.float32)
            valid = (arr > 5) & (arr < 250)
            if not valid.any():
                continue
            mean_px = float(np.mean(arr[valid]))
            # Map pixel (0-255) back to DN, then normalise to 0-1 reflectance factor
            mean_dn = mean_px / 255.0 * (hi_r - lo_r) + lo_r
            reflectance = mean_dn / 10000.0  # Sentinel-2 scale factor
            results[band] = {
                "wavelength":       wavelengths.get(band, 0),
                "mean_reflectance": round(reflectance, 4),
            }
        except Exception:
            continue

    return results


# ---------------------------------------------------------------------------
# Spectral signature chart
# ---------------------------------------------------------------------------

def build_spectral_signature_chart(sig: dict, satellite_key: str) -> go.Figure:
    """Build a bar chart of mean reflectance per band — the spectral signature."""
    sat = satellite_catalog.SATELLITES[satellite_key]
    bands_meta = sat["bands"]

    items = sorted(sig.items(), key=lambda x: x[1]["wavelength"])
    labels = []
    values = []
    colors = []

    wavelength_colors = {
        range(400, 500):  "#4169e1",   # blue
        range(500, 600):  "#3cb371",   # green
        range(600, 700):  "#e74c3c",   # red
        range(700, 1000): "#8b0000",   # NIR — dark red
        range(1000, 2000):"#d2691e",   # SWIR1 — brown
        range(2000, 3000):"#8b4513",   # SWIR2 — darker brown
    }

    def get_color(wl):
        for rng, col in wavelength_colors.items():
            if wl in rng:
                return col
        return "#888888"

    for band, data in items:
        name = bands_meta.get(band, {}).get("name", band)
        labels.append(f"{band}\n{name}\n{data['wavelength']} nm")
        values.append(data["mean_reflectance"])
        colors.append(get_color(data["wavelength"]))

    fig = go.Figure(go.Bar(
        x=labels,
        y=values,
        marker_color=colors,
        text=[f"{v:.3f}" for v in values],
        textposition="outside",
    ))
    fig.update_layout(
        title="Spectral signature — mean reflectance per band",
        yaxis_title="Mean reflectance (0–1)",
        height=340,
        margin=dict(l=10, r=20, t=45, b=20),
        plot_bgcolor="#f8f8f8",
        paper_bgcolor="#ffffff",
    )
    return fig


# ---------------------------------------------------------------------------
# AI interpretation — integrated scene analysis
# ---------------------------------------------------------------------------

def get_scene_interpretation(contact_results: list, index_stats: dict,
                              satellite_key: str, location_name: str,
                              scene_date: str, scene_cloud: float) -> tuple:
    """Generate an integrated AI interpretation of the full contact sheet results.

    Takes all rendered combinations + index stats as input.
    Returns (text, model_name).
    """
    # Summarise what each combination showed
    combo_summary = []
    for r in contact_results:
        combo_summary.append(f"- {r['label']}: {r['note']}")

    # Summarise index stats
    idx_summary = []
    for idx, stats in index_stats.items():
        if stats:
            idx_summary.append(
                f"- {idx}: min {stats['min']:+.3f}, mean {stats['mean']:+.3f}, max {stats['max']:+.3f}"
            )

    prompt = f"""You are an Earth Observation analyst interpreting a multi-spectral satellite scene.

Satellite: {satellite_key}
Location: {location_name}
Scene date: {scene_date}
Cloud cover: {scene_cloud:.1f}%

Band combinations rendered:
{chr(10).join(combo_summary)}

Spectral index statistics (5th-95th percentile range):
{chr(10).join(idx_summary) if idx_summary else "No index stats available."}

Write a detailed integrated analysis covering all four sections below. \
Each section should be a full paragraph of 4-6 sentences. \
Do not compress — the reader needs depth, not brevity.

**Section 1 — Landscape Character:** Describe what this scene reveals about the landscape \
based on all the band combinations together. What land cover types are present? \
What is the dominant character of the area? What stands out across multiple views?

**Section 2 — Spectral Signals:** Interpret the index statistics. What do the NDVI, NDWI, \
and other index values tell us about vegetation health, water presence, soil exposure, \
and surface moisture? Reference specific values where they are informative.

**Section 3 — Sensor Insight:** What does this satellite ({satellite_key}) reveal about \
this location that would not be visible in a standard photograph? Which band combinations \
were most diagnostic for this specific landscape, and why?

**Section 4 — Decision Application:** Name one specific stakeholder and describe exactly \
how they would use this multi-spectral analysis to make a real operational decision. \
Be specific about the decision, the data they would act on, and what the gap would be \
without satellite imagery.

Write in plain language. Use the bold section headings. Be direct and thorough."""

    import ai_chain
    text, model = ai_chain.complete(
        prompt,
        groq_key=config.GROQ_API_KEY,
        gemini_key=config.GEMINI_API_KEY,
    )
    return text, model


def _fallback_scene_interpretation(location_name: str, satellite_key: str) -> str:
    return f"""**Multi-spectral scene analysis — {location_name}**

This scene was analysed across multiple band combinations using {satellite_key} imagery. Each combination reveals a different physical property of the surface — from chlorophyll absorption in the red band to soil moisture in the shortwave infrared.

To enable AI interpretation, add a GROQ_API_KEY or GEMINI_API_KEY to your .env file."""


# ---------------------------------------------------------------------------
# Word document builder
# ---------------------------------------------------------------------------

def build_spectral_docx(contact_results, index_stats, spectral_sig,
                         location_name, scene_date, scene_cloud,
                         satellite_key, ai_text, ai_model):
    """Build a Word document for the Spectral Explorer contact sheet.

    Sections:
      1. Title and scene metadata
      2. Contact sheet thumbnails — 3-column grid
      3. Spectral index statistics table
      4. Band reference table
      5. AI Interpretation
    """
    import io as _io
    from docx import Document as _Document
    from docx.shared import Inches as _Inches, Pt as _Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _ALIGN
    from docx.oxml import OxmlElement as _OxmlElement
    from docx.oxml.ns import qn as _qn
    from PIL import Image as _PIL

    doc = _Document()

    # ---- Title ----
    title = doc.add_heading("Spectral Explorer Report", 0)
    title.alignment = _ALIGN.CENTER
    sub = doc.add_paragraph()
    sub.alignment = _ALIGN.CENTER
    sub.add_run(f"{location_name}   |   {scene_date}   |   {satellite_key}   |   Cloud {scene_cloud:.1f}%").bold = True
    doc.add_paragraph()

    # ---- Contact sheet thumbnails — 3 per row ----
    doc.add_heading("Band Combinations and Indices", level=1)
    cols_per_row = 3
    for row_start in range(0, len(contact_results), cols_per_row):
        row_items = contact_results[row_start:row_start + cols_per_row]
        # Two Word table rows per group: images + captions
        tbl = doc.add_table(rows=2, cols=cols_per_row)
        # Remove borders
        for row in tbl.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = _OxmlElement("w:tcBorders")
                for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    border = _OxmlElement(f"w:{side}")
                    border.set(_qn("w:val"), "none")
                    tcBorders.append(border)
                tcPr.append(tcBorders)
        for col_idx, result in enumerate(row_items):
            img_cell = tbl.rows[0].cells[col_idx]
            cap_cell = tbl.rows[1].cells[col_idx]
            arr = result.get("array")
            if arr is not None:
                buf = _io.BytesIO()
                _PIL.fromarray(arr).save(buf, format="PNG")
                buf.seek(0)
                run = img_cell.paragraphs[0].add_run()
                run.add_picture(buf, width=_Inches(1.9))
            cap_p = cap_cell.paragraphs[0]
            cap_p.add_run(result["label"]).bold = True
            cap_cell.add_paragraph(result["note"]).style
        doc.add_paragraph()

    # ---- Index statistics table ----
    if index_stats:
        doc.add_heading("Spectral Index Statistics", level=1)
        doc.add_paragraph(
            "Values represent the 5th–95th percentile range of valid pixels in the scene."
        )
        idx_tbl = doc.add_table(rows=1, cols=4)
        idx_tbl.style = "Table Grid"
        for cell, txt in zip(idx_tbl.rows[0].cells, ["Index", "Min", "Mean", "Max"]):
            cell.text = txt
            cell.paragraphs[0].runs[0].bold = True
        for idx_name, stats in index_stats.items():
            if stats:
                row = idx_tbl.add_row().cells
                row[0].text = idx_name
                row[1].text = f"{stats['min']:+.3f}"
                row[2].text = f"{stats['mean']:+.3f}"
                row[3].text = f"{stats['max']:+.3f}"
        doc.add_paragraph()

    # ---- Band reference table ----
    doc.add_heading("Band Reference", level=1)
    sat_meta = satellite_catalog.SATELLITES.get(satellite_key, {})
    bands_meta = sat_meta.get("bands", {})
    if bands_meta:
        band_tbl = doc.add_table(rows=1, cols=4)
        band_tbl.style = "Table Grid"
        for cell, txt in zip(band_tbl.rows[0].cells, ["Band", "Name", "Wavelength", "What it measures"]):
            cell.text = txt
            cell.paragraphs[0].runs[0].bold = True
        for band_id, meta in bands_meta.items():
            row = band_tbl.add_row().cells
            row[0].text = band_id
            row[1].text = meta.get("name", "")
            row[2].text = meta.get("wavelength", "")
            row[3].text = meta.get("description", "")
        doc.add_paragraph()

    # ---- AI Interpretation ----
    if ai_text:
        doc.add_heading("AI Interpretation", level=1)
        if ai_model:
            doc.add_paragraph(f"Model: {ai_model}").italic = True
        import re as _re
        for line in ai_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            parts = _re.split(r"(\*\*.*?\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                elif part:
                    p.add_run(part)

    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Scene timeline chart
# ---------------------------------------------------------------------------

def build_timeline(items: list, cloud_field: str | None,
                   coverage_scores: list = None) -> go.Figure:
    """
    Build a Plotly bar chart showing available scenes over time.
    Bars are coloured by cloud cover: green = clear, yellow = some cloud, red = cloudy.
    Hovering shows the scene date, cloud cover, and coverage score if available.
    """
    if not items:
        return go.Figure()

    scores_map = {}
    if coverage_scores:
        for date, cloud, pct, item in coverage_scores:
            scores_map[item.id] = pct

    dates  = []
    clouds = []
    colors = []
    texts  = []

    for item in items:
        date  = item.datetime.strftime("%Y-%m-%d")
        cloud = item.properties.get(cloud_field, 0) if cloud_field else 0
        pct   = scores_map.get(item.id, None)

        dates.append(date)
        clouds.append(cloud)

        # Colour by cloud cover
        if cloud < 5:
            colors.append("#2ecc71")     # green
        elif cloud < 15:
            colors.append("#f39c12")     # amber
        else:
            colors.append("#e74c3c")     # red

        hover = f"Date: {date}<br>Cloud: {cloud:.1f}%"
        if pct is not None:
            hover += f"<br>Coverage: {pct:.0f}%"
        texts.append(hover)

    fig = go.Figure(go.Bar(
        x=dates,
        y=[1] * len(dates),
        marker_color=colors,
        hovertext=texts,
        hoverinfo="text",
    ))

    fig.update_layout(
        title="Available scenes (green = clear, amber = some cloud, red = cloudy)",
        xaxis_title="Date",
        yaxis_visible=False,
        height=160,
        margin=dict(l=10, r=10, t=40, b=40),
        plot_bgcolor="#f8f8f8",
        paper_bgcolor="#ffffff",
    )
    return fig


# ---------------------------------------------------------------------------
# AI explanation
# ---------------------------------------------------------------------------

def explain_combination(r_band: str, g_band: str, b_band: str,
                         satellite_key: str, location_name: str,
                         is_index: bool = False, index_name: str = "") -> str:
    """
    Generate an AI explanation of what the selected band combination reveals
    for the given location. Uses the same fallback chain as the rest of the app.
    """
    sat   = satellite_catalog.SATELLITES[satellite_key]
    bands = sat["bands"]

    if is_index:
        prompt = (
            f"Explain what the {index_name} index shows when applied to {satellite_key} "
            f"imagery over {location_name}. Cover: what the index measures, what the "
            f"colour scale means (bright vs dark), and one specific thing to look for "
            f"in this location. Plain language, under 150 words."
        )
    else:
        r_name = bands.get(r_band, {}).get("name", r_band)
        g_name = bands.get(g_band, {}).get("name", g_band)
        b_name = bands.get(b_band, {}).get("name", b_band)
        prompt = (
            f"Explain what a satellite image with Red={r_name}, Green={g_name}, "
            f"Blue={b_name} reveals when looking at {location_name} from {satellite_key}. "
            f"Cover: what stands out in each colour, what land cover types are identifiable, "
            f"and one specific insight for this location. Plain language, under 150 words."
        )

    return ai_assistant.auto_explain(
        theme="EO Basics",
        dataset=satellite_key,
        location=location_name,
        mode="Explain selected dataset",
    ) if not config.has_any_key() else ai_assistant.ask(
        question=prompt,
        theme="EO Basics",
        dataset=satellite_key,
        location=location_name,
        mode="Explain selected dataset",
    )


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _pad_to_ratio(arr: np.ndarray, max_ratio: float = 2.0) -> np.ndarray:
    """
    Pad a rendered array with black pixels to prevent extreme aspect ratios.

    When a large map picker radius is selected, the clipped region can be
    much taller than it is wide (or vice versa), producing a very long image
    in the display. This pads the shorter dimension with black so the final
    image is never more than max_ratio:1 in either direction.

    max_ratio=2.0 means: width can be at most 2× the height, and vice versa.
    """
    h, w = arr.shape[:2]
    if h == 0 or w == 0:
        return arr

    ratio = max(w / h, h / w)
    if ratio <= max_ratio:
        return arr

    channels = arr.shape[2] if arr.ndim == 3 else 1

    if w > h:
        # Wider than tall — pad top and bottom
        new_h = int(w / max_ratio)
        pad   = (new_h - h) // 2
        out   = np.zeros((new_h, w, channels), dtype=arr.dtype)
        out[pad:pad + h, :] = arr
    else:
        # Taller than wide — pad left and right
        new_w = int(h / max_ratio)
        pad   = (new_w - w) // 2
        out   = np.zeros((h, new_w, channels), dtype=arr.dtype)
        out[:, pad:pad + w] = arr

    return out


def _clip_to_bbox(arr: np.ndarray, item_bbox: list, clip_bbox: list) -> np.ndarray:
    """
    Clip a rendered satellite image array to a geographic sub-bbox.

    The PC Render API returns the full item tile. This function converts the
    desired geographic bbox into pixel row/column positions within the image
    and crops to that section.

    arr:       numpy H×W×3 array (the full rendered tile from the API)
    item_bbox: [west, south, east, north] of the full tile — from item.bbox
    clip_bbox: [west, south, east, north] of the desired clip area

    Returns the cropped array. Returns the original array if the clip bbox
    is outside the item extent or produces a zero-size result.
    """
    h, w = arr.shape[:2]

    item_west, item_south, item_east, item_north = item_bbox
    clip_west, clip_south, clip_east, clip_north = clip_bbox

    lon_span = item_east  - item_west
    lat_span = item_north - item_south

    if lon_span <= 0 or lat_span <= 0:
        return arr

    # x (column) increases eastward; y (row) increases southward in image coords
    x0 = int((clip_west  - item_west)  / lon_span * w)
    x1 = int((clip_east  - item_west)  / lon_span * w)
    y0 = int((item_north - clip_north) / lat_span * h)
    y1 = int((item_north - clip_south) / lat_span * h)

    # Clamp to image bounds
    x0 = max(0, min(x0, w - 1))
    x1 = max(0, min(x1, w))
    y0 = max(0, min(y0, h - 1))
    y1 = max(0, min(y1, h))

    if x1 <= x0 or y1 <= y0:
        return arr

    return arr[y0:y1, x0:x1]


def _crop_to_valid(arr: np.ndarray, threshold: int = 5,
                    skip_ratio_guard: bool = False) -> np.ndarray:
    """
    Crop a numpy image array to the bounding box of non-black pixels.
    Removes the large black nodata borders that satellite tiles contain.

    Safety guard: if the crop would produce an extreme aspect ratio (more than
    3:1 in either direction), the original uncropped array is returned instead.
    This prevents thin-strip thumbnails when a tile only marginally overlaps
    the search area and valid pixels are confined to one edge.

    skip_ratio_guard: set True when cropping a bbox-clipped result. In that
    case we always crop (the caller applies _pad_to_ratio afterward to handle
    any resulting extreme aspect ratio).
    """
    valid = arr.max(axis=2) > threshold
    if not valid.any():
        return arr
    rows = np.where(valid.any(axis=1))[0]
    cols = np.where(valid.any(axis=0))[0]
    cropped = arr[rows[0]:rows[-1] + 1, cols[0]:cols[-1] + 1]

    h, w = cropped.shape[:2]
    if h == 0 or w == 0:
        return arr

    # Reject degenerate crops — return original if aspect ratio exceeds 3:1.
    # Skipped when caller handles aspect ratio itself (e.g. after _clip_to_bbox).
    if not skip_ratio_guard:
        ratio = max(w / h, h / w)
        if ratio > 3.0:
            return arr

    return cropped
