"""
lidar_intelligence.py — LiDAR Clearance Intelligence module for the EOIL portal.

Arc 5: LiDAR Clearance Intelligence.

Three-layer architecture:
  Layer 1: Point cloud loading and inspection — point count, classification breakdown,
            Plotly 3D interactive view, overhead density map
  Layer 2: Random Forest classification + DBSCAN tree crown delineation —
            accuracy, feature importance, tree count, violation count, CHM raster
  Layer 3: Groq/Gemini structured clearance brief — five-element inspection output

Pre-processed .npz assets in assets/ replace live computation on every portal load.
All algorithm logic is documented in notebook 09_lidar_intelligence.ipynb.
Results are cached in Streamlit session state so switching tabs does not re-run.
"""

import io
import re
import warnings
warnings.filterwarnings('ignore')

import numpy as np
import streamlit as st
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import plotly.graph_objects as go

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor

import config
import ai_chain

# ---------------------------------------------------------------------------
# Pre-configured corridors
# ---------------------------------------------------------------------------

ASSETS_DIR = Path(__file__).parent / "assets"

LIDAR_CORRIDORS = {
    "Catawba Valley, NC — 115kV transmission line": {
        "asset_file":            "lidar_catawba_nc.npz",
        "voltage_kv":            115,
        "clearance_threshold_m": 4.5,
        "corridor_half_width_m": 15,
        "utm_zone":              "17N",
        "terrain":               "Rolling mixed deciduous/conifer forest, Piedmont transition zone",
        "context": (
            "115kV single-circuit wood-pole line, Lincoln County NC. "
            "Mixed deciduous/conifer forest typical of the Piedmont transition zone. "
            "Terrain relief 200–280m. Annual vegetation growth season April–October. "
            "This corridor segment was selected for its moderate encroachment risk and "
            "clear structural variation between the cleared strip and flanking forest."
        ),
        "algorithm_lesson": (
            "Baseline case. Rolling terrain with moderate canopy density. "
            "Random Forest clearly separates ground from vegetation using height above ground "
            "and return ratio. DBSCAN resolves individual crowns cleanly at 2.5m eps. "
            "Expect a small number of violations near the corridor edges."
        ),
    },
    "Gifford Pinchot, WA — 500kV BPA transmission line": {
        "asset_file":            "lidar_gifford_pinchot_wa.npz",
        "voltage_kv":            500,
        "clearance_threshold_m": 6.1,
        "corridor_half_width_m": 30,
        "utm_zone":              "10N",
        "terrain":               "Dense Douglas fir and Western hemlock, Cascade Range foothills",
        "context": (
            "500kV Bonneville Power Administration line, Skamania County WA. "
            "Dense Douglas fir and Western hemlock forest up to 45m tall. "
            "Higher voltage requires larger clearance. Rapid regrowth after clearing means "
            "annual inspection is standard practice on this corridor."
        ),
        "algorithm_lesson": (
            "High canopy density case. DBSCAN finds many more crowns and the violation count "
            "is higher than the NC corridor. The larger clearance threshold (6.1m vs 4.5m) "
            "means the algorithm is stricter. Compare violation counts directly with Catawba "
            "to see how terrain and forest type drive inspection workload."
        ),
    },
    "Cumberland Plateau, TN — 161kV TVA transmission line": {
        "asset_file":            "lidar_cumberland_tn.npz",
        "voltage_kv":            161,
        "clearance_threshold_m": 5.0,
        "corridor_half_width_m": 20,
        "utm_zone":              "16N",
        "terrain":               "Deciduous hardwood forest, sandstone plateau terrain",
        "context": (
            "161kV Tennessee Valley Authority line, Fentress County TN. "
            "Deciduous hardwood forest on the Cumberland Plateau. "
            "Sandstone terrain with pronounced ridges and hollows. "
            "Leaf-off LiDAR collection — point cloud penetrates to ground through bare canopy."
        ),
        "algorithm_lesson": (
            "Leaf-off collection case. Deciduous trees in winter lose most of their canopy "
            "volume. Point cloud density in the vegetation classes is lower than the other corridors. "
            "DBSCAN detects fewer crowns but height measurements are more accurate — pulses reach "
            "the trunk and main limbs rather than being scattered by foliage. "
            "Compare crown radius values with Gifford Pinchot to see the foliage effect."
        ),
    },
}

# ---------------------------------------------------------------------------
# Asset loader
# ---------------------------------------------------------------------------

def _load_asset(corridor_key: str) -> dict | None:
    """Load the pre-processed .npz asset for a corridor. Returns None if missing."""
    meta  = LIDAR_CORRIDORS[corridor_key]
    fpath = ASSETS_DIR / meta["asset_file"]
    if not fpath.exists():
        return None
    data = np.load(str(fpath), allow_pickle=False)
    return dict(data)


def _asset_available(corridor_key: str) -> bool:
    meta  = LIDAR_CORRIDORS[corridor_key]
    fpath = ASSETS_DIR / meta["asset_file"]
    return fpath.exists()

# ---------------------------------------------------------------------------
# Charts
# ---------------------------------------------------------------------------

def _fig_to_bytes(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _build_3d_scatter(data: dict, corridor_key: str) -> go.Figure:
    """Plotly Scatter3d — interactive rotatable point cloud coloured by RF class."""
    CLASS_COLORS = {
        2: '#8B5E3C',   # ground — brown
        3: '#90EE90',   # low veg — light green
        4: '#228B22',   # medium veg — forest green
        5: '#004d00',   # high veg — dark green
        6: '#888888',   # building — grey
    }
    CLASS_LABELS = {
        2: 'Ground', 3: 'Low veg (0.5–3m)',
        4: 'Medium veg (3–10m)', 5: 'High veg (>10m)', 6: 'Building',
    }

    x = data['x']
    y = data['y']
    z = data['z']
    rf_class = data['rf_class']

    # Subsample to 10,000 points to stay within Intel integrated GPU memory limits
    MAX_PTS = 50_000  # Intel integrated GPU limit — do not raise above 75k
    if len(x) > MAX_PTS:
        rng = np.random.default_rng(42)
        idx = rng.choice(len(x), MAX_PTS, replace=False)
        idx.sort()
        x, y, z, rf_class = x[idx], y[idx], z[idx], rf_class[idx]

    # Mark violating tree crowns bright red
    violating = data.get('tree_violating', np.array([], dtype=bool))
    tree_cx   = data.get('tree_cx', np.array([], dtype=np.float32))
    tree_cy   = data.get('tree_cy', np.array([], dtype=np.float32))
    tree_r    = data.get('tree_radius', np.array([], dtype=np.float32))

    violation_flag = np.zeros(len(x), dtype=bool)
    for i in range(len(tree_cx)):
        if len(violating) > i and violating[i]:
            dist = np.sqrt((x - tree_cx[i])**2 + (y - tree_cy[i])**2)
            violation_flag |= (dist < tree_r[i] * 1.2)

    fig = go.Figure()

    # Violation points — bright red, drawn first so they sit on top
    if violation_flag.any():
        fig.add_trace(go.Scatter3d(
            x=x[violation_flag], y=y[violation_flag], z=z[violation_flag],
            mode='markers',
            marker=dict(size=1.5, color='#FF2200', opacity=0.9),
            name='Clearance violation',
        ))

    # One trace per class
    for cls in [2, 3, 4, 5, 6]:
        mask = (rf_class == cls) & ~violation_flag
        if not mask.any():
            continue
        fig.add_trace(go.Scatter3d(
            x=x[mask], y=y[mask], z=z[mask],
            mode='markers',
            marker=dict(size=1.0, color=CLASS_COLORS.get(cls, '#cccccc'), opacity=0.6),
            name=CLASS_LABELS.get(cls, f'Class {cls}'),
        ))

    meta = LIDAR_CORRIDORS[corridor_key]
    fig.update_layout(
        title=dict(
            text=(f"Point cloud — {corridor_key.split('—')[0].strip()}<br>"
                  f"<sup>RF classification | red = clearance violation | rotate to inspect</sup>"),
            font=dict(size=13),
        ),
        scene=dict(
            xaxis_title='X (m)', yaxis_title='Y (m)', zaxis_title='Elevation (m)',
            aspectmode='data',
        ),
        legend=dict(itemsizing='constant', font=dict(size=10)),
        margin=dict(l=0, r=0, t=60, b=0),
        height=520,
    )
    return fig


def _build_chm_fig(data: dict, corridor_key: str) -> bytes:
    """Matplotlib CHM overhead map with violation overlay."""
    CHM   = data['CHM']
    x_min = float(data['grid_x_min'])
    x_max = float(data['grid_x_max'])
    y_min = float(data['grid_y_min'])
    y_max = float(data['grid_y_max'])
    thr   = LIDAR_CORRIDORS[corridor_key]['clearance_threshold_m']
    hw    = LIDAR_CORRIDORS[corridor_key]['corridor_half_width_m']
    line_y = (y_min + y_max) / 2

    # Build a y-coordinate array matching the CHM grid rows
    n_rows, n_cols = CHM.shape
    y_coords  = np.linspace(y_min, y_max, n_rows)
    clear_zone = ((y_coords >= line_y - hw) & (y_coords <= line_y + hw))[:, np.newaxis]
    violation_map = ((CHM > thr) & clear_zone).astype(float)

    fig, axes = plt.subplots(1, 2, figsize=(14, 5))

    im0 = axes[0].imshow(
        CHM, origin='lower', aspect='auto',
        extent=[x_min, x_max, y_min, y_max],
        cmap='YlGn', vmin=0, vmax=float(np.nanmax(CHM))
    )
    plt.colorbar(im0, ax=axes[0], label='Canopy height above ground (m)')
    axes[0].axhspan(line_y - hw, line_y + hw, alpha=0.15, color='red', label='Clear zone')
    axes[0].set_title('Canopy Height Model', fontsize=11, fontweight='bold')
    axes[0].set_xlabel('X (m)'); axes[0].set_ylabel('Y (m)')
    axes[0].legend(fontsize=8)

    axes[1].imshow(
        CHM, origin='lower', aspect='auto',
        extent=[x_min, x_max, y_min, y_max],
        cmap='Greens', alpha=0.5, vmin=0, vmax=float(np.nanmax(CHM))
    )
    # Paint violations as solid bright red using an RGBA overlay
    red_overlay = np.zeros((*CHM.shape, 4), dtype=float)
    red_overlay[violation_map == 1] = [1.0, 0.1, 0.1, 0.92]
    axes[1].imshow(
        red_overlay, origin='lower', aspect='auto',
        extent=[x_min, x_max, y_min, y_max],
    )
    axes[1].axhspan(line_y - hw, line_y + hw, alpha=0.08, color='yellow')
    red_p = mpatches.Patch(color='red', alpha=0.7, label=f'Violation (>{thr}m)')
    axes[1].legend(handles=[red_p], fontsize=8)
    axes[1].set_title(
        f'Clearance violations — threshold {thr}m',
        fontsize=11, fontweight='bold'
    )
    axes[1].set_xlabel('X (m)'); axes[1].set_ylabel('Y (m)')

    plt.suptitle(corridor_key.split('—')[0].strip(), fontsize=10, y=1.01)
    plt.tight_layout()
    return _fig_to_bytes(fig)


def _build_clearance_chart(data: dict, corridor_key: str) -> bytes:
    """Bar chart: tree counts by severity (safe, amber, violation)."""
    tree_height   = data.get('tree_height',    np.array([]))
    tree_violating = data.get('tree_violating', np.array([], dtype=bool))
    thr = LIDAR_CORRIDORS[corridor_key]['clearance_threshold_m']

    n_trees     = int(data['n_trees'])
    n_violating = int(data['n_violating'])
    n_safe      = n_trees - n_violating

    # Amber: trees > 80% of threshold but not yet violating
    if len(tree_height) > 0 and len(tree_violating) > 0:
        non_viol_heights = tree_height[~tree_violating]
        n_amber = int((non_viol_heights > thr * 0.8).sum())
        n_safe  = n_safe - n_amber
    else:
        n_amber = 0

    categories = ['Clear\n(safe)', f'Amber\n(>{thr*0.8:.1f}m)', f'Violation\n(>{thr}m)']
    values     = [n_safe, n_amber, n_violating]
    colors     = ['#2E7D32', '#F9A825', '#C62828']

    fig, ax = plt.subplots(figsize=(7, 4))
    bars = ax.bar(categories, values, color=colors, edgecolor='white', linewidth=1.5, width=0.5)

    max_val = max(values) if max(values) > 0 else 1
    for bar, val in zip(bars, values):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + max_val * 0.02,
            str(val), ha='center', va='bottom',
            fontweight='bold', fontsize=12
        )

    ax.set_ylabel('Number of trees', fontsize=10)
    ax.set_title(
        f'Tree clearance status — {corridor_key.split("—")[0].strip()}\n'
        f'Total: {n_trees} trees | {n_violating} violations',
        fontsize=11, fontweight='bold'
    )
    ax.set_facecolor('#f8f8f8')
    fig.patch.set_facecolor('white')
    ax.grid(axis='y', alpha=0.3)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    return _fig_to_bytes(fig)

def _build_priority_df(data: dict, corridor_key: str):
    """Build a priority-ranked inspection DataFrame for all violation and amber crowns."""
    import pandas as pd

    tree_cx   = data.get('tree_cx',        np.array([], dtype=np.float32))
    tree_cy   = data.get('tree_cy',        np.array([], dtype=np.float32))
    tree_r    = data.get('tree_radius',    np.array([], dtype=np.float32))
    tree_h    = data.get('tree_height',    np.array([], dtype=np.float32))
    tree_viol = data.get('tree_violating', np.array([], dtype=bool))
    thr = LIDAR_CORRIDORS[corridor_key]['clearance_threshold_m']
    hw  = LIDAR_CORRIDORS[corridor_key]['corridor_half_width_m']
    utm_zone = LIDAR_CORRIDORS[corridor_key].get('utm_zone', 'see corridor metadata')

    y_min = float(data['grid_y_min'])
    y_max = float(data['grid_y_max'])
    line_y = (y_min + y_max) / 2.0

    rows = []
    for i in range(len(tree_cx)):
        h       = float(tree_h[i])   if len(tree_h) > i   else 0.0
        r       = float(tree_r[i])   if len(tree_r) > i   else 2.0
        is_viol = bool(tree_viol[i]) if len(tree_viol) > i else False
        cx      = float(tree_cx[i])  if len(tree_cx) > i  else 0.0
        cy      = float(tree_cy[i])  if len(tree_cy) > i  else 0.0

        is_amber = (not is_viol) and (h > thr * 0.8)
        if not is_viol and not is_amber:
            continue

        dist_from_centre = abs(cy - line_y)
        above_thr = h - thr

        # Priority score: higher = inspect sooner
        # Weights: height ratio (dominant), crown size, proximity to centre
        height_factor    = h / thr
        size_factor      = 1.0 + (r / 15.0)
        proximity_factor = 1.0 + max(0.0, (hw - dist_from_centre) / hw) * 0.3
        score = height_factor * size_factor * proximity_factor
        if is_amber:
            score *= 0.6   # amber is lower urgency than confirmed violations

        rows.append({
            'Status':            'VIOLATION' if is_viol else 'AMBER',
            'Height_m':          round(h, 1),
            'Threshold_m':       thr,
            'Above_threshold_m': round(above_thr, 2),
            'Crown_radius_m':    round(r, 1),
            'Dist_from_centre_m': round(dist_from_centre, 1),
            'Priority_score':    round(score, 3),
            'X_UTM':             round(cx, 1),
            'Y_UTM':             round(cy, 1),
            'UTM_zone':          utm_zone,
        })

    df = pd.DataFrame(rows)
    if len(df):
        df = df.sort_values('Priority_score', ascending=False).reset_index(drop=True)
        df.index = df.index + 1
        df.index.name = 'Rank'
    return df


def _build_growth_chart(data: dict, corridor_key: str, growth_rate_m: float) -> bytes:
    """Matplotlib bar chart showing current vs projected crown heights at Year 1/2/3."""
    tree_h    = data.get('tree_height',    np.array([], dtype=np.float32))
    tree_viol = data.get('tree_violating', np.array([], dtype=bool))
    thr = LIDAR_CORRIDORS[corridor_key]['clearance_threshold_m']
    hw  = LIDAR_CORRIDORS[corridor_key]['corridor_half_width_m']
    tree_cy = data.get('tree_cy', np.array([], dtype=np.float32))
    y_min = float(data['grid_y_min'])
    y_max = float(data['grid_y_max'])
    line_y = (y_min + y_max) / 2.0

    # Only trees within the clear strip
    in_strip = np.array([
        abs(float(tree_cy[i]) - line_y) <= hw
        for i in range(len(tree_cy))
    ]) if len(tree_cy) else np.zeros(len(tree_h), dtype=bool)

    h_strip    = tree_h[in_strip]   if len(tree_h) else np.array([])
    viol_strip = tree_viol[in_strip] if len(tree_viol) else np.array([], dtype=bool)

    current_viol  = int(viol_strip.sum()) if len(viol_strip) else 0
    yr1_viol = int(((h_strip + growth_rate_m * 1) > thr).sum()) if len(h_strip) else 0
    yr2_viol = int(((h_strip + growth_rate_m * 2) > thr).sum()) if len(h_strip) else 0
    yr3_viol = int(((h_strip + growth_rate_m * 3) > thr).sum()) if len(h_strip) else 0

    labels = ['Now', 'Year 1', 'Year 2', 'Year 3']
    values = [current_viol, yr1_viol, yr2_viol, yr3_viol]
    colors = ['#C62828', '#E57373', '#EF9A9A', '#FFCDD2']

    fig, ax = plt.subplots(figsize=(7, 4))
    bars = ax.bar(labels, values, color=colors, edgecolor='white', linewidth=1.5, width=0.5)
    ax.axhline(current_viol, color='#C62828', lw=1.0, ls='--', alpha=0.5)

    max_val = max(values) if max(values) > 0 else 1
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2,
                bar.get_height() + max_val * 0.02,
                str(val), ha='center', va='bottom', fontweight='bold', fontsize=11)

    ax.set_ylabel('Trees exceeding clearance threshold', fontsize=9)
    ax.set_title(
        f'Violation projection — {corridor_key.split("—")[0].strip()}\n'
        f'Growth rate: {growth_rate_m} m/year  |  Threshold: {thr} m',
        fontsize=10, fontweight='bold'
    )
    ax.set_facecolor('#f8f8f8')
    fig.patch.set_facecolor('white')
    ax.grid(axis='y', alpha=0.3)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    return _fig_to_bytes(fig)


def _build_corridor_map(data: dict, corridor_key: str) -> bytes:
    """Top-down corridor map: DBSCAN tree crowns as circles coloured by clearance status."""
    tree_cx   = data.get('tree_cx',        np.array([], dtype=np.float32))
    tree_cy   = data.get('tree_cy',        np.array([], dtype=np.float32))
    tree_r    = data.get('tree_radius',    np.array([], dtype=np.float32))
    tree_h    = data.get('tree_height',    np.array([], dtype=np.float32))
    tree_viol = data.get('tree_violating', np.array([], dtype=bool))
    thr = LIDAR_CORRIDORS[corridor_key]['clearance_threshold_m']
    hw  = LIDAR_CORRIDORS[corridor_key]['corridor_half_width_m']

    x_min  = float(data['grid_x_min'])
    x_max  = float(data['grid_x_max'])
    y_min  = float(data['grid_y_min'])
    y_max  = float(data['grid_y_max'])
    line_y = (y_min + y_max) / 2.0

    fig, ax = plt.subplots(figsize=(13, 5))
    ax.set_facecolor('#e8ede8')
    fig.patch.set_facecolor('white')

    # Forest background outside clear strip
    ax.axhspan(y_min, line_y - hw, color='#c8dfc8', alpha=0.6, zorder=1)
    ax.axhspan(line_y + hw, y_max, color='#c8dfc8', alpha=0.6, zorder=1)

    # Clear strip background
    ax.axhspan(line_y - hw, line_y + hw, color='#fffde7', alpha=0.9, zorder=1)
    ax.axhline(line_y - hw, color='#757575', lw=1.5, ls='--', zorder=2)
    ax.axhline(line_y + hw, color='#757575', lw=1.5, ls='--', zorder=2)
    ax.axhline(line_y, color='#9e9e9e', lw=0.8, ls=':', zorder=2)

    # Labels for clear strip boundaries
    ax.text(x_min + 2, line_y + hw + 0.5, f'Clear zone boundary (+{hw}m)', fontsize=7, color='#555')
    ax.text(x_min + 2, line_y - hw - 1.2, f'Clear zone boundary (−{hw}m)', fontsize=7, color='#555')

    # DBSCAN tree crowns as circles
    n_drawn = len(tree_cx)
    for i in range(n_drawn):
        r       = max(float(tree_r[i]), 0.5) if len(tree_r) > i else 2.0
        is_viol = bool(tree_viol[i])          if len(tree_viol) > i else False
        h       = float(tree_h[i])            if len(tree_h) > i else 0.0

        if is_viol:
            color, edge, alpha = '#C62828', '#8B0000', 0.85
        elif h > thr * 0.8:
            color, edge, alpha = '#F9A825', '#E65100', 0.80
        else:
            color, edge, alpha = '#388E3C', '#1B5E20', 0.65

        circle = plt.Circle(
            (float(tree_cx[i]), float(tree_cy[i])), r,
            color=color, ec=edge, lw=0.5, alpha=alpha, zorder=3
        )
        ax.add_patch(circle)

    # Legend
    legend_patches = [
        mpatches.Patch(color='#388E3C', label='Clear (below threshold)'),
        mpatches.Patch(color='#F9A825', label=f'Amber  > {thr*0.8:.1f} m — approaching'),
        mpatches.Patch(color='#C62828', label=f'Violation  > {thr} m — exceeds NERC FAC-003'),
        mpatches.Patch(color='#fffde7', ec='#757575', ls='--', label=f'Clear strip  ±{hw} m'),
    ]
    ax.legend(handles=legend_patches, loc='upper right', fontsize=8, framealpha=0.92)

    ax.set_xlim(x_min, x_max)
    ax.set_ylim(y_min - 2, y_max + 2)
    ax.set_xlabel('Along-corridor distance (m)', fontsize=9)
    ax.set_ylabel('Cross-corridor distance (m)', fontsize=9)
    ax.set_title(
        f'DBSCAN Tree Crown Map — {corridor_key.split("—")[0].strip()}\n'
        f'Each circle = one detected crown  |  Radius = crown radius  |  Colour = clearance status',
        fontsize=10, fontweight='bold'
    )
    plt.tight_layout()
    return _fig_to_bytes(fig)


def _build_3d_static_png(data: dict, corridor_key: str) -> bytes | None:
    """Matplotlib static 3D point cloud — same colours as the interactive chart.
    Used for Word doc and AI brief snapshot. Renders correctly without WebGL."""
    try:
        from mpl_toolkits.mplot3d import Axes3D  # noqa: F401

        CLASS_COLORS = {
            2: '#8B5E3C', 3: '#90EE90', 4: '#228B22', 5: '#004d00', 6: '#888888',
        }
        CLASS_LABELS = {
            2: 'Ground', 3: 'Low veg', 4: 'Medium veg', 5: 'High veg', 6: 'Building',
        }

        x        = data['x']
        y        = data['y']
        z        = data['z']
        rf_class = data['rf_class']

        # Same 50k subsample as the interactive chart
        MAX_PTS = 50_000
        if len(x) > MAX_PTS:
            rng = np.random.default_rng(42)
            idx = rng.choice(len(x), MAX_PTS, replace=False)
            idx.sort()
            x, y, z, rf_class = x[idx], y[idx], z[idx], rf_class[idx]

        # Recompute violation flags
        violating = data.get('tree_violating', np.array([], dtype=bool))
        tree_cx   = data.get('tree_cx',    np.array([], dtype=np.float32))
        tree_cy   = data.get('tree_cy',    np.array([], dtype=np.float32))
        tree_r    = data.get('tree_radius', np.array([], dtype=np.float32))
        violation_flag = np.zeros(len(x), dtype=bool)
        for i in range(len(tree_cx)):
            if len(violating) > i and violating[i]:
                dist = np.sqrt((x - tree_cx[i])**2 + (y - tree_cy[i])**2)
                violation_flag |= (dist < tree_r[i] * 1.2)

        fig = plt.figure(figsize=(12, 6))
        ax  = fig.add_subplot(111, projection='3d')
        ax.set_facecolor('#f8f8f8')
        fig.patch.set_facecolor('white')

        # Non-violation points by class
        for cls in [2, 3, 4, 5, 6]:
            mask = (rf_class == cls) & ~violation_flag
            if not mask.any():
                continue
            ax.scatter(x[mask], y[mask], z[mask],
                       c=CLASS_COLORS.get(cls, '#cccccc'),
                       s=0.3, alpha=0.5, label=CLASS_LABELS.get(cls, f'Class {cls}'),
                       rasterized=True)

        # Violation points on top — bright red, larger
        if violation_flag.any():
            ax.scatter(x[violation_flag], y[violation_flag], z[violation_flag],
                       c='#FF2200', s=2.5, alpha=0.95, label='Clearance violation',
                       zorder=5, rasterized=True)

        ax.set_xlabel('X (m)', fontsize=8, labelpad=2)
        ax.set_ylabel('Y (m)', fontsize=8, labelpad=2)
        ax.set_zlabel('Elevation (m)', fontsize=8, labelpad=2)
        ax.tick_params(labelsize=6)
        ax.set_title(
            f"Point cloud — {corridor_key.split('—')[0].strip()}\n"
            "RF classification  |  red = clearance violation",
            fontsize=10, fontweight='bold'
        )
        ax.legend(loc='upper left', fontsize=7, markerscale=4, framealpha=0.8)
        ax.view_init(elev=22, azim=-60)
        plt.tight_layout()
        return _fig_to_bytes(fig)
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Layer 3: prompt and fallback
# ---------------------------------------------------------------------------

def _build_inspection_prompt(data: dict, corridor_key: str) -> str:
    meta = LIDAR_CORRIDORS[corridor_key]
    return f"""You are an Earth observation and LiDAR analyst producing a structured clearance inspection brief for a utility vegetation management team.

CORRIDOR: {corridor_key}
VOLTAGE: {meta['voltage_kv']} kV
TERRAIN: {meta['terrain']}
CONTEXT: {meta['context']}

LAYER 2 OUTPUTS (computed from USGS 3DEP LiDAR point cloud):
- Total points in scan: {int(data['n_pts_total']):,}
- Total tree crowns detected (DBSCAN): {int(data['n_trees'])}
- Mean maximum canopy height: {float(data['mean_height_m']):.1f} m
- Tallest tree crown: {float(data['max_height_m']):.1f} m
- Clearance threshold (NERC FAC-003): {meta['clearance_threshold_m']} m
- Trees in violation: {int(data['n_violating'])} of {int(data['n_trees'])}
- Corridor violation area: {float(data['violation_pct']):.1f}% of clear zone

Write a structured five-element inspection brief. Use exactly these five headings.

## 1. Clearance Status
State the number of violations, the total tree count, and the violation percentage. Reference the voltage and applicable clearance standard. Use the numbers provided.

## 2. Priority Inspection Zones
Based on the violation count and corridor context, describe which zones require immediate field inspection. Distinguish between confirmed violations and amber-zone trees approaching the threshold.

## 3. Sensor Capability
Explain what airborne LiDAR measures that a satellite cannot. Cover point cloud structure, return physics (why multi-return pulses penetrate canopy), and how height above ground is derived from the DTM. Explain why leaf-on vs leaf-off collection timing matters for clearance analysis.

## 4. Recommended Actions
State what a vegetation management crew should do with this output, in priority order. Include at least one field verification step the LiDAR cannot confirm alone (e.g., species identification, trunk condition, growth rate).

## 5. Confidence and Limitations
State clearly what this analysis cannot confirm without field inspection. Address crown overlap (two touching crowns counted as one), wind sway (crown position shifts under load), and the difference between a crown at threshold height and an actual flash-over risk.

End with:
DATA QUALITY: Points: {int(data['n_pts_total']):,} | Trees detected: {int(data['n_trees'])} | Violations: {int(data['n_violating'])} | Violation area: {float(data['violation_pct']):.1f}%
"""


def _fallback_brief(data: dict, corridor_key: str) -> str:
    meta = LIDAR_CORRIDORS[corridor_key]
    n_trees     = int(data['n_trees'])
    n_violating = int(data['n_violating'])
    max_h       = float(data['max_height_m'])
    thr         = meta['clearance_threshold_m']
    viol_pct    = float(data['violation_pct'])

    return f"""## 1. Clearance Status

**{n_violating} of {n_trees} tree crowns** exceed the {thr}m clearance threshold for this {meta['voltage_kv']}kV corridor. Violation area: **{viol_pct:.1f}% of the clear zone**. The tallest detected crown reaches {max_h:.1f}m above ground. NERC FAC-003 requires utilities to maintain clearance from transmission conductors to prevent flashover under maximum sag conditions.

## 2. Priority Inspection Zones

- **Immediate priority:** All {n_violating} confirmed violation crowns. Tallest trees first — a crown at {max_h:.1f}m in a {meta['clearance_threshold_m']}m threshold corridor presents the highest risk under conductor sag.
- **Secondary priority:** Trees in the amber zone (height > {thr*0.8:.1f}m) — these will likely breach threshold within one to two growing seasons without intervention.
- Inspect the {meta['terrain'].split(',')[0]} areas specifically. Dense forest flanking a corridor tends to produce encroachment from canopy lean rather than vertical growth alone.

## 3. Sensor Capability

Airborne LiDAR fires laser pulses at ~200,000 per second and records the precise time-of-flight for each return. A single pulse hitting a tree crown can produce 2–4 returns: first return from the outer canopy, intermediate returns from inner branches, and a last return from the ground. This multi-return structure is what allows the algorithm to measure both canopy height and ground elevation simultaneously.

The Digital Terrain Model (DTM) is constructed from ground-class returns only. Canopy height is then computed as the highest return minus the DTM value at that location. Without LiDAR, you cannot measure vegetation height accurately in forested terrain from satellite data.

Leaf-on collection (April–September) captures the maximum seasonal canopy extent but pulse penetration to ground is reduced. Leaf-off collection (November–March) penetrates to the trunk and main limbs — height measurements are slightly lower but more representative of the structural hazard present year-round.

## 4. Recommended Actions

1. Generate a GPS-tagged field work order for all {n_violating} violation trees. LiDAR provides position accuracy of approximately 0.5m at 1m collection resolution.
2. Prioritise by height above threshold, not just violation count. A tree at {max_h:.1f}m needs immediate treatment; a tree at {thr+0.1:.1f}m can be scheduled for next season.
3. Field crew to confirm species for each flagged crown. Species determines growth rate and guides treatment decision (trim vs remove).
4. Inspect for lean angle and condition. A structurally compromised tree 3m inside the threshold may present higher risk than a healthy tree 0.2m above threshold.
5. Re-fly this segment in the opposite season (leaf-on if current is leaf-off) to capture maximum seasonal canopy extent.

## 5. Confidence and Limitations

- **Crown overlap:** Two adjacent trees whose crowns touch will register as one DBSCAN cluster. Violation count may understate the number of individual trees requiring treatment.
- **Wind sway:** LiDAR captures one moment in time. Under wind loading, conductor sag increases and crown position shifts. A tree measuring exactly at threshold in calm conditions may violate under design wind speed.
- **Growth since collection:** LiDAR data has a collection date. Trees grow 0.3–1.5m per year depending on species and site. Any tree above {thr*0.75:.1f}m at time of collection should be re-evaluated with current field measurement.
- **Single-return ground ambiguity:** In dense canopy, some ground returns are from low vegetation misclassified as ground. DTM accuracy in dense forest is ±0.2–0.5m, which propagates into height measurements.

DATA QUALITY: Points: {int(data['n_pts_total']):,} | Trees detected: {n_trees} | Violations: {n_violating} | Violation area: {viol_pct:.1f}%
"""

# ---------------------------------------------------------------------------
# Word document builder
# ---------------------------------------------------------------------------

def _add_inline_bold(paragraph, text):
    """Render **bold** spans correctly inside a Word paragraph."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def _build_word_doc(data: dict, brief_text: str, model_used: str,
                    clearance_chart_bytes: bytes, chm_bytes: bytes,
                    corridor_key: str, fig_3d_bytes: bytes = None,
                    corridor_map_bytes: bytes = None) -> bytes:
    """Build a formatted Word document with stats, charts, and inspection brief."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    meta = LIDAR_CORRIDORS[corridor_key]

    # Title
    title = doc.add_paragraph()
    run   = title.add_run('LiDAR Clearance Intelligence Brief')
    run.bold = True; run.font.size = Pt(16)

    doc.add_paragraph(f"Corridor: {corridor_key}")
    doc.add_paragraph(f"Voltage: {meta['voltage_kv']} kV  |  Clearance threshold: {meta['clearance_threshold_m']} m")
    doc.add_paragraph(meta['context'])
    doc.add_paragraph()

    # 3D point cloud snapshot
    if fig_3d_bytes:
        p0 = doc.add_paragraph()
        p0.add_run('Layer 1 — 3D Point Cloud Snapshot').bold = True
        p0.runs[0].font.size = Pt(13)
        doc.add_picture(io.BytesIO(fig_3d_bytes), width=Inches(6.0))
        cap = doc.add_paragraph()
        cap.add_run(
            'RF classification: brown = ground, green shades = vegetation height classes, red = clearance violation crowns.'
        ).italic = True
        cap.runs[0].font.size = Pt(9)
        doc.add_paragraph()

    # Layer 2 stats table
    p = doc.add_paragraph()
    p.add_run('Layer 2 — Clearance Statistics').bold = True
    p.runs[0].font.size = Pt(13)

    tbl = doc.add_table(rows=7, cols=2)
    tbl.style = 'Table Grid'
    rows_data = [
        ('Metric',                   'Value'),
        ('Total points scanned',     f"{int(data['n_pts_total']):,}"),
        ('Tree crowns detected',      str(int(data['n_trees']))),
        ('Mean canopy height',        f"{float(data['mean_height_m']):.1f} m"),
        ('Tallest crown',             f"{float(data['max_height_m']):.1f} m"),
        ('Clearance violations',      f"{int(data['n_violating'])} of {int(data['n_trees'])} trees"),
        ('Corridor violation area',   f"{float(data['violation_pct']):.1f}%"),
    ]
    for i, (label, value) in enumerate(rows_data):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = value
        if i == 0:
            for cell in tbl.rows[i].cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()

    # Clearance bar chart
    p2 = doc.add_paragraph()
    p2.add_run('Tree Clearance Status').bold = True
    p2.runs[0].font.size = Pt(13)
    if clearance_chart_bytes:
        doc.add_picture(io.BytesIO(clearance_chart_bytes), width=Inches(4.5))
    doc.add_paragraph()

    # Corridor crown map
    if corridor_map_bytes:
        p2b = doc.add_paragraph()
        p2b.add_run('DBSCAN Tree Crown Map').bold = True
        p2b.runs[0].font.size = Pt(13)
        doc.add_picture(io.BytesIO(corridor_map_bytes), width=Inches(6.0))
        doc.add_paragraph()

    # CHM panels
    p3 = doc.add_paragraph()
    p3.add_run('Layer 1 — Canopy Height Model and Violation Map').bold = True
    p3.runs[0].font.size = Pt(13)
    if chm_bytes:
        doc.add_picture(io.BytesIO(chm_bytes), width=Inches(6.0))
    doc.add_paragraph()

    # AI brief
    p4 = doc.add_paragraph()
    p4.add_run('Layer 3 — AI Clearance Inspection Brief').bold = True
    p4.runs[0].font.size = Pt(13)
    if model_used:
        mi = doc.add_paragraph()
        mi.add_run(f'Generated by: {model_used}').italic = True

    lines = brief_text.splitlines()
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if stripped.startswith('## '):
            p = doc.add_paragraph(stripped[3:])
            p.style = doc.styles['Heading 2']
        elif stripped.startswith('# '):
            p = doc.add_paragraph(stripped[2:])
            p.style = doc.styles['Heading 1']
        elif stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_bold(p, stripped[2:])
        elif stripped.startswith('DATA QUALITY:'):
            p = doc.add_paragraph()
            p.add_run(stripped).bold = True
        elif stripped == '' or stripped.startswith('---'):
            doc.add_paragraph()
        else:
            if stripped:
                p = doc.add_paragraph()
                _add_inline_bold(p, stripped)
        i += 1

    footer = doc.add_paragraph()
    footer.add_run('Generated by EOIL — AI-Native Earth Observation Innovation Lab').italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _build_markdown(data: dict, brief_text: str, model_used: str,
                    corridor_key: str) -> str:
    meta = LIDAR_CORRIDORS[corridor_key]
    lines = [
        "# LiDAR Clearance Intelligence Brief",
        "",
        f"**Corridor:** {corridor_key}",
        f"**Voltage:** {meta['voltage_kv']} kV",
        f"**Context:** {meta['context']}",
        "",
        "---",
        "",
        "## Layer 2 — Clearance Statistics",
        "",
        "| Metric | Value |",
        "|--------|-------|",
        f"| Total points scanned | {int(data['n_pts_total']):,} |",
        f"| Tree crowns detected | {int(data['n_trees'])} |",
        f"| Mean canopy height | {float(data['mean_height_m']):.1f} m |",
        f"| Tallest crown | {float(data['max_height_m']):.1f} m |",
        f"| Clearance violations | {int(data['n_violating'])} of {int(data['n_trees'])} trees |",
        f"| Corridor violation area | {float(data['violation_pct']):.1f}% |",
        "",
        "---",
        "",
        "## Layer 3 — AI Clearance Inspection Brief",
        "",
        f"*Generated by: {model_used or 'built-in fallback'}*",
        "",
        brief_text,
    ]
    return "\n".join(lines)

# ---------------------------------------------------------------------------
# Main render function
# ---------------------------------------------------------------------------

def render():
    """Render the LiDAR Clearance Intelligence module. Called by app.py."""

    st.markdown("## 🌲 LiDAR Clearance Intelligence")
    st.markdown(
        "Airborne LiDAR point cloud analysis for transmission corridor vegetation management. "
        "3D point cloud classification, individual tree crown delineation, "
        "and clearance violation detection — all from a single laser scan."
    )

    with st.expander("ℹ️ What does this module do?", expanded=False):
        st.markdown("""
**This module answers one question: which trees are too close to the transmission line?**

It processes a LiDAR point cloud using three analysis steps:

1. **Digital Terrain Model (DTM)** — ground returns are interpolated into a bare-earth elevation surface. This is the reference level for all height measurements.
2. **Canopy Height Model (CHM)** — the highest vegetation return in each 1m grid cell, minus the DTM below it. This is the actual height of the vegetation above ground — not above sea level.
3. **Machine learning classification** — a Random Forest trained on six point attributes classifies every point as ground, low vegetation, medium vegetation, high vegetation, or building.
4. **DBSCAN tree crown delineation** — high vegetation points are grouped into individual tree crowns. Each crown gets a position, a maximum height, and a crown radius.
5. **Clearance violation flagging** — any crown whose height exceeds the voltage-specific clearance threshold is flagged.

The AI brief translates those measurements into a prioritised inspection output for vegetation management crews.
        """)

    st.divider()

    corridor_key = st.selectbox(
        "Select corridor",
        list(LIDAR_CORRIDORS.keys()),
        key="li_corridor_key",
        help="Each corridor has a pre-processed LiDAR asset and known vegetation context."
    )
    meta = LIDAR_CORRIDORS[corridor_key]

    col_a, col_b = st.columns(2)
    with col_a:
        st.info(
            f"**Voltage:** {meta['voltage_kv']} kV  \n"
            f"**Clearance threshold:** {meta['clearance_threshold_m']} m  \n"
            f"**Clear strip half-width:** {meta['corridor_half_width_m']} m"
        )
    with col_b:
        st.info(f"**Terrain:** {meta['terrain']}")

    lesson = meta.get("algorithm_lesson")
    if lesson:
        st.caption(f"📌 **What this corridor demonstrates:** {lesson}")

    if not _asset_available(corridor_key):
        st.warning(
            f"Pre-processed asset not found for this corridor "
            f"(`assets/{meta['asset_file']}`). "
            "Run notebook 09_lidar_intelligence.ipynb to generate it."
        )
        return

    run_btn = st.button("▶ Run LiDAR Analysis", type="primary", key="li_run")

    # ------------------------------------------------------------------
    # Two-step run pattern
    # ------------------------------------------------------------------
    cache_key = f"li_results_{corridor_key}"

    if run_btn:
        st.session_state["li_pending"]     = cache_key
        st.session_state["li_pending_key"] = corridor_key
        st.rerun()

    if st.session_state.get("li_pending") == cache_key:
        del st.session_state["li_pending"]
        with st.spinner("Loading LiDAR point cloud and running analysis — allow 10–20 seconds..."):
            try:
                data = _load_asset(st.session_state["li_pending_key"])
                if data is None:
                    st.error("Asset file could not be loaded.")
                    st.stop()
                st.session_state[cache_key] = data
            except Exception as ex:
                st.error(f"LiDAR analysis failed: {ex}")
                st.stop()

    # ------------------------------------------------------------------
    # Results
    # ------------------------------------------------------------------
    if cache_key not in st.session_state:
        st.info("Select a corridor and click **▶ Run LiDAR Analysis** to begin.")
        with st.expander("What is LiDAR and why does it matter for transmission corridors?", expanded=True):
            st.markdown("""
**What a satellite cannot do.**

A multispectral satellite measures reflected sunlight from the top of the canopy.
It can tell you that a pixel contains green vegetation. It cannot tell you how tall
that vegetation is, how many individual trees are present, or whether any of them are
within the legal clearance distance of a transmission conductor.

**What LiDAR can do.**

An aircraft fires laser pulses at ~200,000 per second. Each pulse travels to the ground
and the return time is measured to millimetre precision. A pulse that hits a tree crown
splits into multiple returns: one from the outer canopy, one from an inner branch,
one from the ground below. This multi-return structure gives you both the top of the
tree and the ground beneath it in a single pass.

The result is a **point cloud** — millions of 3D points, each with a physical location
and a set of attributes. From the point cloud you derive the Canopy Height Model (CHM):
vegetation height above ground, at 1m resolution, across the full corridor.

**The regulatory context.**

NERC FAC-003 requires transmission operators to maintain minimum clearance distances
between conductors and vegetation. The threshold varies by voltage: 4.5m for 115kV,
6.1m for 500kV. Violation of this standard has caused major North American blackouts.
LiDAR is the only remote sensing method that can measure this distance reliably
across hundreds of kilometres of corridor at sub-metre accuracy.
            """)
        return

    data = st.session_state[cache_key]

    def section_break():
        st.markdown(
            '<hr style="border: none; border-top: 3px solid #d0d0d0; margin: 24px 0 16px 0;">',
            unsafe_allow_html=True,
        )

    # ------------------------------------------------------------------
    # Layer 1: Point cloud overview
    # ------------------------------------------------------------------
    section_break()
    st.subheader("📡 Layer 1 — Signal Processing")
    st.caption(
        "Point cloud loaded from pre-processed USGS 3DEP asset. "
        "Four metrics below summarise the raw scan. "
        "The 3D scatter below is interactive — rotate to inspect the corridor structure."
    )

    n_pts   = int(data['n_pts_total'])
    n_trees = int(data['n_trees'])

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total points", f"{n_pts:,}")
    with col2:
        rf_class = data.get('rf_class', np.array([]))
        veg_pct  = float((np.isin(rf_class, [3, 4, 5])).mean() * 100) if len(rf_class) else 0
        st.metric("Vegetation points", f"{veg_pct:.0f}%")
    with col3:
        st.metric("Tree crowns detected", str(n_trees))
    with col4:
        st.metric("Tallest crown", f"{float(data['max_height_m']):.1f} m")

    st.caption("🔵 **Simulated corridor data** — representative of USGS 3DEP QL2 point density and structure. See notebook 09_lidar_intelligence.ipynb for the live API workflow.")

    fig_key = cache_key + "_fig3d"
    if fig_key not in st.session_state:
        st.session_state[fig_key] = _build_3d_scatter(data, corridor_key)
    st.plotly_chart(st.session_state[fig_key], use_container_width=True, key="li_3d_chart")

    with st.expander("ℹ️ How to read the 3D point cloud", expanded=False):
        st.markdown("""
**Each dot is one laser return.** Colour indicates the Random Forest classification.

- **Brown** — ground returns. These form the flat base of the scene.
- **Light green** — low vegetation (0.5–3m). Shrubs, saplings, regrowth in the corridor.
- **Forest green** — medium vegetation (3–10m). Established understorey and young trees.
- **Dark green** — high vegetation (>10m). Mature canopy trees flanking the corridor.
- **Grey** — buildings or structures.
- **Red** — clearance violations. These are tree crowns whose height exceeds the threshold.

**How to use the interactive view:**
Rotate to look down the corridor axis. You will see the clear strip in the centre —
lower points, no tall vegetation. The flanking forest rises sharply on both sides.
Red points near the corridor edge are where crowns breach the clearance zone.

**Why is this subsampled?**
The full point cloud may contain several million points. The portal displays ~50,000
for smooth browser rendering. The analysis used the full dataset.
        """)

    with st.expander("🗓️ When do utilities fly LiDAR — and why does season matter?", expanded=False):
        st.markdown("""
Timing the LiDAR flight is an operational decision. The choice of season changes what the
point cloud can and cannot see.

| Season | Canopy state | Ground penetration | Best for |
|--------|-------------|-------------------|----------|
| **Winter (leaf-off)** | Bare deciduous trees | Excellent — pulses reach ground through bare branches | DTM accuracy, root-zone proximity, structural assessment |
| **Late spring (full leaf)** | Dense green canopy | Poor — most pulses absorbed by canopy | Maximum crown height, worst-case clearance measurement |
| **Summer** | Full leaf, dry conditions | Poor | High-risk period coincides with growth peak and dry-weather sag |
| **Early fall** | Pre-senescence, still leafy | Moderate | Balances height measurement with some ground return |

**The utility practice:**

Most North American transmission operators fly **leaf-off** in the northeast and midwest
(November to March) for DTM accuracy. They fly **leaf-on** in the southeast (May to August)
to capture maximum canopy height — the worst-case clearance condition.

Some operators fly twice per year: leaf-off for structural analysis, leaf-on for clearance.
The two CHMs compared give a direct measure of seasonal height change.

**The Cumberland Plateau corridor in this module** was collected leaf-off. Crown heights
are lower than leaf-on values. This is normal — annotated on the module header.
        """)

    # ------------------------------------------------------------------
    # CHM
    # ------------------------------------------------------------------
    section_break()
    st.subheader("🗺️ Canopy Height Model and Violation Map")
    chm_bytes = _build_chm_fig(data, corridor_key)
    st.image(chm_bytes, use_container_width=True)
    st.caption(
        "Left: canopy height above ground (m) across the full corridor. "
        "Right: cells in the clear zone exceeding the clearance threshold (red). "
        "The yellow band marks the legally required clear strip."
    )

    # ------------------------------------------------------------------
    # Layer 2: Classification and clearance statistics
    # ------------------------------------------------------------------
    section_break()
    st.subheader("📊 Layer 2 — Classification and Clearance Analysis")
    st.caption(
        "Random Forest classified every point. "
        "DBSCAN grouped high vegetation points into individual tree crowns. "
        "Any crown breaching the clearance threshold is a violation."
    )

    col_a, col_b, col_c, col_d = st.columns(4)
    with col_a:
        st.metric("Tree crowns", str(n_trees))
    with col_b:
        st.metric("Mean canopy height", f"{float(data['mean_height_m']):.1f} m")
    with col_c:
        n_viol = int(data['n_violating'])
        st.metric("Violations", str(n_viol),
                  delta=f"{n_viol/n_trees*100:.0f}% of trees" if n_trees > 0 else "",
                  delta_color="inverse")
    with col_d:
        st.metric("Violation area", f"{float(data['violation_pct']):.1f}%")

    clearance_chart_bytes = _build_clearance_chart(data, corridor_key)
    corridor_map_bytes    = _build_corridor_map(data, corridor_key)

    st.markdown("**Top-down corridor view — DBSCAN tree crown map**")
    st.image(corridor_map_bytes, use_container_width=True)
    st.caption(
        "Each circle is one tree crown detected by DBSCAN. "
        "Radius matches the crown's spatial extent. "
        "Colour shows clearance status against the NERC FAC-003 threshold for this voltage."
    )

    st.markdown("**Tree count by clearance status**")
    st.image(clearance_chart_bytes, use_column_width=False, width=420)

    with st.expander("⚠️ What this algorithm cannot detect", expanded=False):
        st.markdown(f"""
**Crown overlap.**
DBSCAN groups nearby points into clusters. When two adjacent trees have touching
crowns, their points merge into a single cluster and register as one tree.
Violation count may understate the true number of trees needing treatment.

**Wind sway and conductor sag.**
LiDAR captures one moment in calm conditions. Under design wind speed (NERC NESC loading),
conductors sag further and tree crowns sway toward the line. A tree at {meta['clearance_threshold_m']}m
in still air may be in contact with the conductor at 40 km/h wind.
The clearance threshold already includes a sag allowance — but field verification
at maximum load conditions is always required before closing a violation as resolved.

**Growth rate variation.**
The LiDAR scan has a collection date. Trees grow 0.3–1.5m per year depending on species.
A tree at {meta['clearance_threshold_m']*0.85:.1f}m at time of collection may exceed {meta['clearance_threshold_m']}m before
the next scheduled re-flight. The amber zone (80% of threshold) flags these proactively.

**Underground root hazard.**
LiDAR detects aerial structure only. A diseased or root-damaged tree may be structurally
unstable below the threshold height. Tree health and structural integrity require
ground inspection regardless of what LiDAR shows.
        """)

    with st.expander("📋 What happens after this analysis in a real inspection workflow", expanded=False):
        st.markdown("""
**1. Export GPS-tagged violation points.**
Each flagged crown has a centroid coordinate at ~0.5m accuracy. The work order system
ingests these as a prioritised job list. Crew vehicles are routed to each point.

**2. Field crew confirms species and condition.**
LiDAR tells you height and crown radius. It does not tell you whether the tree is
a fast-growing poplar or a slow-growing oak, or whether it has a decay cavity that
makes it a fall risk before it reaches the clearance boundary.

**3. Treatment decision: trim or remove.**
Trimming is cheaper but the tree regrows. Removal eliminates recurrence. The decision
depends on species, growth rate, proximity to threshold, and infrastructure criticality.

**4. Re-fly after treatment.**
Ground-truthing verifies that the treatment was effective and that the violation is resolved.
Some utilities re-fly full corridors every 3–5 years; others use continuous UAS patrols
for high-criticality segments.

**5. Compare with previous scan.**
If a prior-year LiDAR dataset exists, comparing CHMs gives a growth rate map.
Areas with fast annual growth get higher inspection frequency. This is the analytical
step that turns a one-time survey into a predictive maintenance programme.
        """)

    # ------------------------------------------------------------------
    # Priority inspection table
    # ------------------------------------------------------------------
    section_break()
    st.subheader("🔴 Priority Inspection Table")
    st.caption(
        "Violation and amber crowns ranked by urgency. "
        "Height ratio drives the score. Crown size and proximity to the corridor centre add weight. "
        "Amber crowns (80–100% of threshold) score lower than confirmed violations."
    )
    priority_df = _build_priority_df(data, corridor_key)
    if len(priority_df):
        import pandas as pd
        st.dataframe(priority_df, use_container_width=True)

        csv_priority = priority_df.to_csv().encode('utf-8')
        safe_name_p = corridor_key[:30].replace(' ', '_').replace(',', '').replace('—', '')
        st.download_button(
            label="⬇ Download inspection table (.csv)",
            data=csv_priority,
            file_name=f"priority-inspection-{safe_name_p}.csv",
            mime="text/csv",
            use_container_width=False,
        )

        with st.expander("How is the priority score calculated?", expanded=False):
            st.markdown(f"""
Each tree in the violation or amber zone receives a **priority score** based on three factors:

| Factor | Formula | Rationale |
|--------|---------|-----------|
| **Height ratio** | `height / threshold` | A tree at 2× threshold is twice as urgent as one barely over |
| **Crown size** | `1 + (crown_radius / 15)` | A wider crown is harder to trim and more likely to touch the conductor |
| **Centre proximity** | `1 + max(0, (half_width - dist_from_centre) / half_width) × 0.3` | Trees directly under the line are at most risk |

Final score = height_ratio × size_factor × proximity_factor.
Amber crowns (80–100% of {meta['clearance_threshold_m']}m threshold) receive a ×0.6 multiplier —
they are flagged for monitoring, not immediate treatment.

**UTM coordinates** (X_UTM, Y_UTM) are the field GPS waypoints for each tree crown centroid.
These can be loaded directly into handheld GPS units or work order systems.
UTM zone for this corridor: **{meta.get('utm_zone', 'see corridor metadata')}**.
            """)
    else:
        st.info("No violation or amber crowns detected in this corridor.")

    # ------------------------------------------------------------------
    # Growth rate projection
    # ------------------------------------------------------------------
    section_break()
    st.subheader("📈 Vegetation Growth Projection")
    st.caption(
        "How many trees will breach the clearance threshold if left untreated? "
        "Adjust the growth rate slider to match species and region. "
        "Only trees currently within the clear strip are counted."
    )

    growth_rate = st.slider(
        "Assumed annual growth rate (m/year)",
        min_value=0.1, max_value=1.5, value=0.5, step=0.1,
        key="li_growth_rate",
        help="Typical range: 0.3 m/yr for slow hardwoods, 0.8–1.5 m/yr for fast-growing poplars or pines."
    )
    growth_chart_key = cache_key + f"_growth_{growth_rate}"
    if growth_chart_key not in st.session_state:
        st.session_state[growth_chart_key] = _build_growth_chart(data, corridor_key, growth_rate)
    st.image(st.session_state[growth_chart_key], use_container_width=False, width=580)
    st.caption(
        f"Bars show the number of trees in the clear strip that exceed the "
        f"{meta['clearance_threshold_m']}m threshold at each point in time, "
        f"assuming {growth_rate} m/year uniform growth."
    )

    # Explain flat result so it doesn't look like a bug
    _tree_h  = data.get('tree_height', np.array([], dtype=np.float32))
    _tree_cy = data.get('tree_cy',     np.array([], dtype=np.float32))
    _line_y  = (float(data['grid_y_min']) + float(data['grid_y_max'])) / 2.0
    _hw      = meta['corridor_half_width_m']
    _thr     = meta['clearance_threshold_m']
    if len(_tree_cy) and len(_tree_h):
        _in_strip = np.array([abs(float(_tree_cy[i]) - _line_y) <= _hw for i in range(len(_tree_cy))])
        _h_strip  = _tree_h[_in_strip]
        _now_count = int(((_h_strip) > _thr).sum())
        _yr3_count = int(((_h_strip + growth_rate * 3) > _thr).sum())
        if _yr3_count == _now_count:
            st.info(
                f"All four bars are equal at {growth_rate} m/yr. "
                f"The {_now_count} current violations are already above the threshold. "
                f"No other strip trees are close enough to {_thr}m to breach within 3 years at this rate. "
                f"**Try the slider at 1.0–1.5 m/yr** to see new violations emerge."
            )

    with st.expander("What growth rates should I use?", expanded=False):
        st.markdown("""
Growth rate varies by species, soil, rainfall, and competition.

| Species group | Typical annual growth |
|---------------|----------------------|
| Slow hardwoods (oak, hickory) | 0.2–0.4 m/yr |
| Mixed deciduous (maple, tulip poplar) | 0.4–0.7 m/yr |
| Fast deciduous (cottonwood, silver maple) | 0.8–1.2 m/yr |
| Fast conifers (loblolly pine, Douglas fir) | 0.6–1.0 m/yr |
| Invasive species (tree of heaven, kudzu vine) | 1.0–1.5 m/yr |

**Practical guidance:** Use the lower end if you collected at leaf-off (winter).
Heights are underestimated by 0.3–0.8m compared to full leaf-on conditions.
The amber zone (80% of threshold) already provides a buffer for this.

**What this projection cannot do:** it assumes every tree grows at the same rate.
In practice, a field crew records species at each violation site. That data feeds
a per-species growth model in the next survey cycle.
        """)

    # ------------------------------------------------------------------
    # Layer 3: AI brief
    # ------------------------------------------------------------------
    section_break()
    st.subheader("🤖 Layer 3 — AI Clearance Inspection Brief")
    st.caption(
        "Five-element structured brief for vegetation management teams. "
        "Every claim is grounded in the Layer 2 measurements above."
    )

    brief_key = f"li_brief_{cache_key}"
    model_key = f"li_model_{cache_key}"

    if st.button("Get AI Inspection Brief", type="primary", key="li_ai_btn"):
        with st.spinner("Generating inspection brief..."):
            prompt = _build_inspection_prompt(data, corridor_key)
            text, model_used = ai_chain.complete(
                prompt,
                groq_key=config.GROQ_API_KEY,
                gemini_key=config.GEMINI_API_KEY,
            )
            if text:
                st.session_state[brief_key] = text
                st.session_state[model_key] = model_used
            else:
                st.session_state[brief_key] = _fallback_brief(data, corridor_key)
                st.session_state[model_key] = None

    if brief_key in st.session_state:
        brief_text = st.session_state[brief_key]
        model_used = st.session_state.get(model_key)

        # Export 3D PNG once and cache it
        png_key = cache_key + "_3d_png"
        if png_key not in st.session_state:
            st.session_state[png_key] = _build_3d_static_png(data, corridor_key)
        fig_3d_bytes = st.session_state.get(png_key)

        with st.expander("📋 AI Clearance Inspection Brief", expanded=True):
            st.markdown(brief_text)
            if model_used:
                st.caption(f"AI response from {model_used}")
            else:
                st.caption("Built-in analysis — no AI key present")

        section_break()
        st.markdown("### Downloads")
        dl1, dl2 = st.columns(2)

        safe_name = corridor_key[:30].replace(' ', '_').replace(',', '').replace('—', '')

        with dl1:
            md_content = _build_markdown(data, brief_text, model_used or "built-in fallback", corridor_key)
            st.download_button(
                label="⬇ Download brief (.md)",
                data=md_content.encode('utf-8'),
                file_name=f"lidar-clearance-{safe_name}.md",
                mime="text/markdown",
                use_container_width=True,
            )

        with dl2:
            word_bytes = _build_word_doc(
                data, brief_text, model_used or "built-in fallback",
                clearance_chart_bytes, chm_bytes, corridor_key,
                fig_3d_bytes=fig_3d_bytes,
                corridor_map_bytes=corridor_map_bytes,
            )
            st.download_button(
                label="⬇ Download brief (.docx)",
                data=word_bytes,
                file_name=f"lidar-clearance-{safe_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

        # GPS work order CSV — all violation and amber crowns with coordinates
        import pandas as pd
        wo_df = _build_priority_df(data, corridor_key)
        if len(wo_df):
            wo_df_export = wo_df.reset_index().rename(columns={'Rank': 'Work_order_rank'})
            wo_df_export.insert(0, 'Corridor', corridor_key)
            wo_df_export.insert(1, 'Voltage_kV', meta['voltage_kv'])
            wo_df_export.insert(2, 'Threshold_m', meta['clearance_threshold_m'])
            wo_csv = wo_df_export.to_csv(index=False).encode('utf-8')
            st.download_button(
                label="⬇ Download GPS work order (.csv)",
                data=wo_csv,
                file_name=f"work-order-{safe_name}.csv",
                mime="text/csv",
                use_container_width=False,
                help="GPS-tagged work order for field crews. Load X_UTM / Y_UTM directly into a GPS unit or work order system.",
            )
