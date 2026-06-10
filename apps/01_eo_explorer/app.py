"""
app.py — EOIL Portal v1.9
Streamlit layout and component wiring only. No business logic here.
All logic is imported from other modules.

v1.0  Day 1:  EO Explorer foundational app
v1.1  Day 3:  Spectral Explorer tab added
v1.2  Day 6:  Time Series Explorer module added; sidebar navigation replaces tab navigation
v1.3  Day 6:  Spectral Explorer promoted to standalone sidebar module; tabs removed
v1.4  Day 7:  SAR Explorer added; EO Explorer replaced with Welcome panel
v1.5  Day 9:  Change Detection module added (fifth sidebar entry)
v1.6  Day 10: AI Imagery Interpreter added (sixth sidebar entry)
v1.7  Day 11: Shared map picker added to all five modules; aspect ratio fix for contact sheet and large radius renders
v1.8  Day 12: Emissions Explorer added (TROPOMI CH4/NO2/CO/SO2 via GEE)
v1.9  Arc 1:  Land Cover Intelligence added (K-means + Random Forest on Sentinel-2 via Planetary Computer)
"""

import streamlit as st
from streamlit_folium import st_folium
from datetime import date, timedelta
import numpy as np

import config
import data_catalog
import sample_layers
import map_builder
import ai_assistant
import satellite_catalog
import geocoder
import spectral_explorer
import gee_timeseries
import gee_sar
import gee_change
import imagery_interpreter
import map_picker
import methane_explorer
import land_cover
import corridor_risk

# ---------------------------------------------------------------------------
# Page config
# ---------------------------------------------------------------------------
st.set_page_config(
    page_title="EOIL Portal",
    page_icon="🌍",
    layout="wide",
)

st.title("🌍 EOIL Portal")
st.caption("Earth Observation Innovation Lab — satellite data analysis and AI interpretation.")

# Pointer cursor on all interactive controls
st.markdown("""
<style>
div[data-baseweb="select"] > div,
div[data-baseweb="select"] input,
button[kind="primary"],
button[kind="secondary"],
.stButton > button,
.stDownloadButton > button,
input[type="range"] {
    cursor: pointer !important;
}
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------------------------
# Sidebar — module navigation (always visible at the top)
# ---------------------------------------------------------------------------
with st.sidebar:
    st.markdown("### Module")
    selected_module = st.radio(
        "Navigate",
        ["🏠 Welcome", "🔬 Spectral Explorer", "📈 Time Series Explorer", "🔀 Change Detection", "🔍 AI Imagery Interpreter", "📡 SAR Explorer", "🌫️ Emissions Explorer", "🌿 Land Cover Intelligence", "⚠️ Corridor Risk Intelligence"],
        label_visibility="collapsed",
    )
    st.divider()

# ---------------------------------------------------------------------------
# GEE status — check once per session and cache the result
# If GEE credentials are in Streamlit secrets, live queries are available.
# If not, the Time Series Explorer uses sample data.
# ---------------------------------------------------------------------------
if "gee_available" not in st.session_state:
    st.session_state.gee_available = gee_timeseries.init_gee()

# ---------------------------------------------------------------------------
# Helper: "What am I looking at?" expander content
# ---------------------------------------------------------------------------

def render_what_am_i_looking_at(sat_key: str, r_band: str, g_band: str, b_band: str,
                                  preset_name: str = "", location_label: str = "",
                                  is_contact_sheet: bool = False):
    """
    Render an expandable explanation panel below a satellite image.
    Contains two sections:
      1. A fixed primer on what satellite imagery is (vs a photograph)
      2. A dynamic explanation of the specific band combination displayed
    """
    with st.expander("ℹ️ What am I looking at? — Understanding satellite imagery", expanded=False):

        # -------------------------------------------------------------------
        # Section 1: What satellite imagery is
        # -------------------------------------------------------------------
        st.markdown("### Satellite imagery is not a photograph")
        st.markdown("""
A camera — in a phone, a drone, or an aircraft — captures what the human eye sees.
It records reflected light in three wavelength ranges: red, green, and blue. That is
the full extent of what a photograph captures.

A satellite multispectral sensor is fundamentally different. It measures reflected
energy across many separate wavelength ranges simultaneously — including ranges that
are completely invisible to the human eye. Sentinel-2 measures 12 separate wavelength
ranges. Landsat measures 11. Each range reveals different physical properties of the
ground surface.

When you look at a satellite image on this screen, you are not looking at a photograph.
You are looking at a scientific measurement rendered as colour to make it interpretable.
The colour assigned to each pixel is a choice — it represents the intensity of a
particular wavelength measurement, mapped onto a colour channel so human eyes can
perceive the difference.
        """)

        st.markdown("### What is a band?")
        st.markdown("""
Each wavelength range that a sensor measures is called a **band**. Think of it like
a filter on a camera — except instead of blocking colour, each band captures energy
at a precise part of the electromagnetic spectrum.

Some bands measure visible light (blue at 490 nm, green at 560 nm, red at 665 nm).
Others measure near-infrared (842 nm), which healthy vegetation reflects strongly but
human eyes cannot see at all. Others measure shortwave infrared (1610 nm, 2190 nm),
which reveals moisture content in soil, snow cover, and burn scars. Landsat even
includes a thermal infrared band that measures surface temperature directly.

Every band has a specific physical meaning. Choosing which band to display in which
colour channel is how you ask the satellite a specific question.
        """)

        st.markdown("### True color vs. false color")
        st.markdown("""
**True color** is when you assign the red band to the Red channel, the green band to
the Green channel, and the blue band to the Blue channel. The result mimics what a
camera would capture — it looks natural to the human eye. You see the ground roughly
as you would from an aircraft.

**False color** is when you assign any other combination of bands to the RGB channels.
The result does not look natural, but it reveals things a camera cannot show.

The most common example: assign the near-infrared band (NIR) to the Red channel.
Healthy vegetation reflects NIR strongly, so it appears bright red. Bare soil reflects
NIR weakly, so it appears dark. Open water absorbs NIR almost completely, so it appears
black. In a true color image, all three of these might look similar shades of brown and
green. In false color NIR, they are immediately and unambiguously distinguishable.

This is the core power of multispectral imaging. You are not making the image look
different for aesthetic reasons. You are selecting which physical measurement to
visualise in order to answer a specific question about the ground.
        """)

        st.markdown("### The tilt you may notice")
        st.markdown("""
Satellite images often appear tilted or diamond-shaped rather than as neat rectangles.
This is normal and reflects real satellite geometry.

Landsat orbits at approximately 705 km altitude on a sun-synchronous orbit inclined at
98 degrees relative to the equator. The satellite travels diagonally across the Earth
while the planet rotates underneath it. Each pass captures a diagonal strip of ground.
Sentinel-2 follows the same orbital geometry.

The black triangles in the corners of tilted images are simply areas the sensor did not
cover on that particular orbital pass. They contain no data. Professional GIS software
hides this by mosaicking many scenes together and reprojecting them into a flat grid —
but that requires additional processing. What you see here is the raw scene exactly as
the satellite captured it.
        """)

        # -------------------------------------------------------------------
        # Section 2: Dynamic explanation of this specific combination
        # -------------------------------------------------------------------
        st.divider()

        if is_contact_sheet:
            st.markdown("### How to read the comparison grid")
            st.markdown(f"""
The grid below shows **{sat_key}** imagery of **{location_label}** rendered in every
available band combination simultaneously. Each thumbnail uses a different assignment
of bands to the Red, Green, and Blue display channels.

Look at the same feature — a forest patch, a city, a river — across all thumbnails.
Notice how dramatically its colour changes between combinations. That is not an
artefact. Each colour change represents a real difference in what that combination
is measuring about the surface.

The last two thumbnails (NDVI and NDWI) are different from the others. They are not
RGB composites — they are mathematical indices computed from band ratios and displayed
using a fixed colour scale (green-to-red for vegetation health, blue for water
extent). These are the most analytically useful outputs: they give you a direct numeric
signal rather than an interpretive colour image.
            """)

            st.markdown("**What each preset combination is measuring:**")
            sat_info_local = satellite_catalog.SATELLITES[sat_key]
            for preset_name_item, preset_vals in sat_info_local["presets"].items():
                r = preset_vals["r"]
                g = preset_vals["g"]
                b = preset_vals["b"]
                note = preset_vals["note"]
                r_name = sat_info_local["bands"].get(r, {}).get("name", r)
                g_name = sat_info_local["bands"].get(g, {}).get("name", g)
                b_name = sat_info_local["bands"].get(b, {}).get("name", b)
                st.markdown(
                    f"**{preset_name_item}** — R: {r} ({r_name}), G: {g} ({g_name}), "
                    f"B: {b} ({b_name}). *{note}*"
                )

        else:
            sat_info_local = satellite_catalog.SATELLITES[sat_key]
            r_info = sat_info_local["bands"].get(r_band, {})
            g_info = sat_info_local["bands"].get(g_band, {})
            b_info = sat_info_local["bands"].get(b_band, {})

            r_name = r_info.get("name", r_band)
            g_name = g_info.get("name", g_band)
            b_name = b_info.get("name", b_band)
            r_wl   = r_info.get("wavelength", "")
            g_wl   = g_info.get("wavelength", "")
            b_wl   = b_info.get("wavelength", "")
            r_desc = r_info.get("description", "")
            g_desc = g_info.get("description", "")
            b_desc = b_info.get("description", "")

            is_true_color = (
                r_band in ("B04", "red") and
                g_band in ("B03", "green") and
                b_band in ("B02", "blue")
            )

            if preset_name and preset_name != "Manual":
                st.markdown(f"### This combination: {preset_name}")
            else:
                st.markdown("### This combination")

            if is_true_color:
                st.markdown(f"""
You are viewing a **true color** composite. This is the closest satellite imagery
gets to a photograph. The Red channel carries {r_band} ({r_name}, {r_wl}), the Green
channel carries {g_band} ({g_name}, {g_wl}), and the Blue channel carries {b_band}
({b_name}, {b_wl}) — matching the three colour ranges the human eye uses.

The result is a natural-looking image where vegetation appears green, bare soil appears
brown or tan, water appears dark blue, and built-up areas appear grey. Snow and clouds
appear white because they reflect strongly across all three visible bands.

Even though it looks like a photograph, it is still a processed scientific product.
The raw sensor data has been atmospherically corrected (to remove the effect of the
atmosphere on reflectance values), radiometrically calibrated (to convert raw digital
numbers to physically meaningful reflectance percentages), and display-stretched
(to map the actual reflectance range onto the 0-255 display range).
                """)
            else:
                st.markdown(f"""
You are viewing a **false color** composite. The colour in this image does not
correspond to what the human eye would see — each channel carries a measurement
from outside the visible spectrum, or a different visible band than usual.

**Red channel — {r_band} ({r_name}, {r_wl}):**
{r_desc}
In this image, anything that appears **red or bright** is a surface that reflects
strongly in the {r_wl} wavelength range.

**Green channel — {g_band} ({g_name}, {g_wl}):**
{g_desc}
In this image, anything that contributes to a **green tone** is reflecting strongly
in the {g_wl} range.

**Blue channel — {b_band} ({b_name}, {b_wl}):**
{b_desc}
In this image, anything that appears **blue** is reflecting strongly in the {b_wl}
range.

When all three channels combine, the colour of any given pixel tells you the relative
reflectance of that surface across all three wavelength ranges simultaneously. A pixel
that appears white is reflecting strongly in all three. A pixel that appears black is
absorbing energy in all three. A pixel that appears bright red is reflecting strongly
only in the {r_wl} range.
                """)

            if location_label:
                st.markdown(f"### What to look for over {location_label}")
                if "zanzibar" in location_label.lower() or "tanzania" in location_label.lower():
                    st.markdown("""
- **The island outline:** The boundary between land and ocean is sharp. Ocean absorbs
  near-infrared strongly — it will appear very dark in any combination that includes NIR.
- **Coral reefs:** In true color, shallow reef areas appear turquoise or light blue due
  to the sandy substrate and clear shallow water. In NIR combinations they disappear
  entirely as water absorbs all NIR.
- **Vegetation:** The island has a mix of dense coastal forest, smallholder farms, and
  the urban area of Stone Town. In false color NIR, the forest appears brighter red than
  the farmland, which appears paler. Urban Stone Town appears grey-dark regardless of
  combination.
- **Cloud cover:** Clouds appear white in every combination because they reflect all
  wavelengths strongly. Cloud shadows on the ground appear dark.
                    """)
                elif "moscow" in location_label.lower() or "russia" in location_label.lower():
                    st.markdown("""
- **The urban core:** Moscow is one of the largest cities in Europe. Dense urban areas
  appear grey in true color and purple-magenta in urban false color combinations. The
  contrast with surrounding forest is stark.
- **The Moscow River:** The river running through the city appears dark in any NIR
  combination. Water absorbs NIR almost completely.
- **Seasonal vegetation:** If the scene was captured in summer, the surrounding mixed
  forest appears bright green (true color) or bright red (NIR false color). Winter scenes
  will show snow (white across all combinations) and bare deciduous trees.
                    """)
                elif "alpharetta" in location_label.lower() or "georgia" in location_label.lower():
                    st.markdown("""
- **Suburban development pattern:** Alpharetta is a rapidly growing suburb north of
  Atlanta. The development pattern — road networks, residential subdivisions, commercial
  zones — is clearly visible in true color as a grey-tan grid overlaid on forest.
- **Urban heat:** In Landsat thermal band, developed areas appear warmer (brighter) than
  the surrounding forest. This is the urban heat island effect.
- **Forest fragments:** The Atlanta metro area retains significant forest between
  developments. In NIR false color, these fragments appear bright red against the grey
  built-up areas.
                    """)
                elif "serengeti" in location_label.lower() or "central tanzania" in location_label.lower():
                    st.markdown("""
- **Savanna texture:** The Serengeti ecosystem appears as a mosaic of grass, shrub, and
  occasional forest. In true color it looks uniformly tan-brown. In NIR false color the
  variation in vegetation density becomes clearly visible.
- **Seasonal water:** Ephemeral rivers and seasonal lakes appear as dark lines and patches.
  In NIR and SWIR combinations, moisture in the soil around these features makes them
  stand out.
- **Burned areas:** Controlled burns are common in East African savanna management.
  Recent burn scars appear very dark in NIR and bright in SWIR combinations.
                    """)
                else:
                    st.markdown("""
- Look for the **boundary between land and water** — water absorbs NIR almost completely
  and appears very dark in any combination that includes a NIR band.
- Look for **vegetation density gradients** — dense healthy vegetation reflects NIR more
  strongly than stressed or sparse vegetation. In NIR false color, this appears as varying
  intensities of red.
- Look for **built-up areas** — urban surfaces reflect SWIR more strongly than vegetation.
  In SWIR combinations, cities and roads appear brighter than surrounding land.
                    """)


# ===========================================================================
# MODULE ROUTING
# ---------------------------------------------------------------------------
# Time Series Explorer helpers (Word document builder)
# ---------------------------------------------------------------------------

def _ts_timeseries_png(df, stats, dataset, region, start_year, end_year):
    """Matplotlib time series chart: raw values, smoothed, and trend line."""
    import io as _io
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as _np

    unit  = gee_timeseries.DATASETS[dataset]["unit"]
    color = gee_timeseries.DATASETS[dataset]["chart_color"]

    fig, ax = plt.subplots(figsize=(7.2, 3.2))

    dates = df["date"].values
    vals  = df["value"].values

    # Raw values
    ax.plot(dates, vals, color=color, linewidth=0.8, alpha=0.35, label=f"{unit} (raw)")

    # Smoothed
    smooth = df.dropna(subset=["value_smooth"])
    ax.plot(smooth["date"].values, smooth["value_smooth"].values,
            color=color, linewidth=2.2, label=f"{unit} (smoothed)")

    # Trend line
    trend_y = _np.poly1d(stats["polyfit_z"])(stats["x_days"])
    ax.plot(stats["valid_dates"], trend_y, color="red", linewidth=1.5,
            linestyle="--", label="Linear trend")

    direction = "Increasing" if stats["slope_per_year"] > 0 else "Decreasing"
    ax.text(0.01, 0.97,
            f"Trend: {direction} ({stats['slope_per_year']:+.4f} {unit}/year)",
            transform=ax.transAxes, fontsize=8, va="top", color="red",
            bbox=dict(boxstyle="round,pad=0.3", facecolor="white",
                      edgecolor="red", linewidth=0.8))

    ax.set_title(f"{region} — {dataset} {start_year}–{end_year}", fontsize=9.5)
    ax.set_xlabel("Date", fontsize=8)
    ax.set_ylabel(unit, fontsize=8)
    ax.tick_params(labelsize=7)
    ax.legend(fontsize=7, loc="upper right")
    ax.grid(True, alpha=0.3, linestyle="--")
    fig.autofmt_xdate(rotation=30, ha="right")

    buf = _io.BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _ts_seasonal_png(stats, dataset):
    """Matplotlib seasonal cycle bar chart: average value by calendar month."""
    import io as _io
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    unit    = gee_timeseries.DATASETS[dataset]["unit"]
    color   = gee_timeseries.DATASETS[dataset]["chart_color"]
    monthly = stats["monthly_avg"]
    labels  = ["Jan","Feb","Mar","Apr","May","Jun",
               "Jul","Aug","Sep","Oct","Nov","Dec"]
    values  = [monthly.get(m, 0) for m in range(1, 13)]

    peak_val   = max(values)
    trough_val = min(values)
    bar_colors = []
    for v in values:
        if v == peak_val:
            bar_colors.append("#004d00" if "NDVI" in dataset else "#8b0000")
        elif v == trough_val:
            bar_colors.append("#c8a060")
        else:
            bar_colors.append(color)

    fig, ax = plt.subplots(figsize=(5.2, 2.8))
    ax.bar(labels, values, color=bar_colors)
    ax.set_title(f"Average seasonal cycle — {dataset}", fontsize=9.5)
    ax.set_xlabel("Month", fontsize=8)
    ax.set_ylabel(f"Mean {unit}", fontsize=8)
    ax.tick_params(labelsize=7)
    ax.grid(True, axis="y", alpha=0.3, linestyle="--")

    buf = _io.BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _ts_annual_png(df, stats, dataset, start_year, end_year):
    """Matplotlib annual means bar chart, colour-coded by period."""
    import io as _io
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as _np

    unit = gee_timeseries.DATASETS[dataset]["unit"]
    early_boundary  = start_year + max(1, (end_year - start_year) // 4)
    recent_boundary = end_year   - max(1, (end_year - start_year) // 4)

    _df = df.copy()
    _df["year"] = _df["date"].dt.year
    annual = _df.groupby("year")["value"].mean().reset_index()
    annual.columns = ["year", "mean"]
    overall_mean = annual["mean"].mean()

    def bar_colour(y):
        if y <= early_boundary:   return "#4a90d9"
        elif y >= recent_boundary: return "#2d7a2d"
        else:                      return "#b0b8c1"

    colours = [bar_colour(y) for y in annual["year"]]
    x_labels = [str(y) for y in annual["year"]]

    fig, ax = plt.subplots(figsize=(7.2, 2.8))
    ax.bar(x_labels, annual["mean"], color=colours)
    ax.axhline(overall_mean, color="#e67e22", linewidth=1.4, linestyle="--",
               label=f"Overall mean ({overall_mean:.3f})")

    from matplotlib.patches import Patch
    legend_els = [
        Patch(facecolor="#4a90d9", label=f"Early period (≤{early_boundary})"),
        Patch(facecolor="#2d7a2d", label=f"Recent period (≥{recent_boundary})"),
        Patch(facecolor="#b0b8c1", label="Transitional"),
    ]
    ax.legend(handles=legend_els, fontsize=6.5, loc="upper right")

    ax.set_title(f"Annual mean {dataset}", fontsize=9.5)
    ax.set_xlabel("Year", fontsize=8)
    ax.set_ylabel(f"Mean {unit}", fontsize=8)
    ax.tick_params(labelsize=7)
    plt.xticks(rotation=45, ha="right")
    ax.grid(True, axis="y", alpha=0.3, linestyle="--")

    buf = _io.BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def build_timeseries_docx(ai_text, dataset, region, start_year, end_year,
                           stats, unit, df, model_name=""):
    """Build a Word document for the Time Series Explorer result.

    Sections:
      1. Title and metadata
      2. Statistics table (mean, trend, peak/trough, amplitude)
      3. Time Series chart
      4. Seasonal Cycle chart
      5. Annual Comparison chart
      6. AI Interpretation (markdown → Word formatting)

    Charts are rendered on demand from df + stats using kaleido.
    Returns bytes ready for st.download_button.
    """
    import io as _io
    import re as _re
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()

    # Margins and font
    sec = doc.sections[0]
    sec.top_margin    = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin   = Inches(1.1)
    sec.right_margin  = Inches(1.1)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ---- Title ----
    t = doc.add_paragraph()
    r = t.add_run("Time Series Intelligence Report")
    r.bold = True
    r.font.size = Pt(16)

    # ---- Metadata ----
    for label, value in [
        ("Dataset:",   dataset),
        ("Region:",    region),
        ("Period:",    f"{start_year} – {end_year}"),
        ("AI model:",  model_name or "Built-in interpretation"),
    ]:
        p = doc.add_paragraph()
        p.add_run(label + " ").bold = True
        p.add_run(str(value))

    doc.add_paragraph()

    # ---- Statistics table ----
    h = doc.add_paragraph("Statistics")
    h.style = doc.styles["Heading 1"]

    stat_rows = [
        ("Observations",        str(stats.get("count", "—"))),
        (f"Mean ({unit})",      f"{stats['mean']:.3f}"),
        (f"Min ({unit})",       f"{stats['min']:.3f}"),
        (f"Max ({unit})",       f"{stats['max']:.3f}"),
        ("Trend direction",     stats.get("trend_direction", "—").capitalize()),
        (f"Trend ({unit}/year)",f"{stats['slope_per_year']:+.4f}"),
        ("Peak month",          f"{stats['peak_month']} ({stats['peak_value']:.3f} {unit})"),
        ("Trough month",        f"{stats['trough_month']} ({stats['trough_value']:.3f} {unit})"),
        (f"Amplitude ({unit})", f"{stats['amplitude']:.3f}"),
    ]
    tbl = doc.add_table(rows=len(stat_rows), cols=2)
    tbl.style = "Table Grid"
    for ri, (lbl, val) in enumerate(stat_rows):
        cells = tbl.rows[ri].cells
        cells[0].text = lbl
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
        cells[1].text = val

    doc.add_paragraph()

    # ---- Charts — rendered with matplotlib (no browser required) ----
    chart_specs = [
        ("Time Series",
         lambda: _ts_timeseries_png(df, stats, dataset, region, start_year, end_year)),
        ("Seasonal Cycle",
         lambda: _ts_seasonal_png(stats, dataset)),
        ("Annual Comparison",
         lambda: _ts_annual_png(df, stats, dataset, start_year, end_year)),
    ]

    for heading, render_fn in chart_specs:
        h = doc.add_paragraph(heading)
        h.style = doc.styles["Heading 1"]
        try:
            png = render_fn()
            doc.add_picture(_io.BytesIO(png), width=Inches(4.8))
        except Exception as _e:
            doc.add_paragraph(f"[Chart could not be rendered: {_e}]")
        doc.add_paragraph()

    # ---- AI Interpretation ----
    h = doc.add_paragraph("AI Interpretation")
    h.style = doc.styles["Heading 1"]

    def _ts_inline_bold(paragraph, text):
        parts = _re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                paragraph.add_run(part[2:-2]).bold = True
            else:
                paragraph.add_run(part)

    if ai_text:
        lines = ai_text.splitlines()
        i = 0
        while i < len(lines):
            stripped = lines[i].strip()
            if stripped.startswith("## ") or stripped.startswith("# "):
                p = doc.add_paragraph(stripped.lstrip("# ").strip())
                p.style = doc.styles["Heading 1"]
                i += 1
            elif stripped.startswith("|"):
                block = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    block.append(lines[i])
                    i += 1
                data_lines = [l for l in block if not _re.match(r"^\s*\|[-| :]+\|\s*$", l)]
                if len(data_lines) >= 2:
                    def _split(ln):
                        return [c.strip() for c in ln.strip().strip("|").split("|")]
                    hdrs = _split(data_lines[0])
                    rws  = [_split(l) for l in data_lines[1:]]
                    t2   = doc.add_table(rows=1 + len(rws), cols=len(hdrs))
                    t2.style = "Table Grid"
                    for ci, hdr in enumerate(hdrs):
                        t2.rows[0].cells[ci].text = hdr
                        for run in t2.rows[0].cells[ci].paragraphs[0].runs:
                            run.bold = True
                    for ri2, rd in enumerate(rws):
                        for ci, ct in enumerate(rd):
                            if ci < len(t2.rows[ri2+1].cells):
                                t2.rows[ri2+1].cells[ci].text = ct
                    doc.add_paragraph()
                else:
                    for bl in block:
                        doc.add_paragraph(bl.strip())
            elif _re.match(r"^\d+\.\s", stripped):
                p = doc.add_paragraph(style="List Number")
                _ts_inline_bold(p, stripped[stripped.index(". ") + 2:])
                i += 1
            elif stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                _ts_inline_bold(p, stripped[2:])
                i += 1
            elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
                p = doc.add_paragraph()
                p.add_run(stripped.strip("*")).bold = True
                i += 1
            elif stripped == "" or stripped.startswith("---"):
                doc.add_paragraph()
                i += 1
            else:
                if stripped:
                    p = doc.add_paragraph()
                    _ts_inline_bold(p, stripped)
                i += 1
    else:
        doc.add_paragraph("No AI interpretation available.")

    # ---- Footer ----
    f = doc.add_paragraph()
    f.add_run("Generated by EOIL — AI-Native Earth Observation Innovation Lab").italic = True

    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# The Time Series Explorer renders first and calls st.stop() when active.
# That prevents the EO Explorer code below from running.
# The EO Explorer code needs no indentation changes — it just runs when
# the Time Series module is not selected.
# ===========================================================================

if selected_module == "📈 Time Series Explorer":

    # -----------------------------------------------------------------------
    # MODULE 2 — Time Series Explorer
    # All logic lives in gee_timeseries.py. This block handles layout only.
    # -----------------------------------------------------------------------

    gee_available = st.session_state.gee_available

    # --- Header ---
    st.subheader("📈 Time Series Explorer")
    st.caption(
        "Select a dataset, location, and year range, then click Run Analysis."
    )

    with st.expander("ℹ️ What does this module do?", expanded=False):
        st.markdown("""
**This module answers the question: how has this location changed over time?**

It pulls 10+ years of satellite measurements for any location on Earth and shows
you how a chosen index — vegetation greenness, land surface temperature, or others —
has shifted across the full period.

**What you will see after running an analysis:**

- **Summary metrics** — mean value, long-term trend per year, and seasonal amplitude at a glance
- **Time series chart** — every data point plotted across the full date range, with a smoothed trend line and a linear trend overlay
- **Seasonal cycle chart** — the average pattern by month, showing when values peak and trough across a typical year
- **Annual comparison chart** — one bar per year, colour-coded to highlight how early years compare to recent years
- **AI interpretation** — a plain-language explanation of what the data shows, why the pattern exists, and what it means in practice

**How to use it:**

1. Choose a dataset from the dropdown
2. Type any city, country, or region in the location box
3. Set the year range
4. Click Run Analysis
        """)

    # --- Row 1: Dataset + Location (mirrors Spectral Explorer layout) ---
    col_ds, col_loc = st.columns([1, 3])

    with col_ds:
        ts_dataset = st.selectbox(
            "Dataset",
            list(gee_timeseries.DATASETS.keys()),
            key="ts_dataset",
        )

    with col_loc:
        ts_custom_place = st.text_input(
            "Location — type any city, country, or region",
            placeholder="e.g. Sahel, West Africa   |   Tanzania   |   Patagonia   |   Fergana Valley",
            key="ts_custom_place",
        )

    # --- Row 2: Year range + Run button ---
    yr_col1, yr_col2, yr_col3 = st.columns([1, 1, 1], vertical_alignment="bottom")

    with yr_col1:
        ts_start_year = st.number_input(
            "From year", min_value=2000, max_value=2023,
            value=2014, step=1, key="ts_start_year",
        )

    with yr_col2:
        ts_end_year = st.number_input(
            "To year", min_value=2001, max_value=2024,
            value=2024, step=1, key="ts_end_year",
        )

    with yr_col3:
        run_btn = st.button(
            "▶ Run Analysis", type="primary",
            use_container_width=True, key="ts_run",
        )

    # Dataset educational info — collapsible, sits just below the controls
    with st.expander("ℹ️ About this dataset", expanded=False):
        ds_info = gee_timeseries.DATASETS[ts_dataset]
        col_a, col_b = st.columns(2)
        with col_a:
            st.markdown(f"**Sensor:** {ds_info['sensor']}")
            st.markdown(f"**Resolution:** {ds_info['resolution']}")
            st.markdown(f"**Revisit:** {ds_info['revisit']}")
            st.markdown(f"**Available from:** {ds_info['available_from']}")
        with col_b:
            st.markdown(f"**Measures:** {ds_info['measures']}")
            st.markdown(f"**Best for:** {ds_info['best_for']}")
            st.markdown(f"**Limitation:** {ds_info['limitation']}")

    # Landsat datasets are at 30 m — 70x more pixels than MODIS.
    # Warn the user before they run a slow query on a large region.
    if "Landsat" in ts_dataset:
        st.warning(
            "⏱ **Landsat is a slow dataset for large regions.** "
            "At 30 m resolution it has ~70x more pixels than MODIS for the same area. "
            "Queries for a country or large region can take 1–3 minutes. "
            "For country-level trends use MODIS NDVI instead. "
            "Landsat works best for small areas: a field, a wetland, a river reach."
        )

    st.divider()

    # --- Resolve region from typed location ---
    ts_bbox        = None
    ts_region_name = ""

    if ts_custom_place.strip():
        cached_place = st.session_state.get("ts_geocoded_place", "")
        cached_bbox  = st.session_state.get("ts_geocoded_bbox",  None)

        if ts_custom_place.strip() != cached_place:
            # New location typed — clear any stale map click
            map_picker.clear_click("ts")
            with st.spinner(f"Looking up '{ts_custom_place}'..."):
                result_bbox = geocoder.geocode_place(ts_custom_place)
            if result_bbox:
                st.session_state.ts_geocoded_place = ts_custom_place.strip()
                st.session_state.ts_geocoded_bbox  = result_bbox
                cached_bbox = result_bbox
            else:
                st.session_state.ts_geocoded_place = ""
                st.session_state.ts_geocoded_bbox  = None
                cached_bbox = None
                st.error(f"Could not geocode '{ts_custom_place}'. Using preset region.")

        if cached_bbox:
            ts_bbox        = cached_bbox
            ts_region_name = ts_custom_place.strip()
            st.caption(f"📍 {ts_region_name}")
        else:
            # Location changed and geocode failed — clear stale map click
            map_picker.clear_click("ts")

    # Map picker — optional override shown once a geocoded centre exists
    if ts_bbox:
        with st.expander("📍 Refine location — click map to set exact area", expanded=False):
            picked = map_picker.render_map_picker(
                centre_bbox     = ts_bbox,
                picker_key      = "ts",
                default_size_km = 100,
            )
            if picked:
                ts_bbox = picked

    # GEE status — compact caption so it doesn't dominate the page
    if gee_available:
        st.caption("🟢 GEE connected — live data active.")
    else:
        st.caption(
            "🔵 No GEE credentials — showing sample data modeled from MODIS climatology. "
            "Add GEE_SERVICE_ACCOUNT_JSON to Streamlit secrets to enable live queries."
        )

    # --- Session state initialisation for Time Series results ---
    for _k, _v in [
        ("ts_df",     None), ("ts_stats",   None),
        ("ts_result_dataset", None), ("ts_result_region", None),
        ("ts_result_start",   None), ("ts_result_end",    None),
        ("ts_is_sample", True),
        ("ts_ai_result", None), ("ts_ai_model", None),
    ]:
        if _k not in st.session_state:
            st.session_state[_k] = _v

    # --- Run Analysis ---
    # Two-step pattern: button click clears results and queues a run, then
    # st.rerun() forces a full re-render (blank page). The pending run block
    # below picks up on the next render and executes the analysis.
    if run_btn:
        if not ts_bbox:
            st.warning("Enter a location first.")
        elif ts_start_year >= ts_end_year:
            st.error("Start year must be before end year.")
        else:
            # Step 1: clear old results and queue the analysis parameters
            st.session_state.ts_df        = None
            st.session_state.ts_ai_result = None
            st.session_state.ts_pending_run = {
                "bbox":    ts_bbox,
                "dataset": ts_dataset,
                "start":   ts_start_year,
                "end":     ts_end_year,
                "region":  ts_region_name,
            }
            st.rerun()

    # Step 2: execute the queued analysis (runs on the render AFTER the button click)
    if st.session_state.get("ts_pending_run"):
        p = st.session_state.ts_pending_run
        st.session_state.ts_pending_run = None  # consume the queue entry
        with st.spinner(f"Fetching {p['dataset']} data for {p['region']}..."):
            try:
                if gee_available:
                    df       = gee_timeseries.extract_time_series_gee(
                        p["bbox"], p["dataset"], p["start"], p["end"]
                    )
                    is_sample = False
                else:
                    df       = gee_timeseries.generate_sample_data(
                        p["region"], p["dataset"], p["start"], p["end"]
                    )
                    is_sample = True

                stats = gee_timeseries.compute_statistics(df, p["dataset"])
                st.session_state.ts_df             = df
                st.session_state.ts_stats          = stats
                st.session_state.ts_result_dataset = p["dataset"]
                st.session_state.ts_result_region  = p["region"]
                st.session_state.ts_result_start   = p["start"]
                st.session_state.ts_result_end     = p["end"]
                st.session_state.ts_is_sample      = is_sample
                st.session_state.ts_ai_result      = None
                st.session_state.ts_ai_model       = None

                # Reset AI prompt to match the new region and dataset
                st.session_state.ts_ai_prompt = (
                    f"Interpret this {p['dataset']} time series for {p['region']} "
                    f"from {p['start']} to {p['end']}. "
                    f"Cover: what the data shows, why the pattern exists, "
                    f"one practical application, and one limitation."
                )
                st.success(
                    f"Analysis complete — {stats['count']} data points, "
                    f"{p['start']} to {p['end']}."
                )
            except Exception as e:
                st.error(f"Analysis failed: {e}")

    # --- Display results (from session state so they survive widget interactions) ---
    if st.session_state.ts_df is not None:
        df      = st.session_state.ts_df
        stats   = st.session_state.ts_stats
        r_ds    = st.session_state.ts_result_dataset
        r_reg   = st.session_state.ts_result_region
        r_start = st.session_state.ts_result_start
        r_end   = st.session_state.ts_result_end

        if st.session_state.ts_is_sample:
            st.caption("Showing sample data. Add GEE credentials to switch to live data.")

        # Reusable thick section divider — heavier than st.divider()
        def section_break():
            st.markdown(
                '<hr style="border: none; border-top: 3px solid #d0d0d0; margin: 28px 0 20px 0;">',
                unsafe_allow_html=True,
            )

        # --- SECTION 1: Summary metrics ---
        unit = gee_timeseries.DATASETS[r_ds]["unit"]
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Data points", stats["count"])
        c2.metric(f"Mean {unit}", f"{stats['mean']:.3f}")
        c3.metric("Trend / year",
                  f"{stats['slope_per_year']:+.4f} {unit}",
                  delta_color="normal")
        c4.metric("Seasonal amplitude", f"{stats['amplitude']:.3f} {unit}")

        section_break()

        # --- SECTION 2: Time series chart ---
        st.subheader("📈 Time Series")
        fig_ts = gee_timeseries.build_timeseries_chart(
            df, stats, r_ds, r_reg, r_start, r_end
        )
        st.plotly_chart(fig_ts, use_container_width=True)

        section_break()

        # --- SECTION 3: Seasonal cycle chart + statistics panel ---
        st.subheader("🗓️ Seasonal Cycle")
        col_seasonal, col_stats = st.columns([2, 1])

        with col_seasonal:
            fig_seasonal = gee_timeseries.build_seasonal_chart(stats, r_ds)
            st.plotly_chart(fig_seasonal, use_container_width=True)

        with col_stats:
            st.markdown("**Seasonal statistics**")
            st.markdown(f"- Peak month: **{stats['peak_month']}** ({stats['peak_value']:.3f} {unit})")
            st.markdown(f"- Trough month: **{stats['trough_month']}** ({stats['trough_value']:.3f} {unit})")
            st.markdown(f"- Amplitude: **{stats['amplitude']:.3f} {unit}**")
            st.divider()
            st.markdown("**Long-term trend**")
            direction_emoji = "📈" if stats["slope_per_year"] > 0 else "📉"
            st.markdown(
                f"{direction_emoji} {stats['trend_direction'].capitalize()} "
                f"at {stats['slope_per_year']:+.4f} {unit}/year"
            )

        section_break()

        # --- SECTION 4: Annual means chart ---
        st.subheader("📊 Annual Comparison")
        early_boundary  = r_start + max(1, (r_end - r_start) // 4)
        recent_boundary = r_end   - max(1, (r_end - r_start) // 4)

        # Compute change statistics for the summary panel
        early_mean  = df[df["date"].dt.year <= early_boundary]["value"].mean()
        recent_mean = df[df["date"].dt.year >= recent_boundary]["value"].mean()
        diff_val    = recent_mean - early_mean
        pct_change  = (diff_val / early_mean * 100) if early_mean != 0 else 0

        col_chart, col_summary = st.columns([3, 1])

        with col_chart:
            fig_annual = gee_timeseries.build_annual_means_chart(
                df, r_ds, r_start, r_end, early_boundary, recent_boundary
            )
            st.plotly_chart(fig_annual, use_container_width=True)

        with col_summary:
            st.markdown("**Period comparison**")
            st.metric(
                label=f"Early mean ({r_start}–{early_boundary})",
                value=f"{early_mean:.3f} {unit}",
            )
            st.metric(
                label=f"Recent mean ({recent_boundary}–{r_end})",
                value=f"{recent_mean:.3f} {unit}",
                delta=f"{diff_val:+.3f} {unit}",
            )
            direction = "increase" if diff_val > 0 else "decrease"
            emoji     = "🟢" if diff_val > 0 else "🔴"
            st.markdown(
                f"{emoji} **{abs(pct_change):.1f}% {direction}** "
                f"over the full period."
            )

        section_break()

        # --- SECTION 5: AI Interpretation ---
        st.subheader("🤖 AI Interpretation")
        default_prompt = (
            f"Interpret this {r_ds} time series for {r_reg} "
            f"from {r_start} to {r_end}. "
            f"Cover: what the data shows, why the pattern exists, "
            f"one practical application, and one limitation."
        )
        ts_user_prompt = st.text_area(
            "Edit the prompt before submitting (optional):",
            value=default_prompt,
            height=100,
            key="ts_ai_prompt",
        )
        if st.button("Get AI Interpretation", type="primary", key="ts_ai_btn"):
            with st.spinner("Thinking..."):
                interpretation, model_used = gee_timeseries.get_ai_interpretation(
                    stats, r_ds, r_reg, r_start, r_end,
                    custom_prompt=ts_user_prompt,
                    groq_key=config.GROQ_API_KEY,
                    gemini_key=config.GEMINI_API_KEY,
                )
                st.session_state.ts_ai_result = interpretation
                st.session_state.ts_ai_model  = model_used

        if st.session_state.get("ts_ai_result"):
            with st.expander("📋 AI Interpretation", expanded=True):
                st.markdown(st.session_state.ts_ai_result)
                _ts_model = st.session_state.get("ts_ai_model")
                if _ts_model:
                    st.caption(f"AI response from **{_ts_model}**")
                else:
                    st.caption("Showing built-in interpretation. Add GROQ_API_KEY or GEMINI_API_KEY to enable AI.")

            # Download buttons
            _ts_safe = f"{r_ds.replace(' ', '_').replace('/', '-')}_{r_reg.replace(' ', '_').replace(',', '')[:25]}_{r_start}_{r_end}"
            _ts_dc1, _ts_dc2 = st.columns([1, 1])
            with _ts_dc1:
                st.download_button(
                    label="⬇ Download as Markdown",
                    data=st.session_state.ts_ai_result,
                    file_name=f"timeseries_{_ts_safe}.md",
                    mime="text/markdown",
                    key="ts_dl_md",
                )
            with _ts_dc2:
                _ts_docx = build_timeseries_docx(
                    ai_text=st.session_state.ts_ai_result,
                    dataset=r_ds,
                    region=r_reg,
                    start_year=r_start,
                    end_year=r_end,
                    stats=stats,
                    unit=unit,
                    df=df,
                    model_name=st.session_state.get("ts_ai_model", ""),
                )
                st.download_button(
                    label="⬇ Download as Word (.docx)",
                    data=_ts_docx,
                    file_name=f"timeseries_{_ts_safe}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="ts_dl_docx",
                )

        section_break()

        # --- SECTION 6: Data Quality ---
        st.subheader("🔍 Data Quality")
        _ds_info = gee_timeseries.DATASETS[r_ds]
        _obs     = stats["count"]
        if _obs >= 100:
            _ts_conf = "High — large sample. Trend and seasonal estimates are reliable."
        elif _obs >= 50:
            _ts_conf = "Moderate — adequate sample. Treat long-term trend slope with care."
        else:
            _ts_conf = "Limited — fewer than 50 observations. Trends may not be reliable."
        st.info(
            f"**Sensor:** {_ds_info['sensor']}  \n"
            f"**Spatial resolution:** {_ds_info['resolution']}  \n"
            f"**Revisit cadence:** {_ds_info['revisit']}  \n"
            f"**Observations used:** {_obs} data points ({r_start}–{r_end})  \n"
            f"**Cloud handling:** Pre-composited by Google Earth Engine. "
            f"Cloud pixels excluded before aggregation. No gaps from cloud cover.  \n"
            f"**Confidence:** {_ts_conf}"
        )

        section_break()

        # --- SECTION 7: Export ---
        st.subheader("📥 Export")
        export_df = df[["date", "value"]].copy()
        export_df.columns = ["Date", r_ds]
        csv_data  = export_df.to_csv(index=False).encode()
        safe_reg  = r_reg.replace(" ", "_").replace(",", "")
        safe_ds   = r_ds.replace(" ", "_").replace("/", "-")
        fname_ts  = f"time_series_{safe_reg}_{safe_ds}_{r_start}_{r_end}.csv"
        st.download_button(
            label="⬇️ Download time series as CSV",
            data=csv_data,
            file_name=fname_ts,
            mime="text/csv",
            key="ts_export_csv",
        )
        st.caption(f"{len(export_df)} data points — Date and {r_ds} value columns.")

    else:
        # No analysis run yet — show a prompt to get started
        st.markdown("---")
        st.markdown(
            "**Type a location above, then click Run Analysis.**\n\n"
            "The module will produce a time series chart, a seasonal cycle chart, "
            "two comparison maps, and an AI interpretation of what the data shows."
        )

    # Stop here — do not render the EO Explorer below
    st.stop()

# ---------------------------------------------------------------------------
# MODULE 2 — Spectral Explorer (reached only when Spectral Explorer is selected)
# ---------------------------------------------------------------------------

if selected_module == "🔬 Spectral Explorer":

    st.subheader("🔬 Spectral Explorer")
    st.caption(
        "Choose any location, satellite, and band combination. "
        "Fetch real imagery and compare how the same area looks in different spectral views."
    )

    with st.expander("ℹ️ What does this module do?", expanded=False):
        st.markdown("""
**This module answers the question: what can satellites see that cameras cannot?**

It fetches real satellite imagery from NASA and ESA archives for any location on Earth
and lets you render the same scene through different combinations of spectral bands.
Each combination reveals different physical properties of the ground surface — vegetation
health, water extent, urban heat, burn scars, soil moisture, and more.

**What you will see after fetching a scene:**

- **Scene timeline** — a chart of all available satellite passes in your date range, scored by how much of your area they cover
- **Rendered image** — the actual satellite data displayed using your chosen band combination, with a download button
- **Scene details** — date, cloud cover, satellite, and an explanation of what each channel is measuring
- **Compare all views** — every available band combination rendered side by side so you can see how the same place looks across the full electromagnetic spectrum
- **AI explanation** — a plain-language interpretation of what the chosen combination reveals and why

**How to use it:**

1. Choose a satellite from the dropdown
2. Type any location in the search box
3. Set a date range and maximum cloud cover
4. Click Search for scenes
5. Choose a band combination and click Render
        """)

    # -----------------------------------------------------------------------
    # Initialise all session state keys up front so they always exist.
    # This separates state management from display logic.
    # -----------------------------------------------------------------------
    for key, default in [
        ("se_items",           []),
        ("se_best_item",       None),
        ("se_coverage",        []),
        ("se_timeline_items",  []),
        ("se_prev_satellite",  None),
        ("se_prev_preset",     None),
        ("se_rendered_arr",    None),
        ("se_rendered_info",   None),
        ("se_contact_results",   None),
        ("se_contact_info",      None),
        ("se_selected_item_id",  None),
    ]:
        if key not in st.session_state:
            st.session_state[key] = default

    # -----------------------------------------------------------------------
    # Row 1: Satellite + Location
    # -----------------------------------------------------------------------
    col_sat, col_place = st.columns([1, 3])

    with col_sat:
        sat_key  = st.selectbox("Satellite", list(satellite_catalog.SATELLITES.keys()), key="se_satellite")
        sat_info = satellite_catalog.SATELLITES[sat_key]

    with col_place:
        place_input = st.text_input(
            "Location — type any city, country, or region",
            placeholder="e.g. Zanzibar   |   Moscow, Russia   |   Serengeti, Tanzania   |   Alpharetta, Georgia",
            key="se_place_input",
        )

    # Resolve bounding box from the typed place name.
    # Geocoding runs on every render while the text box has content.
    # The spinner only shows on the first lookup — subsequent renders
    # use the cached result already stored in session state.
    bbox_se        = None
    location_label = ""

    if place_input.strip():
        # Only call the geocoder when the place text has changed
        cached_place = st.session_state.get("se_geocoded_place", "")
        cached_bbox  = st.session_state.get("se_geocoded_bbox",  None)

        if place_input.strip() != cached_place:
            # New location typed — clear any stale map click
            map_picker.clear_click("se")
            with st.spinner(f"Looking up '{place_input}'..."):
                result_bbox = geocoder.geocode_place(place_input)
            if result_bbox:
                st.session_state.se_geocoded_place = place_input.strip()
                st.session_state.se_geocoded_bbox  = result_bbox
                cached_bbox = result_bbox
            else:
                st.session_state.se_geocoded_place = ""
                st.session_state.se_geocoded_bbox  = None
                cached_bbox = None
                st.error(f"Could not find '{place_input}'. Check the spelling or try a broader name (e.g. country instead of city).")

        if cached_bbox:
            bbox_se        = cached_bbox
            location_label = place_input.strip()
            area = geocoder.bbox_area_km2(bbox_se)
            st.caption(f"📍 {location_label} — bbox {[round(x, 3) for x in bbox_se]} — {area:,.0f} km²")
    else:
        # Text box is empty — clear any cached geocode result
        st.session_state.se_geocoded_place = ""
        st.session_state.se_geocoded_bbox  = None

    # Map picker — optional override shown once a geocoded centre exists
    if bbox_se:
        with st.expander("📍 Refine location — click map to set exact area", expanded=False):
            picked = map_picker.render_map_picker(
                centre_bbox     = bbox_se,
                picker_key      = "se",
                default_size_km = 50,
            )
            if picked:
                bbox_se = picked

    # Clear all search results when the user changes the location.
    # Compare the typed text against what was last searched — not against
    # previous widget states, which was the source of the clearing bugs.
    last_searched = st.session_state.get("se_last_searched_location", "")
    if last_searched and location_label and location_label.lower() != last_searched.lower():
        st.session_state.se_items            = []
        st.session_state.se_timeline_items   = []
        st.session_state.se_coverage         = []
        st.session_state.se_best_item        = None
        st.session_state.se_rendered_arr     = None
        st.session_state.se_contact_results  = None
        st.session_state.se_selected_item_id = None

    # -----------------------------------------------------------------------
    # Row 2: Date range + Cloud cover + Search button
    # The Search button lives here because it acts on these parameters.
    # Once the user has set satellite, location, dates and cloud cover,
    # the next natural action is to search — not scroll past band controls.
    # -----------------------------------------------------------------------
    # vertical_alignment="bottom" aligns the button with the bottom edge of
    # the date inputs and slider, matching their visual baseline precisely.
    col_d1, col_d2, col_cloud, col_search = st.columns([1, 1, 1, 1], vertical_alignment="bottom")
    with col_d1:
        date_start = st.date_input("From date", value=date.today() - timedelta(days=180), key="se_date_start")
    with col_d2:
        date_end   = st.date_input("To date",   value=date.today(), key="se_date_end")
    with col_cloud:
        if sat_info.get("sar"):
            st.info("SAR is cloud-independent. No cloud filter applied.")
            max_cloud = 100
        else:
            max_cloud = st.slider("Max cloud cover %", 0, 50, 20, key="se_cloud")
    with col_search:
        fetch_btn = st.button("🔍 Search for scenes", type="primary", key="se_fetch", use_container_width=True)

    date_range_str = f"{date_start.isoformat()}/{date_end.isoformat()}"

    # -----------------------------------------------------------------------
    # SEARCH — runs when button clicked, stores results in session state
    # -----------------------------------------------------------------------
    if fetch_btn:
        if not bbox_se:
            st.warning("Select a location first.")
        else:
            with st.spinner("Searching Planetary Computer..."):
                try:
                    cat   = spectral_explorer.get_catalog()
                    items = spectral_explorer.search_scenes(
                        cat, sat_info["collection"], bbox_se,
                        date_range_str, max_cloud, sat_info["cloud_field"],
                    )
                    st.session_state.se_items          = items
                    st.session_state.se_timeline_items = items
                    st.session_state.se_selected_item_id = None  # reset on new search
                except Exception as e:
                    st.error(f"Search failed: {e}")
                    items = []

            if items:
                with st.spinner("Scoring scenes for coverage..."):
                    # Use True Color bands for scoring — they are always valid
                    # for the current satellite and don't depend on what the
                    # user has selected in the band dropdowns below.
                    _tc      = sat_info["presets"].get("True Color", {})
                    _band_ids = list(sat_info["bands"].keys())
                    _score_r  = _tc.get("r", _band_ids[0])
                    _score_g  = _tc.get("g", _band_ids[0])
                    _score_b  = _tc.get("b", _band_ids[0])
                    best_item, best_pct, scores = spectral_explorer.find_best_scene(
                        items, _score_r, _score_g, _score_b, sat_key, max_to_check=6
                    )
                st.session_state.se_best_item = best_item
                st.session_state.se_coverage  = scores
                st.session_state.se_rendered_arr    = None
                st.session_state.se_contact_results = None
                # Stamp the location so future renders know what was searched
                st.session_state.se_last_searched_location = location_label
                st.success(
                    f"Found {len(items)} scenes. "
                    f"Best: {best_item.datetime.strftime('%Y-%m-%d')} — {best_pct:.0f}% coverage. "
                    f"Now choose a band combination below and click Render."
                )
            else:
                st.warning("No scenes found. Try a wider date range or higher cloud cover limit.")

    # -----------------------------------------------------------------------
    # DISPLAY SEARCH RESULTS — shown here, directly below the search section,
    # so the user can see what is available before choosing how to visualise it
    # -----------------------------------------------------------------------
    if st.session_state.se_timeline_items:
        fig = spectral_explorer.build_timeline(
            st.session_state.se_timeline_items,
            sat_info["cloud_field"],
            st.session_state.se_coverage,
        )
        st.plotly_chart(fig, use_container_width=True)

        if st.session_state.se_coverage:
            st.markdown("**Top scenes scored by coverage:**")
            for date_str, cloud, pct, _ in st.session_state.se_coverage:
                bar       = "█" * int(pct / 10) + "░" * (10 - int(pct / 10))
                cloud_str = f"Cloud {cloud:.0f}%" if sat_info["cloud_field"] else "SAR"
                st.caption(f"`{date_str}`  {cloud_str}  Coverage: {bar} {pct:.0f}%")

        # -------------------------------------------------------------------
        # SCENE SELECTOR — lets the user pick any scene from the search results
        # instead of always using the auto-selected best scene.
        # Useful when the best scene is from a tile that covers a neighbouring
        # area (e.g. Zanzibar when you searched for mainland Tanzania).
        # -------------------------------------------------------------------
        _all_items = st.session_state.se_items
        if _all_items:
            # Build a label for each scene: "Dec 12, 2025 — Cloud 1.6%"
            # For SAR satellites there is no cloud cover field.
            def _scene_label(it):
                d = it.datetime.strftime("%b %d, %Y")
                cf = sat_info.get("cloud_field")
                if cf:
                    c = it.properties.get(cf, 0)
                    return f"{d}  —  Cloud {c:.1f}%"
                return d

            _labels   = [_scene_label(it) for it in _all_items]
            _item_ids = [it.id for it in _all_items]

            # Default to whichever scene was auto-selected (best coverage)
            _best = st.session_state.get("se_best_item")
            if _best and _best.id in _item_ids:
                _default_idx = _item_ids.index(_best.id)
            else:
                _default_idx = 0

            # If user previously selected a scene, keep that selection
            _prev_id = st.session_state.get("se_selected_item_id")
            if _prev_id and _prev_id in _item_ids:
                _default_idx = _item_ids.index(_prev_id)

            _chosen_label = st.selectbox(
                "Scene — select which satellite pass to render",
                options=_labels,
                index=_default_idx,
                key="se_scene_selector",
            )
            _chosen_idx = _labels.index(_chosen_label)
            st.session_state.se_selected_item_id = _item_ids[_chosen_idx]

    # -----------------------------------------------------------------------
    # Row 3: Band assignment
    # -----------------------------------------------------------------------
    st.divider()
    st.markdown("**Band combination — choose a preset or assign bands manually**")

    band_ids      = list(sat_info["bands"].keys())
    band_names    = [f"{k} — {v['name']} ({v['wavelength']})" for k, v in sat_info["bands"].items()]
    band_map      = dict(zip(band_names, band_ids))
    id_to_display = {v: k for k, v in band_map.items()}

    # True color defaults for this satellite
    tc_preset    = sat_info["presets"].get("True Color", {})
    default_r_id = tc_preset.get("r", band_ids[min(2, len(band_ids)-1)])
    default_g_id = tc_preset.get("g", band_ids[min(1, len(band_ids)-1)])
    default_b_id = tc_preset.get("b", band_ids[0])

    # Reset R/G/B when satellite changes
    satellite_changed = (sat_key != st.session_state.se_prev_satellite)
    if satellite_changed:
        st.session_state.se_prev_satellite = sat_key
        st.session_state.se_r = id_to_display.get(default_r_id, band_names[0])
        st.session_state.se_g = id_to_display.get(default_g_id, band_names[0])
        st.session_state.se_b = id_to_display.get(default_b_id, band_names[0])
        st.session_state.se_prev_preset = None
        # Clear stored results — they belong to the old satellite
        st.session_state.se_rendered_arr    = None
        st.session_state.se_contact_results = None

    # Initialise R/G/B to true color on very first load
    if "se_r" not in st.session_state:
        st.session_state.se_r = id_to_display.get(default_r_id, band_names[0])
    if "se_g" not in st.session_state:
        st.session_state.se_g = id_to_display.get(default_g_id, band_names[0])
    if "se_b" not in st.session_state:
        st.session_state.se_b = id_to_display.get(default_b_id, band_names[0])

    # Preset dropdown.
    # "— Custom —" is a read-only indicator that appears automatically when
    # the R/G/B channels no longer match any named preset. The user never
    # needs to pick it — it switches on its own.
    #
    # Streamlit rule: a widget's session state key can only be set BEFORE
    # the widget is instantiated in a given render cycle. So the mismatch
    # check must happen here, before st.selectbox() is called.
    CUSTOM_LABEL = "— Custom —"
    preset_names = [CUSTOM_LABEL] + list(sat_info["presets"].keys())

    # Read the current R/G/B band IDs from session state (set by previous render)
    cur_r = band_map.get(st.session_state.get("se_r", band_names[0]), band_ids[0])
    cur_g = band_map.get(st.session_state.get("se_g", band_names[0]), band_ids[0])
    cur_b = band_map.get(st.session_state.get("se_b", band_names[0]), band_ids[0])

    stored_preset = st.session_state.get("se_preset", "True Color")
    prev_preset   = st.session_state.get("se_prev_preset")

    # Only check for a channel/preset mismatch when the preset is STABLE —
    # i.e. the user did not just change it in this render cycle.
    # If the preset just changed, the R/G/B channels still hold the OLD values
    # (they haven't been updated yet). Running the check at that moment would
    # wrongly detect a mismatch and flip straight to Custom.
    preset_just_changed = (stored_preset != prev_preset)

    if not preset_just_changed and stored_preset != CUSTOM_LABEL and stored_preset in sat_info["presets"]:
        pv = sat_info["presets"][stored_preset]
        if cur_r != pv["r"] or cur_g != pv["g"] or cur_b != pv["b"]:
            st.session_state.se_preset      = CUSTOM_LABEL
            st.session_state.se_prev_preset = CUSTOM_LABEL

    col_preset, col_r, col_g, col_b = st.columns([2, 1, 1, 1])

    with col_preset:
        # Default index points to True Color, not Custom (index 0).
        # The index is only used on first load when se_preset is not yet
        # in session state. After that, session state controls the value.
        tc_idx = preset_names.index("True Color") if "True Color" in preset_names else 1
        chosen_preset = st.selectbox(
            "Named preset (auto-fills R/G/B below)",
            preset_names,
            index=tc_idx,
            key="se_preset",
        )

    # When a real preset is chosen, push its bands into session state before
    # the R/G/B dropdowns render so they display the updated values immediately.
    preset_switched = (chosen_preset != st.session_state.se_prev_preset)
    if preset_switched and chosen_preset != CUSTOM_LABEL:
        st.session_state.se_prev_preset = chosen_preset
        preset_vals = sat_info["presets"][chosen_preset]
        if preset_vals["r"] in id_to_display:
            st.session_state.se_r = id_to_display[preset_vals["r"]]
        if preset_vals["g"] in id_to_display:
            st.session_state.se_g = id_to_display[preset_vals["g"]]
        if preset_vals["b"] in id_to_display:
            st.session_state.se_b = id_to_display[preset_vals["b"]]
    elif preset_switched:
        st.session_state.se_prev_preset = chosen_preset

    # Caption: description when a preset is active, instruction when Custom
    if chosen_preset != CUSTOM_LABEL:
        st.caption(f"ℹ️ {sat_info['presets'][chosen_preset]['note']}")
    else:
        st.caption("ℹ️ Channels set manually. Pick a named preset above to reset to a standard combination.")

    with col_r:
        r_display = st.selectbox("Red channel",   band_names, key="se_r")
        r_band    = band_map[r_display]
    with col_g:
        g_display = st.selectbox("Green channel", band_names, key="se_g")
        g_band    = band_map[g_display]
    with col_b:
        b_display = st.selectbox("Blue channel",  band_names, key="se_b")
        b_band    = band_map[b_display]

    st.caption(
        f"🔴 **{r_band}:** {sat_info['bands'][r_band]['description']}  |  "
        f"🟢 **{g_band}:** {sat_info['bands'][g_band]['description']}  |  "
        f"🔵 **{b_band}:** {sat_info['bands'][b_band]['description']}"
    )

    # Render and Compare buttons sit here — after the band controls,
    # because they act on the band selection the user just made.
    st.divider()
    col_btn1, col_btn2 = st.columns(2)
    with col_btn1:
        render_btn  = st.button("🖼️ Render selected combination", key="se_render",  use_container_width=True)
    with col_btn2:
        contact_btn = st.button("📋 Compare All Views",           key="se_contact", use_container_width=True)

    # -----------------------------------------------------------------------
    # RENDER — runs only when button clicked, stores result in session state
    # -----------------------------------------------------------------------
    if render_btn:
        # Use the user's chosen scene if set; otherwise fall back to best/first
        _sel_id  = st.session_state.get("se_selected_item_id")
        _all     = st.session_state.se_items
        item     = next((it for it in _all if it.id == _sel_id), None)
        if item is None:
            item = st.session_state.se_best_item
        if item is None and _all:
            item = _all[0]

        if item is None:
            st.warning("Search for scenes first.")
        else:
            scene_date = item.datetime.strftime("%B %d, %Y")
            cloud      = item.properties.get("eo:cloud_cover", 0)

            with st.spinner(f"Rendering {r_band}/{g_band}/{b_band} for {scene_date}..."):
                arr = spectral_explorer.render_combination(item, r_band, g_band, b_band, sat_key, width=700, bbox=bbox_se)

            if arr is not None and arr.max() > 0:
                # Store everything needed to redisplay after any widget interaction
                st.session_state.se_rendered_arr  = arr
                st.session_state.se_rendered_info = {
                    "scene_date":    scene_date,
                    "cloud":         cloud,
                    "r_band":        r_band,
                    "g_band":        g_band,
                    "b_band":        b_band,
                    "sat_key":       sat_key,
                    "location_label":location_label,
                    "preset_name":   chosen_preset,
                    "item_date":     item.datetime.strftime("%Y-%m-%d"),
                }
            else:
                st.warning("No valid data returned. Try a different scene or location.")

    # -----------------------------------------------------------------------
    # DISPLAY RENDERED IMAGE — always shown if a result exists in session state
    # -----------------------------------------------------------------------
    if st.session_state.se_rendered_arr is not None:
        info = st.session_state.se_rendered_info
        arr  = st.session_state.se_rendered_arr

        # Use stored info so it does not change when user fiddles with dropdowns
        r_b   = info["r_band"];   g_b  = info["g_band"];   b_b  = info["b_band"]
        s_key = info["sat_key"];  s_inf = satellite_catalog.SATELLITES[s_key]

        st.divider()
        col_img, col_info = st.columns([3, 2])

        with col_img:
            st.image(arr, caption=f"{r_b} / {g_b} / {b_b} — {info['scene_date']}", use_container_width=True)

            from PIL import Image as PILImage
            import io as _io
            buf = _io.BytesIO()
            PILImage.fromarray(arr).save(buf, format="PNG")
            fname = (
                f"{info['location_label'].replace(' ','_').replace(',','')}_"
                f"{s_key.replace(' ','_').replace('/','_')}_"
                f"{r_b}-{g_b}-{b_b}_{info['item_date']}.png"
            )
            st.download_button("⬇️ Download image", data=buf.getvalue(),
                               file_name=fname, mime="image/png", key="se_dl_single")

        with col_info:
            st.markdown("**Scene details**")
            st.markdown(f"- **Date:** {info['scene_date']}")
            st.markdown(f"- **Cloud cover:** {info['cloud']:.1f}%")
            st.markdown(f"- **Satellite:** {s_key}")
            st.markdown(f"- **Location:** {info['location_label']}")
            st.markdown(f"- **R channel:** {r_b} — {s_inf['bands'][r_b]['name']} ({s_inf['bands'][r_b]['wavelength']})")
            st.markdown(f"- **G channel:** {g_b} — {s_inf['bands'][g_b]['name']} ({s_inf['bands'][g_b]['wavelength']})")
            st.markdown(f"- **B channel:** {b_b} — {s_inf['bands'][b_b]['name']} ({s_inf['bands'][b_b]['wavelength']})")
            st.divider()
            st.markdown("**What this combination reveals**")
            st.markdown(f"🔴 {s_inf['bands'][r_b]['description']}")
            st.markdown(f"🟢 {s_inf['bands'][g_b]['description']}")
            st.markdown(f"🔵 {s_inf['bands'][b_b]['description']}")

            if config.has_any_key():
                st.divider()
                if st.button("🤖 AI: Explain this view", key="se_explain_single"):
                    with st.spinner("Asking AI..."):
                        explanation = spectral_explorer.explain_combination(
                            r_b, g_b, b_b, s_key, info["location_label"]
                        )
                    st.markdown(explanation)

        st.divider()
        render_what_am_i_looking_at(
            s_key, r_b, g_b, b_b,
            preset_name=info["preset_name"],
            location_label=info["location_label"],
            is_contact_sheet=False,
        )

    # -----------------------------------------------------------------------
    # COMPARE ALL VIEWS — runs only when button clicked
    # -----------------------------------------------------------------------
    if contact_btn:
        _sel_id  = st.session_state.get("se_selected_item_id")
        _all     = st.session_state.se_items
        item     = next((it for it in _all if it.id == _sel_id), None)
        if item is None:
            item = st.session_state.se_best_item
        if item is None and _all:
            item = _all[0]

        if item is None:
            st.warning("Search for scenes first.")
        else:
            with st.spinner("Rendering all combinations (30-60 seconds)..."):
                results = spectral_explorer.render_contact_sheet(item, sat_key, bbox=bbox_se)

            valid = [r for r in results if r["array"] is not None and r["array"].max() > 0]
            st.session_state.se_contact_results = valid
            st.session_state.se_contact_info    = {
                "scene_date":    item.datetime.strftime("%B %d, %Y"),
                "sat_key":       sat_key,
                "location_label":location_label,
            }

    # -----------------------------------------------------------------------
    # DISPLAY COMPARE ALL VIEWS — always shown if results exist in session state
    # -----------------------------------------------------------------------
    if st.session_state.se_contact_results:
        c_info  = st.session_state.se_contact_info
        valid   = st.session_state.se_contact_results

        st.divider()
        st.markdown(
            f"**All spectral views — {c_info['location_label']} — "
            f"{c_info['scene_date']} — {c_info['sat_key']}**"
        )
        st.caption("Every named band combination + NDVI + NDWI + NDMI + NBR + SAVI + EVI + BSI rendered for comparison.")

        cols_per_row = 3
        for i in range(0, len(valid), cols_per_row):
            row_items = valid[i:i + cols_per_row]
            cols = st.columns(cols_per_row)
            for col, result in zip(cols, row_items):
                with col:
                    st.image(result["array"], caption=result["label"], use_container_width=True)
                    st.caption(result["note"])
                    if result.get("channels"):
                        st.caption(f"📡 {result['channels']}")

    # Stop here — do not render the EO Explorer below
    st.stop()

# ---------------------------------------------------------------------------
# MODULE 3 — SAR Explorer (reached only when SAR Explorer is selected)
# ---------------------------------------------------------------------------

if selected_module == "📡 SAR Explorer":

    gee_available = st.session_state.gee_available

    st.subheader("📡 SAR Explorer")
    st.caption(
        "Select a location and two dates. "
        "The module fetches real Sentinel-1 radar data and shows how the surface changed."
    )

    with st.expander("ℹ️ What does this module do?", expanded=False):
        st.markdown("""
**This module answers the question: what can radar see that cameras cannot?**

It fetches Sentinel-1 SAR (Synthetic Aperture Radar) data from Google Earth Engine
for any location on Earth and compares two dates — even dates when cloud cover would
make optical satellite imagery completely unusable.

**What SAR is:**

SAR satellites transmit their own microwave pulses toward the ground and record the
energy that bounces back. Microwaves pass straight through clouds and work at night.
This makes SAR the sensor of choice for flood mapping, ship detection, infrastructure
monitoring, and tropical deforestation — anywhere optical sensors are blocked.

**What you will see after running an analysis:**

- **VV Polarization map** — surface roughness and vertical structures. Urban areas and
  ships appear bright. Calm water appears very dark.
- **VH Polarization map** — vegetation and volume scattering. Forest and crops appear
  relatively brighter than in VV.
- **False Color composite** — VV assigned to Red, VH to Green, VV/VH ratio to Blue.
  Pink = urban, green = vegetation, dark = water, white = ships.
- **Change Map** — VV difference between the two dates. Blue = backscatter increased
  (new structures or ships arrived). Red = backscatter decreased (ships left or surface smoothed).
- **Backscatter statistics** — VV and VH min/max/mean in dB for both dates
- **AI interpretation** — plain-language explanation of what the values reveal

**How to use it:**

1. Type any location in the search box
2. Pick two dates (the module searches a 10-day window around each date)
3. Click Run Analysis
        """)

    # --- Controls ---
    col_loc, col_d1, col_d2 = st.columns([3, 1, 1])

    with col_loc:
        sar_place = st.text_input(
            "Location — port, city, coastline, or river delta (not a whole country)",
            placeholder="e.g. Maasvlakte Rotterdam   |   Lagos harbour   |   Ganges delta   |   Singapore Strait",
            key="sar_place",
        )

    with col_d1:
        from datetime import date as _date, timedelta as _td
        sar_date1 = st.date_input(
            "Date 1",
            value=_date(2024, 1, 10),
            min_value=_date(2014, 1, 1),
            max_value=_date.today(),
            key="sar_date1",
        )

    with col_d2:
        sar_date2 = st.date_input(
            "Date 2",
            value=_date(2024, 3, 15),
            min_value=_date(2014, 1, 1),
            max_value=_date.today(),
            key="sar_date2",
        )

    _, col_run = st.columns([4, 1])
    with col_run:
        sar_run_btn = st.button(
            "▶ Run Analysis", type="primary",
            use_container_width=True, key="sar_run",
        )

    # --- Geocode location ---
    sar_bbox        = None
    sar_region_name = ""

    if sar_place.strip():
        cached_place = st.session_state.get("sar_geocoded_place", "")
        cached_bbox  = st.session_state.get("sar_geocoded_bbox",  None)

        if sar_place.strip() != cached_place:
            # New location typed — clear any stale map click
            map_picker.clear_click("sar")
            with st.spinner(f"Looking up '{sar_place}'..."):
                result_bbox = geocoder.geocode_place(sar_place)
            if result_bbox:
                st.session_state.sar_geocoded_place = sar_place.strip()
                st.session_state.sar_geocoded_bbox  = result_bbox
                cached_bbox = result_bbox
            else:
                st.session_state.sar_geocoded_place = ""
                st.session_state.sar_geocoded_bbox  = None
                cached_bbox = None
                st.error(f"Could not find '{sar_place}'. Try a broader name.")

        if cached_bbox:
            sar_bbox        = cached_bbox
            sar_region_name = sar_place.strip()

            # Show the area size so the user can judge before running
            w_km, h_km = geocoder.bbox_dims_km(sar_bbox)
            st.caption(f"📍 {sar_region_name} — {w_km:.0f} km × {h_km:.0f} km")

            # Warn if either dimension exceeds the useful SAR display threshold.
            # Sentinel-1 at 10-20 m resolution rendered into a 400px image:
            # anything wider than ~150 km loses all meaningful detail.
            SAR_MAX_KM = 150
            if w_km > SAR_MAX_KM or h_km > SAR_MAX_KM:
                st.warning(
                    f"**This area is too large for SAR detail.** "
                    f"{sar_region_name} spans {w_km:.0f} km × {h_km:.0f} km. "
                    f"Sentinel-1 resolution is 10-20 metres. At this scale, "
                    f"individual features — ships, buildings, rivers — are invisible. "
                    f"Try a specific port, harbour, city district, or coastal stretch "
                    f"no wider than {SAR_MAX_KM} km."
                )

    # Map picker — optional override shown once a geocoded centre exists
    if sar_bbox:
        with st.expander("📍 Refine location — click map to set exact area", expanded=False):
            picked = map_picker.render_map_picker(
                centre_bbox     = sar_bbox,
                picker_key      = "sar",
                default_size_km = 50,
            )
            if picked:
                sar_bbox = picked

    # GEE status
    if gee_available:
        st.caption("🟢 GEE connected — live Sentinel-1 data active.")
    else:
        st.caption("🔴 GEE not connected. SAR Explorer requires live GEE data.")

    st.divider()

    # --- Session state initialisation ---
    for _k, _v in [
        ("sar_maps",          None), ("sar_stats1",       None), ("sar_stats2", None),
        ("sar_result_region", None), ("sar_result_date1", None),
        ("sar_result_date2",  None), ("sar_ai_result",    None),
    ]:
        if _k not in st.session_state:
            st.session_state[_k] = _v

    # --- Two-step run pattern (same as Time Series Explorer) ---
    if sar_run_btn:
        if not sar_bbox:
            st.warning("Enter a location first.")
        elif not gee_available:
            st.error("SAR Explorer requires GEE credentials. Add GEE_SERVICE_ACCOUNT_JSON to Streamlit secrets.")
        elif sar_date1 >= sar_date2:
            st.error("Date 1 must be before Date 2.")
        else:
            st.session_state.sar_maps        = None
            st.session_state.sar_stats1      = None
            st.session_state.sar_stats2      = None
            st.session_state.sar_ai_result   = None
            st.session_state.sar_pending_run = {
                "bbox":   sar_bbox,
                "date1":  str(sar_date1),
                "date2":  str(sar_date2),
                "region": sar_region_name,
            }
            st.rerun()

    if st.session_state.get("sar_pending_run"):
        p = st.session_state.sar_pending_run
        st.session_state.sar_pending_run = None

        with st.spinner(f"Fetching Sentinel-1 images for {p['region']}..."):
            img1, count1, bbox1 = gee_sar.fetch_sar_image(p["bbox"], p["date1"], gee_available)
            img2, count2, bbox2 = gee_sar.fetch_sar_image(p["bbox"], p["date2"], gee_available)

        # Use the padded bbox for all downstream calls.
        # fetch_sar_image expands point geocodes (e.g. "Port of Rotterdam") to
        # MIN_BBOX_DEG so the SAR image covers a usable area.
        eff_bbox = bbox1 if bbox1 else bbox2

        if img1 is None or img2 is None:
            st.error(
                f"No Sentinel-1 images found near the selected dates for this location. "
                f"Try dates that are further apart, or a different location. "
                f"(Date 1: {count1} scenes found, Date 2: {count2} scenes found)"
            )
        else:
            with st.spinner("Computing backscatter statistics..."):
                stats1 = gee_sar.get_backscatter_stats(img1, eff_bbox)
                stats2 = gee_sar.get_backscatter_stats(img2, eff_bbox)

            with st.spinner("Building interactive SAR map (this takes 20-40 seconds)..."):
                sar_map = gee_sar.build_sar_map(img1, img2, eff_bbox, p["date1"], p["date2"])

            st.session_state.sar_maps         = sar_map
            st.session_state.sar_stats1       = stats1
            st.session_state.sar_stats2       = stats2
            st.session_state.sar_result_region = p["region"]
            st.session_state.sar_result_date1  = p["date1"]
            st.session_state.sar_result_date2  = p["date2"]
            st.success(
                f"Analysis complete — {p['date1']} and {p['date2']} over {p['region']}."
            )

    # --- Display results ---
    if st.session_state.sar_maps is not None:
        maps   = st.session_state.sar_maps
        stats1 = st.session_state.sar_stats1
        stats2 = st.session_state.sar_stats2
        r_reg  = st.session_state.sar_result_region
        r_d1   = st.session_state.sar_result_date1
        r_d2   = st.session_state.sar_result_date2

        def sar_section_break():
            st.markdown(
                '<hr style="border: none; border-top: 3px solid #d0d0d0; margin: 28px 0 20px 0;">',
                unsafe_allow_html=True,
            )

        # --- SECTION 1: Backscatter statistics ---
        st.subheader("📊 Backscatter Statistics")
        col_s1, col_s2 = st.columns(2)

        with col_s1:
            st.markdown(f"**Date 1 — {r_d1}**")
            c1, c2, c3 = st.columns(3)
            c1.metric("VV min",  f"{stats1['VV_min']} dB")
            c2.metric("VV mean", f"{stats1['VV_mean']} dB")
            c3.metric("VV max",  f"{stats1['VV_max']} dB")
            c4, c5, c6 = st.columns(3)
            c4.metric("VH min",  f"{stats1['VH_min']} dB")
            c5.metric("VH mean", f"{stats1['VH_mean']} dB")
            c6.metric("VH max",  f"{stats1['VH_max']} dB")

        with col_s2:
            st.markdown(f"**Date 2 — {r_d2}**")
            c1, c2, c3 = st.columns(3)
            vv_delta = stats2['VV_mean'] - stats1['VV_mean']
            vh_delta = stats2['VH_mean'] - stats1['VH_mean']
            c1.metric("VV min",  f"{stats2['VV_min']} dB")
            c2.metric("VV mean", f"{stats2['VV_mean']} dB", delta=f"{vv_delta:+.1f} dB")
            c3.metric("VV max",  f"{stats2['VV_max']} dB")
            c4, c5, c6 = st.columns(3)
            c4.metric("VH min",  f"{stats2['VH_min']} dB")
            c5.metric("VH mean", f"{stats2['VH_mean']} dB", delta=f"{vh_delta:+.1f} dB")
            c6.metric("VH max",  f"{stats2['VH_max']} dB")

        sar_section_break()

        # --- SECTION 2: Stats chart ---
        st.subheader("📈 Backscatter Comparison Chart")
        fig_stats = gee_sar.build_stats_chart(stats1, stats2, r_d1, r_d2)
        st.plotly_chart(fig_stats, use_container_width=True)

        st.caption(
            "**Reading the chart:** Values near -20 dB = calm water. "
            "-10 to -15 dB = vegetation. Above -5 dB = urban areas or ships. "
            "VV responds to vertical structures. VH responds to vegetation volume."
        )

        sar_section_break()

        # --- SECTION 3: Interactive SAR Map ---
        # GEE tile URLs include a time-limited auth token in the URL itself.
        # The browser can fetch them directly — no special headers needed.
        # This gives us full zoom/pan interactivity via Folium TileLayers.
        st.subheader("🗺️ SAR Views")

        st.caption(
            "Toggle layers using the control in the top-right corner of the map. "
            "VV Date 1 is shown by default. "
            "**VV** = surface roughness, ships, buildings bright. "
            "**VH** = vegetation volume bright. "
            "**False Color** = Red:VV Green:VH Blue:ratio — pink=urban, green=vegetation, dark=water. "
            "**Change Map** = blue: backscatter increased, red: decreased."
        )

        if maps is not None:
            st_folium(maps, width=700, height=500, returned_objects=[])
        else:
            st.warning("SAR map could not be built. Check GEE connection.")

        sar_section_break()

        # --- SECTION 4: AI Interpretation ---
        st.subheader("🤖 AI Interpretation")
        if st.button("Get AI Interpretation", type="primary", key="sar_ai_btn"):
            with st.spinner("Thinking..."):
                interpretation = gee_sar.get_sar_interpretation(
                    stats1, stats2, r_d1, r_d2, r_reg,
                    groq_key=config.GROQ_API_KEY,
                    gemini_key=config.GEMINI_API_KEY,
                )
                st.session_state.sar_ai_result = interpretation

        if st.session_state.get("sar_ai_result"):
            st.markdown(st.session_state.sar_ai_result)

    else:
        st.markdown("---")
        st.markdown(
            "**Type a location above, pick two dates, then click Run Analysis.**\n\n"
            "The module will fetch live Sentinel-1 radar data and show VV, VH, "
            "false color, and change maps alongside backscatter statistics."
        )

    st.stop()

# ---------------------------------------------------------------------------
# MODULE 4 — Change Detection (reached only when Change Detection is selected)
# ---------------------------------------------------------------------------

if selected_module == "🔀 Change Detection":

    from streamlit_folium import st_folium as _st_folium
    from datetime import date as _date

    gee_available = st.session_state.gee_available

    st.subheader("🔀 Change Detection")
    st.caption(
        "Select a location and two dates. "
        "The module computes the NDVI difference and shows where vegetation gained or lost."
    )

    with st.expander("ℹ️ What does this module do?", expanded=False):
        st.markdown("""
**This module answers the question: what changed between two dates?**

It fetches Sentinel-2 satellite imagery for two dates and computes the NDVI difference —
a pixel-by-pixel comparison of vegetation greenness. Where NDVI increased, vegetation
grew or recovered. Where NDVI decreased, vegetation died, burned, or was removed.

**How to use it:**

1. Type any location in the search box
2. Pick two dates (the module searches a 30-day window around each date for cloud-free imagery)
3. Click Run Analysis

**Data source:** Sentinel-2 Surface Reflectance via Google Earth Engine.
MODIS is used as a fallback if Sentinel-2 has insufficient cloud-free coverage.

---

**What is NDVI?**

NDVI stands for Normalized Difference Vegetation Index. It is the most widely used
measure of vegetation health from satellite data.

Formula: **(NIR − Red) / (NIR + Red)**

Healthy plants absorb red light for photosynthesis and reflect near-infrared strongly.
Bare soil reflects both at similar levels. Water absorbs both almost completely.

| Surface | Typical NDVI |
|---|---|
| Dense healthy forest | 0.6 – 0.9 |
| Crops mid-season | 0.4 – 0.7 |
| Sparse or stressed vegetation | 0.2 – 0.4 |
| Bare soil | 0.1 – 0.2 |
| Water | 0 or negative |

---

**How to read the statistics**

**Mean NDVI — Date 1 and Date 2 (baseline and endpoint)**
These are the average NDVI values across the whole region for each date.
They give the change context. A shift of +0.3 from a baseline of 0.1 (bare soil recovering)
is very different from +0.3 from a baseline of 0.5 (already-dense forest getting denser).
Without the baseline, the difference alone can mislead.

**Mean NDVI change**
The average pixel-by-pixel difference: Date 2 minus Date 1.
Positive = net greening across the region. Negative = net browning.

**Spatial variability (standard deviation of change)**
This is the single most diagnostic number for understanding *why* a region changed.
A low standard deviation means every pixel moved in the same direction by roughly the same
amount — almost always a seasonal signal (wet season arriving, dry season setting in).
A high standard deviation means change is concentrated in patches — more likely fire,
deforestation, flood inundation, or agricultural clearing.

**Net change (gain area minus loss area)**
One number for a briefing headline. Tells you whether the region is net gaining or
net losing vegetation, in km².

**Gain / Loss ratio**
Gain area divided by loss area. Above 1 means more land gained vegetation than lost it.
Below 1 means the opposite. Useful for comparing regions or time periods.

**Gain and Loss areas (±0.1 threshold)**
Area in km² where NDVI changed by more than ±0.1. This threshold separates real
vegetation change from sensor noise and minor cloud contamination artefacts.

**Stable area**
Area where NDVI changed by less than ±0.1. This is the land that did not meaningfully
change between the two dates.

**Extreme gain and Extreme loss (±0.3 threshold)**
Area where NDVI changed by more than ±0.3. This separates dramatic events from ordinary
seasonal movement. Large extreme gain areas suggest post-fire regrowth, crop establishment
after irrigation, or a very strong wet season flush. Large extreme loss areas suggest
fire, clearcut deforestation, or severe drought stress.

---

**Change map colour guide**

- 🟢 **Green** — vegetation increased between Date 1 and Date 2
- 🔴 **Red** — vegetation decreased
- ⬜ **White** — no significant change

Toggle between NDVI Date 1, NDVI Date 2, and the change map using the layer control
in the top-right corner of the map.
        """)

    # --- Controls ---
    col_loc, col_d1, col_d2 = st.columns([3, 1, 1])

    with col_loc:
        cd_place = st.text_input(
            "Location — type any region, country, or ecosystem",
            placeholder="e.g. Sahel, West Africa   |   Amazon, Brazil   |   California, USA",
            key="cd_place",
        )

    with col_d1:
        cd_date1 = st.date_input(
            "Date 1",
            value=_date(2023, 2, 1),
            min_value=_date(2017, 1, 1),
            max_value=_date.today(),
            key="cd_date1",
        )

    with col_d2:
        cd_date2 = st.date_input(
            "Date 2",
            value=_date(2023, 9, 1),
            min_value=_date(2017, 1, 1),
            max_value=_date.today(),
            key="cd_date2",
        )

    _, col_run = st.columns([4, 1])
    with col_run:
        cd_run_btn = st.button(
            "▶ Run Analysis", type="primary",
            use_container_width=True, key="cd_run",
        )

    # --- Geocode ---
    cd_bbox        = None
    cd_region_name = ""

    if cd_place.strip():
        cached_place = st.session_state.get("cd_geocoded_place", "")
        cached_bbox  = st.session_state.get("cd_geocoded_bbox",  None)

        if cd_place.strip() != cached_place:
            # New location typed — clear any stale map click
            map_picker.clear_click("cd")
            with st.spinner(f"Looking up '{cd_place}'..."):
                result_bbox = geocoder.geocode_place(cd_place)
            if result_bbox:
                st.session_state.cd_geocoded_place = cd_place.strip()
                st.session_state.cd_geocoded_bbox  = result_bbox
                cached_bbox = result_bbox
            else:
                st.session_state.cd_geocoded_place = ""
                st.session_state.cd_geocoded_bbox  = None
                cached_bbox = None
                st.error(f"Could not find '{cd_place}'. Try a broader name.")

        if cached_bbox:
            cd_bbox        = cached_bbox
            cd_region_name = cd_place.strip()
            st.caption(f"📍 {cd_region_name}")

    # Map picker — optional override shown once a geocoded centre exists
    if cd_bbox:
        with st.expander("📍 Refine location — click map to set exact area", expanded=False):
            picked = map_picker.render_map_picker(
                centre_bbox     = cd_bbox,
                picker_key      = "cd",
                default_size_km = 100,
            )
            if picked:
                cd_bbox = picked

    if gee_available:
        st.caption("🟢 GEE connected — live Sentinel-2 / MODIS data active.")
    else:
        st.caption("🔴 GEE not connected. Change Detection requires live GEE data.")

    st.divider()

    # --- Session state ---
    for _k, _v in [
        ("cd_map",           None), ("cd_stats",         None),
        ("cd_result_region", None), ("cd_result_date1",  None),
        ("cd_result_date2",  None), ("cd_result_src1",   None),
        ("cd_result_src2",   None), ("cd_ai_result",     None), ("cd_ai_model", None),
    ]:
        if _k not in st.session_state:
            st.session_state[_k] = _v

    # --- Two-step run pattern ---
    if cd_run_btn:
        if not cd_bbox:
            st.warning("Enter a location first.")
        elif not gee_available:
            st.error("Change Detection requires GEE credentials. Add GEE_SERVICE_ACCOUNT_JSON to Streamlit secrets.")
        elif cd_date1 >= cd_date2:
            st.error("Date 1 must be before Date 2.")
        else:
            st.session_state.cd_map        = None
            st.session_state.cd_stats      = None
            st.session_state.cd_ai_result  = None
            st.session_state.cd_pending_run = {
                "bbox":   cd_bbox,
                "date1":  str(cd_date1),
                "date2":  str(cd_date2),
                "region": cd_region_name,
            }
            st.rerun()

    if st.session_state.get("cd_pending_run"):
        p = st.session_state.cd_pending_run
        st.session_state.cd_pending_run = None

        with st.spinner(f"Fetching NDVI Date 1 ({p['date1']}) for {p['region']}..."):
            img1, src1 = gee_change.fetch_ndvi_image(p["bbox"], p["date1"], gee_available)

        with st.spinner(f"Fetching NDVI Date 2 ({p['date2']}) for {p['region']}..."):
            img2, src2 = gee_change.fetch_ndvi_image(p["bbox"], p["date2"], gee_available)

        if img1 is None or img2 is None:
            st.error(
                "Could not retrieve NDVI imagery for one or both dates. "
                "Try a wider date range, a different location, or check GEE credentials."
            )
        else:
            with st.spinner("Computing change statistics..."):
                diff_img = img2.subtract(img1).rename("NDVI_diff")
                stats = gee_change.compute_change_stats(img1, img2, diff_img, p["bbox"], gee_available)

            with st.spinner("Building change map (20-40 seconds)..."):
                cd_map = gee_change.build_change_map(
                    img1, img2, p["bbox"], p["date1"], p["date2"], gee_available
                )

            st.session_state.cd_map          = cd_map
            st.session_state.cd_stats        = stats
            st.session_state.cd_result_region = p["region"]
            st.session_state.cd_result_date1  = p["date1"]
            st.session_state.cd_result_date2  = p["date2"]
            st.session_state.cd_result_src1   = src1
            st.session_state.cd_result_src2   = src2
            st.success(
                f"Analysis complete — {p['date1']} and {p['date2']} over {p['region']}. "
                f"Date 1 source: {src1}. Date 2 source: {src2}."
            )

    # --- Display results ---
    if st.session_state.cd_stats is not None:
        stats  = st.session_state.cd_stats
        r_reg  = st.session_state.cd_result_region
        r_d1   = st.session_state.cd_result_date1
        r_d2   = st.session_state.cd_result_date2

        def cd_section_break():
            st.markdown(
                '<hr style="border: none; border-top: 3px solid #d0d0d0; margin: 28px 0 20px 0;">',
                unsafe_allow_html=True,
            )

        # --- SECTION 1: Summary statistics ---
        st.subheader("📊 Change Statistics")

        # Row 1: Headline change numbers
        c1, c2, c3, c4 = st.columns(4)
        delta_color = "normal" if stats["mean_change"] >= 0 else "inverse"
        c1.metric(
            "Mean NDVI change",
            f"{stats['mean_change']:+.4f}",
            delta="Greening" if stats["mean_change"] > 0 else "Browning",
            delta_color=delta_color,
        )
        c2.metric(
            "NDVI — Date 1 (baseline)",
            f"{stats['mean_ndvi1']:.3f}",
        )
        c3.metric(
            "NDVI — Date 2 (endpoint)",
            f"{stats['mean_ndvi2']:.3f}",
            delta=f"{stats['mean_change']:+.3f}",
            delta_color=delta_color,
        )
        c4.metric(
            "Spatial variability (std dev)",
            f"{stats['std_change']:.4f}",
            help="Low = uniform change (typically seasonal). High = patchy change (fire, clearing, flood).",
        )

        st.markdown("")

        # Row 2: Area breakdown
        thr = gee_change.CHANGE_THRESHOLD
        ext = gee_change.EXTREME_THRESHOLD
        a1, a2, a3, a4 = st.columns(4)
        a1.metric(
            f"🟢 Gain area (>{thr} NDVI)",
            f"{stats['area_gain_km2']:,.0f} km²",
            delta=f"{stats['pct_gain']:.1f}% of region",
        )
        a2.metric(
            f"🔴 Loss area (<-{thr} NDVI)",
            f"{stats['area_loss_km2']:,.0f} km²",
            delta=f"{stats['pct_loss']:.1f}% of region",
        )
        a3.metric(
            "⬜ Stable area",
            f"{stats['area_stable_km2']:,.0f} km²",
            delta=f"{stats['pct_stable']:.1f}% of region",
        )
        a4.metric(
            "Total area analysed",
            f"{stats['area_total_km2']:,.0f} km²",
        )

        st.markdown("")

        # Row 3: Derived indicators
        ratio = stats["gain_loss_ratio"]
        ratio_str = f"{ratio:.2f}×" if ratio is not None else "∞ (no loss)"
        net_dir   = "net gain" if stats["net_change_km2"] >= 0 else "net loss"
        b1, b2, b3, b4 = st.columns(4)
        b1.metric(
            "Net change (gain − loss)",
            f"{stats['net_change_km2']:+,.0f} km²",
            delta=net_dir,
            delta_color="normal" if stats["net_change_km2"] >= 0 else "inverse",
        )
        b2.metric(
            "Gain / Loss ratio",
            ratio_str,
            help=">1 means more land gained vegetation than lost it.",
        )
        b3.metric(
            f"⚡ Extreme gain (>{ext} NDVI)",
            f"{stats['area_extreme_gain_km2']:,.0f} km²",
            help=f"Area where NDVI increased by more than {ext} — regrowth, crop establishment, or post-rain flush.",
        )
        b4.metric(
            f"⚡ Extreme loss (<-{ext} NDVI)",
            f"{stats['area_extreme_loss_km2']:,.0f} km²",
            help=f"Area where NDVI dropped by more than {ext} — likely fire, clearcut, or severe drought.",
        )

        cd_section_break()

        # --- SECTION 2: Change map ---
        st.subheader("🗺️ Change Map")
        st.caption(
            "Toggle layers using the control in the top-right corner of the map. "
            "**Change map** is shown by default: green = vegetation gain, red = vegetation loss, white = stable. "
            "Switch to NDVI Date 1 or Date 2 to see the raw greenness values."
        )

        if st.session_state.cd_map is not None:
            _st_folium(st.session_state.cd_map, width=700, height=520, returned_objects=[])
        else:
            st.warning("Map could not be built. Check GEE connection.")

        cd_section_break()

        # --- SECTION 3: AI Interpretation ---
        st.subheader("🤖 AI Interpretation")
        if st.button("Get AI Interpretation", type="primary", key="cd_ai_btn"):
            with st.spinner("Thinking..."):
                interpretation, model_used = gee_change.get_change_interpretation(
                    stats,
                    r_d1, r_d2, r_reg,
                    st.session_state.cd_result_src1,
                    st.session_state.cd_result_src2,
                    groq_key=config.GROQ_API_KEY,
                    gemini_key=config.GEMINI_API_KEY,
                )
                st.session_state.cd_ai_result = interpretation
                st.session_state.cd_ai_model  = model_used

        if st.session_state.get("cd_ai_result"):
            st.markdown(st.session_state.cd_ai_result)
            model_used = st.session_state.get("cd_ai_model")
            if model_used:
                st.caption(f"AI response from **{model_used}**")
            else:
                st.caption("Showing built-in fallback interpretation. Add GROQ_API_KEY or GEMINI_API_KEY to enable AI.")

        cd_section_break()

        # --- SECTION 4: Data Quality ---
        st.subheader("🔍 Data Quality")
        _src1 = st.session_state.get("cd_result_src1", "")
        _src2 = st.session_state.get("cd_result_src2", "")
        _both_s2 = "Sentinel-2" in str(_src1) and "Sentinel-2" in str(_src2)
        _cd_res  = "10 m (Sentinel-2)" if _both_s2 else "250 m (MODIS fallback)"
        if _both_s2:
            _cd_conf = "High — both dates used Sentinel-2 at 10 m. Change map resolves field-scale features."
        else:
            _cd_conf = ("Moderate — one or both dates used MODIS (250 m). "
                        "Change map shows regional patterns only. Fine-scale features are not resolved.")
        st.info(
            f"**Date 1 source:** {_src1 or 'GEE composite'}  \n"
            f"**Date 2 source:** {_src2 or 'GEE composite'}  \n"
            f"**Spatial resolution:** {_cd_res}  \n"
            f"**Cloud handling:** GEE composites — cloud pixels excluded during aggregation.  \n"
            f"**Scenes per date:** Multiple scenes composited within ±{gee_change.WINDOW_DAYS} days of each target date.  \n"
            f"**Confidence:** {_cd_conf}"
        )

        cd_section_break()

        # --- SECTION 5: Export ---
        st.subheader("📥 Export")
        import pandas as pd
        export_stats = {
            "Region":              r_reg,
            "Date 1":              str(r_d1),
            "Date 2":              str(r_d2),
            "Area increased km2":  round(stats["area_increased_km2"], 2),
            "Area decreased km2":  round(stats["area_decreased_km2"], 2),
            "Area stable km2":     round(stats["area_stable_km2"], 2),
            "Area total km2":      round(stats["area_total_km2"], 2),
            "Net change km2":      round(stats["net_change_km2"], 2),
            "Pct increased":       round(stats["pct_increased"], 2),
            "Pct decreased":       round(stats["pct_decreased"], 2),
            "Pct stable":          round(stats["pct_stable"], 2),
        }
        cd_df    = pd.DataFrame([export_stats])
        csv_data = cd_df.to_csv(index=False).encode()
        safe_reg = r_reg.replace(" ", "_").replace(",", "")
        fname_cd = f"change_detection_{safe_reg}_{r_d1}_{r_d2}.csv"
        st.download_button(
            label="⬇️ Download change statistics as CSV",
            data=csv_data,
            file_name=fname_cd,
            mime="text/csv",
            key="cd_export_csv",
        )
        st.caption("CSV contains area statistics for increased, decreased, and stable vegetation.")

    else:
        st.markdown("---")
        st.markdown(
            "**Type a location above, pick two dates, then click Run Analysis.**\n\n"
            "The module will compute the NDVI difference and show an interactive change map "
            "alongside summary statistics and an AI interpretation."
        )

    st.stop()

# ---------------------------------------------------------------------------
# MODULE 5 — AI Imagery Interpreter
# ---------------------------------------------------------------------------

if selected_module == "🔍 AI Imagery Interpreter":

    from datetime import date as _date

    st.subheader("🔍 AI Imagery Interpreter")
    st.caption(
        "Pick a location and a date range. The module finds the best available "
        "Sentinel-2 scene in that window and asks a vision AI to describe what it sees."
    )

    with st.expander("ℹ️ What does this module do?", expanded=False):
        st.markdown("""
**This module answers the question: what does this place look like from space — and what does it mean?**

Every other module in this portal sends numbers and statistics to an AI text model.
This module sends an actual satellite image to a vision model. The AI reads the picture
directly — identifying land cover, surface features, seasonal state, and what the image
could inform in a real-world context.

**What you will see after running an analysis:**

- **Sentinel-2 true-color image** — the actual satellite chip for your location and date,
  rendered in natural colour (Red=B4, Green=B3, Blue=B2)
- **Scene metadata** — exact acquisition date, cloud cover, and scene ID
- **Vision AI interpretation** — 3-4 paragraphs describing land cover types, notable
  features, ecological state, and potential decision uses
- **Model attribution** — which model in the 8-model vision chain produced the response

**How to use it:**

1. Type any location in the search box
2. Set a date range — this is the window searched for cloud-free Sentinel-2 scenes
3. The module picks the best available scene within your range (lowest cloud cover, best coverage)
4. Click Analyse
5. Wait 15-30 seconds while the chip is fetched and interpreted

**Tips for setting the date range:**
- A 1-3 month window gives enough scenes to find a clean one in most regions
- Narrow the range when you want a specific season (e.g. June-August for northern summer)
- Widen the range for persistently cloudy regions (tropics, maritime climates)

**How to read a true-color satellite image:**

- **Dark to bright green** — healthy vegetation. Dense forest is darker. Crops and grass are lighter.
- **Dark blue to black** — open water. Deep water absorbs almost all light.
- **Grey or white** — urban surfaces, bare concrete, or cloud cover.
- **Brown or tan** — bare soil, dry grassland, or harvested fields.
- **Bright white patches** — clouds or snow and ice.

**Vision chain:** 7 multimodal models in fallback order — 5 Gemini (gemini-2.5-flash first)
then 2 Groq Llama-4 models. Text-only models are excluded — they cannot receive images.
        """)

    # --- Controls ---
    col_loc, col_d1, col_d2 = st.columns([3, 1, 1])

    with col_loc:
        ii_place = st.text_input(
            "Location — type any city, region, or ecosystem",
            placeholder="e.g. Zanzibar   |   Mekong Delta   |   Sahara, Algeria   |   Swiss Alps",
            key="ii_place",
        )

    with col_d1:
        ii_date_start = st.date_input(
            "From date",
            value=_date(2024, 1, 1),
            min_value=_date(2017, 1, 1),
            max_value=_date.today(),
            key="ii_date_start",
        )

    with col_d2:
        ii_date_end = st.date_input(
            "To date",
            value=_date(2024, 4, 1),
            min_value=_date(2017, 1, 1),
            max_value=_date.today(),
            key="ii_date_end",
        )

    _, col_run = st.columns([4, 1])
    with col_run:
        ii_run_btn = st.button(
            "▶ Analyse", type="primary",
            use_container_width=True, key="ii_run",
        )

    # --- Geocode ---
    ii_bbox        = None
    ii_region_name = ""

    if ii_place.strip():
        cached_place = st.session_state.get("ii_geocoded_place", "")
        cached_bbox  = st.session_state.get("ii_geocoded_bbox",  None)

        if ii_place.strip() != cached_place:
            # New location typed — clear any stale map click
            map_picker.clear_click("ii")
            with st.spinner(f"Looking up '{ii_place}'..."):
                result_bbox = geocoder.geocode_place(ii_place)
            if result_bbox:
                st.session_state.ii_geocoded_place = ii_place.strip()
                st.session_state.ii_geocoded_bbox  = result_bbox
                cached_bbox = result_bbox
            else:
                st.session_state.ii_geocoded_place = ""
                st.session_state.ii_geocoded_bbox  = None
                cached_bbox = None
                st.error(f"Could not find '{ii_place}'. Try a broader name.")

        if cached_bbox:
            ii_bbox        = cached_bbox
            ii_region_name = ii_place.strip()
            area = geocoder.bbox_area_km2(cached_bbox)
            st.caption(f"📍 {ii_region_name} — {area:,.0f} km²")

            # Size warnings — imagery interpretation loses detail above ~25,000 km²
            # At 600px image width, 25,000 km² = ~8m per pixel (city-level detail lost)
            if area > 100_000:
                st.error(
                    f"**This area is too large for meaningful imagery interpretation.** "
                    f"{ii_region_name} spans {area:,.0f} km². "
                    f"At this scale the satellite chip will show an entire country — "
                    f"individual features are invisible. "
                    f"Try a city name, a national park, an island, or a specific coastal area. "
                    f"For river features, search for the nearest town instead "
                    f"(e.g. 'Jinja, Uganda' for the Nile source at Lake Victoria)."
                )
                ii_bbox = None  # block the run
            elif area > 25_000:
                st.warning(
                    f"**Large area — detail will be limited.** "
                    f"{ii_region_name} spans {area:,.0f} km². "
                    f"Small features like rivers, roads, and fields may not be visible. "
                    f"For better results, try a more specific location within this region."
                )

    # Map picker — optional override shown once a geocoded centre exists
    # Especially useful here: eliminates the geocoding ambiguity problem where
    # descriptive phrases (e.g. "Nile river north of lake victoria") return huge bboxes.
    if ii_bbox:
        with st.expander("📍 Refine location — click map to set exact area", expanded=False):
            picked = map_picker.render_map_picker(
                centre_bbox     = ii_bbox,
                picker_key      = "ii",
                default_size_km = 50,
            )
            if picked:
                ii_bbox = picked
                ii_region_name = ii_region_name  # keep the typed name as the label

    st.caption("🌐 Data source: Sentinel-2 L2A via Planetary Computer (Microsoft). No GEE required.")
    st.divider()

    # --- Session state ---
    for _k, _v in [
        ("ii_image_arr",    None), ("ii_metadata",       None),
        ("ii_result_region", None), ("ii_result_date",   None),
        ("ii_ai_result",    None), ("ii_ai_model",       None),
    ]:
        if _k not in st.session_state:
            st.session_state[_k] = _v

    # --- Two-step run pattern ---
    if ii_run_btn:
        if not ii_bbox:
            st.warning("Enter a location first.")
        elif ii_date_start >= ii_date_end:
            st.error("From date must be before To date.")
        else:
            st.session_state.ii_image_arr    = None
            st.session_state.ii_metadata     = None
            st.session_state.ii_ai_result    = None
            st.session_state.ii_ai_model     = None
            st.session_state.ii_pending_run  = {
                "bbox":       ii_bbox,
                "date_start": str(ii_date_start),
                "date_end":   str(ii_date_end),
                "region":     ii_region_name,
            }
            st.rerun()

    if st.session_state.get("ii_pending_run"):
        p = st.session_state.ii_pending_run
        st.session_state.ii_pending_run = None

        with st.spinner(f"Searching Planetary Computer for {p['region']} ({p['date_start']} to {p['date_end']})..."):
            arr, metadata = imagery_interpreter.fetch_chip(p["bbox"], p["date_start"], p["date_end"])

        if arr is None:
            st.error(
                f"No Sentinel-2 scene found for {p['region']} between "
                f"{p['date_start']} and {p['date_end']}. "
                "Try widening the date range or increasing cloud cover tolerance."
            )
        else:
            st.session_state.ii_image_arr    = arr
            st.session_state.ii_metadata     = metadata
            st.session_state.ii_result_region = p["region"]
            st.session_state.ii_result_date   = p["date_start"]

            # Immediately run vision AI interpretation
            with st.spinner(f"Asking vision AI to interpret the image of {p['region']}..."):
                image_bytes = imagery_interpreter.array_to_jpeg_bytes(arr)
                interpretation, model_used = imagery_interpreter.interpret_image(
                    image_bytes,
                    location   = p["region"],
                    date_str   = metadata["date"],
                    gemini_key = config.GEMINI_API_KEY,
                    groq_key   = config.GROQ_API_KEY,
                )

            if interpretation:
                st.session_state.ii_ai_result = interpretation
                st.session_state.ii_ai_model  = model_used
            else:
                st.session_state.ii_ai_result = imagery_interpreter.get_fallback_interpretation(
                    p["region"], metadata["date"]
                )
                st.session_state.ii_ai_model  = None

            st.success(
                f"Scene found: {metadata['date']} — cloud cover {metadata['cloud_cover']:.1f}%"
            )

    # --- Display results ---
    if st.session_state.ii_image_arr is not None:
        arr      = st.session_state.ii_image_arr
        metadata = st.session_state.ii_metadata
        r_reg    = st.session_state.ii_result_region
        r_date   = st.session_state.ii_result_date

        st.subheader("🛰️ Satellite Image")

        col_img, col_interp = st.columns([1, 1])

        with col_img:
            st.image(
                arr,
                caption=(
                    f"Sentinel-2 true-color — {r_reg} — {metadata['date']} "
                    f"(cloud {metadata['cloud_cover']:.1f}%)"
                ),
                use_container_width=True,
            )
            st.markdown("**Scene details**")
            st.markdown(f"- **Acquired:** {metadata['date']}")
            st.markdown(f"- **Cloud cover:** {metadata['cloud_cover']:.1f}%")
            st.markdown(f"- **Sensor:** Sentinel-2 L2A")
            st.markdown(f"- **Bands:** B4 (Red), B3 (Green), B2 (Blue) — true color")
            st.markdown(f"- **Scene ID:** `{metadata['scene_id'][:40]}...`")

        with col_interp:
            st.markdown("**🤖 Vision AI Interpretation**")
            if st.session_state.ii_ai_result:
                st.markdown(st.session_state.ii_ai_result)
                model_used = st.session_state.get("ii_ai_model")
                if model_used:
                    st.caption(f"Vision AI response from **{model_used}**")
                else:
                    st.caption("Showing built-in reading guide. Add GEMINI_API_KEY or GROQ_API_KEY to enable AI.")

    else:
        st.markdown("---")
        st.markdown(
            "**Type a location above, pick a date, then click Analyse.**\n\n"
            "The module will find the best available Sentinel-2 scene near your target date "
            "and ask a vision AI to describe what it sees."
        )

    st.stop()

# ---------------------------------------------------------------------------
# MODULE 6 — Emissions Explorer helpers (Word document builder)
# ---------------------------------------------------------------------------

def _lonlat_to_webmercator(lon, lat):
    """Convert WGS-84 lon/lat degrees to Web Mercator (EPSG:3857) metres."""
    import math
    x = lon * 20037508.34 / 180.0
    y = math.log(math.tan((90.0 + lat) * math.pi / 360.0)) / (math.pi / 180.0)
    y = y * 20037508.34 / 180.0
    return x, y


def _em_static_map_bytes(bbox, region_name=""):
    """Render a basemap tile map with the analysis region highlighted.

    Uses contextily to fetch OpenStreetMap tiles so the output looks like
    the interactive Folium map (country borders, place names, roads).
    bbox is (min_lon, min_lat, max_lon, max_lat) in WGS-84.
    Returns PNG bytes for embedding in the Word document.
    """
    import io as _io
    import math
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import matplotlib.patheffects as pe
    import contextily as ctx

    min_lon, min_lat, max_lon, max_lat = bbox
    span_lon = max_lon - min_lon
    span_lat = max_lat - min_lat

    # Convert bbox corners to Web Mercator
    x_min, y_min = _lonlat_to_webmercator(min_lon, min_lat)
    x_max, y_max = _lonlat_to_webmercator(max_lon, max_lat)

    # Padding in Web Mercator units — enough context to show surrounding area
    pad_x = max((x_max - x_min) * 0.6, 30000)
    pad_y = max((y_max - y_min) * 0.6, 30000)

    fig, ax = plt.subplots(figsize=(5.5, 4.2))

    # Set axis limits (Web Mercator) before adding basemap
    ax.set_xlim(x_min - pad_x, x_max + pad_x)
    ax.set_ylim(y_min - pad_y, y_max + pad_y)

    # Add OpenStreetMap basemap tiles
    try:
        ctx.add_basemap(
            ax,
            crs="EPSG:3857",
            source=ctx.providers.CartoDB.Positron,
            zoom="auto",
            attribution=False,
        )
    except Exception:
        # Fallback if tiles can't be fetched (no network)
        ax.set_facecolor("#d6e4f0")

    # Highlighted region rectangle (semi-transparent fill + solid border)
    rect_fill = mpatches.Rectangle(
        (x_min, y_min), x_max - x_min, y_max - y_min,
        linewidth=0, facecolor="#e74c3c", alpha=0.20, zorder=3,
    )
    rect_border = mpatches.Rectangle(
        (x_min, y_min), x_max - x_min, y_max - y_min,
        linewidth=2.2, edgecolor="#c0392b", facecolor="none", zorder=4,
    )
    ax.add_patch(rect_fill)
    ax.add_patch(rect_border)

    # Centre marker
    cx_m = (x_min + x_max) / 2
    cy_m = (y_min + y_max) / 2
    ax.plot(cx_m, cy_m, "o", color="#c0392b", markersize=6,
            markeredgecolor="white", markeredgewidth=1.2, zorder=5)

    # Region name label just above the box
    if region_name:
        short = region_name[:40] + ("…" if len(region_name) > 40 else "")
        ax.text(
            cx_m, y_max + pad_y * 0.08,
            short,
            ha="center", va="bottom", fontsize=8,
            color="#1a1a1a", fontweight="bold",
            path_effects=[pe.withStroke(linewidth=2.5, foreground="white")],
            zorder=6,
        )

    ax.set_axis_off()
    ax.set_title("Analysis Region", fontsize=9.5, fontweight="bold",
                 color="#222222", pad=6)

    buf = _io.BytesIO()
    plt.tight_layout(pad=0.5)
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _em_combined_map_bytes(thumb_bytes, bbox, region_name="", gas_cfg=None):
    """Overlay the GEE concentration thumbnail on a contextily basemap.

    The GEE thumbnail is rendered by GEE for pad_bbox(bbox) — the expanded
    region used in all TROPOMI fetches. We must use that exact same extent
    for the contextily basemap so the two layers align correctly.

    Steps:
      1. Expand bbox with pad_bbox() — same expansion GEE used for the thumb.
      2. Make near-white (no-data/cloud) pixels transparent using Pillow.
      3. Draw a contextily basemap covering exactly the padded extent.
      4. Overlay the concentration image with imshow using the same extent.
      5. Return the combined PNG bytes.
    """
    import io as _io
    import numpy as _np
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patheffects as pe
    import matplotlib.patches as _mp
    from PIL import Image
    import contextily as ctx

    # Use the same padded bbox that GEE used when rendering the thumbnail
    padded = methane_explorer.pad_bbox(list(bbox))
    min_lon, min_lat, max_lon, max_lat = padded

    # Convert padded bbox corners to Web Mercator (EPSG:3857)
    x_min, y_min = _lonlat_to_webmercator(min_lon, min_lat)
    x_max, y_max = _lonlat_to_webmercator(max_lon, max_lat)

    fig, ax = plt.subplots(figsize=(5.5, 4.2))
    # Set extent to exactly the padded region — no extra padding
    ax.set_xlim(x_min, x_max)
    ax.set_ylim(y_min, y_max)

    # Basemap tiles — same extent as the GEE thumbnail
    try:
        ctx.add_basemap(
            ax, crs="EPSG:3857",
            source=ctx.providers.CartoDB.Positron,
            zoom="auto", attribution=False,
        )
    except Exception:
        ax.set_facecolor("#d6e4f0")

    # Make white / near-white pixels in the GEE thumbnail transparent.
    # GEE renders masked pixels (clouds, no data) as white (#ffffff).
    gee_img = Image.open(_io.BytesIO(thumb_bytes)).convert("RGBA")
    data = _np.array(gee_img, dtype=_np.uint8)
    r, g, b = data[:, :, 0], data[:, :, 1], data[:, :, 2]
    white_mask = (r > 235) & (g > 235) & (b > 235)
    data[white_mask, 3] = 0
    gee_rgba = Image.fromarray(data)

    # Overlay the concentration layer — extent matches exactly the basemap
    ax.imshow(
        gee_rgba,
        extent=[x_min, x_max, y_min, y_max],
        origin="upper",
        alpha=0.72,
        zorder=3,
        interpolation="bilinear",
    )

    # Region label centred at top
    if region_name:
        short = region_name[:40] + ("…" if len(region_name) > 40 else "")
        cx_m = (x_min + x_max) / 2
        ax.text(
            cx_m, y_max - (y_max - y_min) * 0.04,
            short, ha="center", va="top",
            fontsize=8.5, color="#1a1a1a", fontweight="bold",
            path_effects=[pe.withStroke(linewidth=2.5, foreground="white")],
            zorder=5,
        )

    ax.set_axis_off()
    ax.set_title("Concentration Map", fontsize=9.5, fontweight="bold",
                 color="#222222", pad=6)

    # Colorbar at the bottom using the GAS_CONFIG palette and value range
    if gas_cfg:
        import matplotlib.colors as mcolors
        import matplotlib.cm as mcm

        palette      = gas_cfg.get("palette", [])
        vmin         = gas_cfg.get("min_val", 0) * gas_cfg.get("display_scale", 1)
        vmax         = gas_cfg.get("max_val", 1) * gas_cfg.get("display_scale", 1)
        display_unit = gas_cfg.get("display_unit", gas_cfg.get("unit", ""))

        cmap = mcolors.LinearSegmentedColormap.from_list("gas_cmap", palette, N=256)
        norm = mcolors.Normalize(vmin=vmin, vmax=vmax)
        sm   = mcm.ScalarMappable(cmap=cmap, norm=norm)
        sm.set_array([])

        cbar = fig.colorbar(sm, ax=ax, orientation="horizontal",
                            fraction=0.025, pad=0.01, aspect=50, shrink=0.7)
        cbar.set_label(display_unit, fontsize=6.5, color="#333333", labelpad=2)
        cbar.ax.tick_params(labelsize=6, colors="#333333", length=2, pad=1)
        cbar.outline.set_edgecolor("#bbbbbb")

    buf = _io.BytesIO()
    plt.tight_layout(pad=0.5)
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _em_add_inline_bold(paragraph, text):
    """Add runs to a Word paragraph, honouring **bold** spans."""
    import re as _re
    parts = _re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def build_emissions_docx(ai_text, gas_name, region, actual_date,
                          mean_val, pass_count, disp_unit, conf_string,
                          model_name="", bbox=None, thumb_bytes=None, gas_cfg=None):
    """Build a Word document for the Emissions Explorer result.

    Sections:
      1. Title and metadata
      2. Statistics block
      3. Concentration map (GEE thumbnail if available, basemap fallback)
      4. AI Interpretation (markdown → Word formatting)

    Returns bytes ready for st.download_button.
    """
    import io as _io
    import re as _re
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Margins and font
    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ---- Title ----
    title_para = doc.add_paragraph()
    run = title_para.add_run("Emissions Intelligence Report")
    run.bold = True
    run.font.size = Pt(16)

    # ---- Metadata block ----
    for label, value in [
        ("Gas measured:", gas_name),
        ("Region:", region),
        ("Composite date:", actual_date),
        ("AI model:", model_name or "Built-in interpretation"),
    ]:
        p = doc.add_paragraph()
        p.add_run(label + " ").bold = True
        p.add_run(value)

    doc.add_paragraph()  # spacer

    # ---- Statistics section ----
    h = doc.add_paragraph("Statistics")
    h.style = doc.styles["Heading 1"]

    val_str = f"{mean_val:.2f} {disp_unit}" if mean_val is not None else "No valid pixels"
    stats_rows = [
        ("Regional mean concentration", val_str),
        ("Orbital passes (7-day window)", str(pass_count) if pass_count else "0"),
        ("Data confidence", conf_string),
        ("Sensor", "TROPOMI (Sentinel-5P)"),
        ("Spatial resolution", "~3.5 × 5.5 km"),
    ]
    tbl = doc.add_table(rows=len(stats_rows), cols=2)
    tbl.style = "Table Grid"
    for r_idx, (lbl, val) in enumerate(stats_rows):
        cells = tbl.rows[r_idx].cells
        cells[0].text = lbl
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
        cells[1].text = val

    doc.add_paragraph()  # spacer

    # ---- Concentration map ----
    # Use GEE thumbnail (same colours as the interactive map) when available.
    # Fall back to the contextily basemap if the thumbnail wasn't generated.
    h2 = doc.add_paragraph("Concentration Map")
    h2.style = doc.styles["Heading 1"]
    map_caption = doc.add_paragraph(
        f"{gas_name} concentration — {actual_date}. "
        "Colour scale: low (cool) → high (warm). Source: TROPOMI Sentinel-5P / GEE."
    )
    map_caption.style = doc.styles["Normal"]
    map_caption.runs[0].font.size = Pt(8)
    map_caption.runs[0].font.italic = True
    try:
        if thumb_bytes and bbox and len(bbox) == 4:
            # Concentration layer overlaid on real basemap tiles
            png = _em_combined_map_bytes(thumb_bytes, bbox, region_name=region, gas_cfg=gas_cfg)
        elif bbox and len(bbox) == 4:
            # Basemap only (fallback if GEE thumb unavailable)
            png = _em_static_map_bytes(bbox, region_name=region)
        else:
            png = None
        if png:
            doc.add_picture(_io.BytesIO(png), width=Inches(4.8))
            doc.add_paragraph()
    except Exception:
        doc.add_paragraph("[Map could not be generated]")

    # ---- AI Interpretation ----
    h3 = doc.add_paragraph("AI Interpretation")
    h3.style = doc.styles["Heading 1"]

    if ai_text:
        lines = ai_text.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if stripped.startswith("## ") or stripped.startswith("# "):
                p = doc.add_paragraph(stripped.lstrip("# ").strip())
                p.style = doc.styles["Heading 1"]
                i += 1

            elif stripped.startswith("|"):
                block = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    block.append(lines[i])
                    i += 1
                data_lines = [l for l in block if not _re.match(r"^\s*\|[-| :]+\|\s*$", l)]
                if len(data_lines) >= 2:
                    def _split(ln):
                        return [c.strip() for c in ln.strip().strip("|").split("|")]
                    headers = _split(data_lines[0])
                    rows_data = [_split(l) for l in data_lines[1:]]
                    t = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
                    t.style = "Table Grid"
                    hdr_cells = t.rows[0].cells
                    for ci, hdr in enumerate(headers):
                        hdr_cells[ci].text = hdr
                        for run in hdr_cells[ci].paragraphs[0].runs:
                            run.bold = True
                    for ri, rd in enumerate(rows_data):
                        rc = t.rows[ri + 1].cells
                        for ci, ct in enumerate(rd):
                            if ci < len(rc):
                                rc[ci].text = ct
                    doc.add_paragraph()
                else:
                    for bl in block:
                        doc.add_paragraph(bl.strip())

            elif _re.match(r"^\d+\.\s", stripped):
                p = doc.add_paragraph(style="List Number")
                _em_add_inline_bold(p, stripped[stripped.index(". ") + 2:])
                i += 1

            elif stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                _em_add_inline_bold(p, stripped[2:])
                i += 1

            elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
                p = doc.add_paragraph()
                p.add_run(stripped.strip("*")).bold = True
                i += 1

            elif stripped == "" or stripped.startswith("---"):
                doc.add_paragraph()
                i += 1

            else:
                if stripped:
                    p = doc.add_paragraph()
                    _em_add_inline_bold(p, stripped)
                i += 1
    else:
        doc.add_paragraph("No AI interpretation available.")

    # ---- Footer ----
    footer = doc.add_paragraph()
    footer.add_run("Generated by EOIL — AI-Native Earth Observation Innovation Lab").italic = True

    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# MODULE 6 — Emissions Explorer (TROPOMI Sentinel-5P atmospheric gases)
# ---------------------------------------------------------------------------

if selected_module == "🌫️ Emissions Explorer":

    from datetime import date as _date

    gee_available = st.session_state.gee_available

    st.subheader("🌫️ Emissions Explorer")
    st.caption(
        "Select a location, gas, and date. "
        "The module fetches TROPOMI satellite data and maps atmospheric concentration "
        "with an AI interpretation for utility and industrial operators."
    )

    with st.expander("ℹ️ What does this module do?", expanded=False):
        st.markdown("""
**This module answers the question: what is the atmospheric concentration of industrial gases over this region?**

It fetches data from the TROPOMI instrument aboard the Copernicus Sentinel-5P satellite.
TROPOMI measures trace gas concentrations across the full globe every day.

**Four gases are available:**

| Gas | What it reveals |
|---|---|
| Methane (CH4) | Pipeline leaks, LNG facilities, coal mines, landfill emissions |
| Nitrogen Dioxide (NO2) | Power plant combustion, vehicle fleets, industrial processes |
| Carbon Monoxide (CO) | Wildfire plumes, industrial combustion, emergency response |
| Sulfur Dioxide (SO2) | Coal plant stacks, industrial sites, volcanic degassing |

**What you will see after running an analysis:**

- **Regional average concentration** — the area-mean value for your selected date
- **Interactive concentration map** — TROPOMI tile layer showing spatial variation across the region
- **AI interpretation** — four-section analysis: pattern, source attribution, regulatory context, and utility action

**How to use it:**

1. Type any region or industrial area in the location box
2. Select the gas you want to analyse
3. Pick a target date (the module searches a 5-day window if that date has no data)
4. Click Run Analysis

**Data source:** Copernicus Sentinel-5P TROPOMI via Google Earth Engine (ESA / EU).
Data is available from 2018 onwards. The effective spatial resolution is approximately
5.5 km × 7 km per TROPOMI pixel, with global daily coverage.
        """)

    # --- Controls ---
    col_loc, col_gas = st.columns([3, 1])

    with col_loc:
        em_place = st.text_input(
            "Location — type any industrial region, basin, or city",
            placeholder="e.g. Permian Basin, Texas   |   Ruhr Valley, Germany   |   Shanxi Province, China",
            key="em_place",
        )

    with col_gas:
        em_gas = st.selectbox(
            "Gas",
            list(methane_explorer.GAS_CONFIG.keys()),
            key="em_gas",
        )

    col_date, _, col_run = st.columns([2, 2, 1])

    with col_date:
        em_date = st.date_input(
            "Target date",
            value=_date.today() - timedelta(days=30),
            min_value=_date(2018, 5, 1),
            max_value=_date.today() - timedelta(days=3),
            key="em_date",
        )

    with col_run:
        em_run_btn = st.button(
            "▶ Run Analysis", type="primary",
            use_container_width=True, key="em_run",
        )

    # --- Geocode ---
    em_bbox        = None
    em_region_name = ""

    if em_place.strip():
        cached_place = st.session_state.get("em_geocoded_place", "")
        cached_bbox  = st.session_state.get("em_geocoded_bbox",  None)

        if em_place.strip() != cached_place:
            map_picker.clear_click("em")
            with st.spinner(f"Looking up '{em_place}'..."):
                result_bbox = geocoder.geocode_place(em_place)
            if result_bbox:
                st.session_state.em_geocoded_place = em_place.strip()
                st.session_state.em_geocoded_bbox  = result_bbox
                cached_bbox = result_bbox
            else:
                st.session_state.em_geocoded_place = ""
                st.session_state.em_geocoded_bbox  = None
                cached_bbox = None
                st.error(f"Could not find '{em_place}'. Try a broader name.")

        if cached_bbox:
            em_bbox        = cached_bbox
            em_region_name = em_place.strip()
            st.caption(f"📍 {em_region_name}")

    # Map picker — optional refinement
    if em_bbox:
        with st.expander("📍 Refine location — click map to set exact area", expanded=False):
            picked = map_picker.render_map_picker(
                centre_bbox     = em_bbox,
                picker_key      = "em",
                default_size_km = 200,
            )
            if picked:
                em_bbox = picked

    # GEE status
    if gee_available:
        st.caption("🟢 GEE connected — live TROPOMI data active.")
    else:
        st.caption("🔴 GEE not connected. Emissions Explorer requires live GEE data.")

    st.divider()

    # --- Session state ---
    for _k, _v in [
        ("em_map",            None), ("em_mean_val",       None),
        ("em_actual_date",    None), ("em_pass_count",     None),
        ("em_result_region",  None), ("em_result_gas",     None),
        ("em_result_bbox",    None), ("em_result_thumb",   None),
        ("em_ai_result",      None), ("em_ai_model",       None),
    ]:
        if _k not in st.session_state:
            st.session_state[_k] = _v

    # --- Two-step run pattern ---
    if em_run_btn:
        if not em_bbox:
            st.warning("Enter a location first.")
        elif not gee_available:
            st.error("Emissions Explorer requires GEE credentials. Add GEE_SERVICE_ACCOUNT_JSON to Streamlit secrets.")
        else:
            st.session_state.em_map         = None
            st.session_state.em_mean_val    = None
            st.session_state.em_ai_result   = None
            st.session_state.em_pending_run = {
                "bbox":   em_bbox,
                "gas":    em_gas,
                "date":   str(em_date),
                "region": em_region_name,
            }
            st.rerun()

    if st.session_state.get("em_pending_run"):
        p = st.session_state.em_pending_run
        st.session_state.em_pending_run = None

        cfg      = methane_explorer.GAS_CONFIG[p["gas"]]
        band     = cfg["band"]

        with st.spinner(f"Fetching TROPOMI {p['gas']} data for {p['region']}..."):
            image, actual_date, pass_count = methane_explorer.fetch_tropomi_mosaic(
                p["gas"], p["bbox"], p["date"]
            )

        if image is None:
            st.error(
                f"No TROPOMI data found for {p['gas']} near {p['date']}. "
                "Try a different date or a wider search window."
            )
        else:
            with st.spinner("Computing regional statistics..."):
                mean_val = methane_explorer.get_regional_mean(image, p["bbox"], band)

            with st.spinner("Building emissions map (20-40 seconds)..."):
                em_map = methane_explorer.build_emissions_map(
                    image, p["gas"], p["bbox"], actual_date
                )

            st.session_state.em_map          = em_map
            st.session_state.em_mean_val     = mean_val
            st.session_state.em_actual_date  = actual_date
            st.session_state.em_pass_count   = pass_count
            st.session_state.em_result_region = p["region"]
            st.session_state.em_result_gas    = p["gas"]
            st.session_state.em_result_bbox   = p["bbox"]
            # Fetch static concentration thumbnail for Word export
            st.session_state.em_result_thumb = methane_explorer.get_concentration_thumb(
                image, p["gas"], p["bbox"]
            )
            st.success(
                f"Data loaded — {p['gas']} over {p['region']} "
                f"(7-day window: {actual_date}, {pass_count} orbital passes)."
            )

    # --- Display results ---
    if st.session_state.em_map is not None:
        cfg        = methane_explorer.GAS_CONFIG[st.session_state.em_result_gas]
        mean_val   = st.session_state.em_mean_val
        actual_date= st.session_state.em_actual_date
        pass_count = st.session_state.em_pass_count
        r_reg      = st.session_state.em_result_region
        r_gas      = st.session_state.em_result_gas

        def em_section_break():
            st.markdown(
                '<hr style="border: none; border-top: 3px solid #d0d0d0; margin: 28px 0 20px 0;">',
                unsafe_allow_html=True,
            )

        # --- SECTION 1: Summary metric ---
        st.subheader("📊 Regional Concentration")
        c1, c2, c3 = st.columns(3)

        disp_scale = cfg.get("display_scale", 1)
        disp_unit  = cfg.get("display_unit",  cfg["unit"])
        if mean_val is not None:
            scaled_val  = mean_val * disp_scale
            # Pick decimal places based on scaled magnitude
            magnitude = abs(scaled_val)
            if magnitude == 0:
                val_display = "0.000"
            elif magnitude >= 100:
                val_display = f"{scaled_val:.1f}"
            elif magnitude >= 1:
                val_display = f"{scaled_val:.3f}"
            else:
                val_display = f"{scaled_val:.4f}"
        else:
            val_display = "No valid pixels"
        c1.metric(f"Regional average ({disp_unit})", val_display)
        c2.metric("Period (7-day composite)", actual_date)
        c3.metric("Orbital passes in window", pass_count)

        st.caption(cfg["description"])

        em_section_break()

        # --- SECTION 2: Emissions map ---
        st.subheader("🗺️ Concentration Map")
        st.caption(
            "Colour scale runs from low (cool) to high (warm). "
            "Use the layer control in the top-right corner to toggle the gas layer. "
            "Zoom in to see spatial variation within the region."
        )

        if st.session_state.em_map is not None:
            from streamlit_folium import st_folium as _st_folium
            _st_folium(st.session_state.em_map, width=700, height=500, returned_objects=[])
        else:
            st.warning("Map could not be built. Check GEE connection.")

        em_section_break()

        # --- SECTION 3: AI Interpretation ---
        st.subheader("🤖 AI Interpretation")
        if st.button("Get AI Interpretation", type="primary", key="em_ai_btn"):
            with st.spinner("Thinking..."):
                interpretation, model_used = methane_explorer.get_emissions_interpretation(
                    r_gas, mean_val, r_reg, actual_date,
                    groq_key=config.GROQ_API_KEY,
                    gemini_key=config.GEMINI_API_KEY,
                )
                st.session_state.em_ai_result = interpretation
                st.session_state.em_ai_model  = model_used

        if st.session_state.get("em_ai_result"):
            with st.expander("📋 AI Interpretation", expanded=True):
                st.markdown(st.session_state.em_ai_result)
                model_used = st.session_state.get("em_ai_model")
                if model_used:
                    st.caption(f"AI response from **{model_used}**")
                else:
                    st.caption("Showing built-in interpretation. Add GROQ_API_KEY or GEMINI_API_KEY to enable AI.")

            # Download buttons
            _em_ai_text   = st.session_state.em_ai_result
            _em_model     = st.session_state.get("em_ai_model", "")
            _em_bbox      = st.session_state.get("em_result_bbox")
            _em_pc        = pass_count if pass_count else 0
            if _em_pc >= 7:
                _em_conf_str = "High"
            elif _em_pc >= 4:
                _em_conf_str = "Moderate"
            else:
                _em_conf_str = "Limited"
            _em_scaled_val = (mean_val * disp_scale) if mean_val is not None else None

            _em_filename = (
                r_gas.split("(")[-1].rstrip(")").strip().lower().replace(" ", "_")
                + "_" + r_reg.lower().replace(" ", "_").replace(",", "")[:30]
            )

            _dc1, _dc2 = st.columns([1, 1])
            with _dc1:
                st.download_button(
                    label="⬇ Download as Markdown",
                    data=_em_ai_text,
                    file_name=f"emissions_{_em_filename}.md",
                    mime="text/markdown",
                    key="em_dl_md",
                )
            with _dc2:
                _docx_bytes = build_emissions_docx(
                    ai_text=_em_ai_text,
                    gas_name=r_gas,
                    region=r_reg,
                    actual_date=actual_date,
                    mean_val=_em_scaled_val,
                    pass_count=_em_pc,
                    disp_unit=disp_unit,
                    conf_string=_em_conf_str,
                    model_name=_em_model,
                    bbox=_em_bbox,
                    thumb_bytes=st.session_state.get("em_result_thumb"),
                    gas_cfg=cfg,
                )
                st.download_button(
                    label="⬇ Download as Word (.docx)",
                    data=_docx_bytes,
                    file_name=f"emissions_{_em_filename}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="em_dl_docx",
                )

        em_section_break()

        # --- Data Quality ---
        st.subheader("🔍 Data Quality")
        _pc = pass_count if pass_count else 0
        if _pc >= 7:
            _em_conf = "High — 7 or more orbital passes. Good spatial coverage for the 7-day window."
        elif _pc >= 4:
            _em_conf = "Moderate — 4–6 passes. Some areas may have sparse coverage."
        else:
            _em_conf = ("Limited — fewer than 4 passes. Coverage may be patchy. "
                        "Interpret regional average with caution.")
        st.info(
            f"**Sensor:** TROPOMI (Sentinel-5P)  \n"
            f"**Gas measured:** {cfg.get('label', r_gas)}  \n"
            f"**Spatial resolution:** ~3.5 × 5.5 km (TROPOMI nadir pixel)  \n"
            f"**Composite window:** 7-day mosaic centred on {actual_date}  \n"
            f"**Orbital passes:** {_pc} passes within the window  \n"
            f"**Cloud handling:** TROPOMI cloud fraction filter applied. "
            f"Pixels with cloud fraction > 0.5 excluded.  \n"
            f"**Confidence:** {_em_conf}"
        )

    else:
        st.markdown("---")
        st.markdown(
            "**Type a location above, select a gas, pick a date, then click Run Analysis.**\n\n"
            "The module will fetch live TROPOMI data and show an interactive concentration map "
            "alongside a regional average and an AI interpretation."
        )

    st.stop()

# ---------------------------------------------------------------------------
# MODULE 7 — Land Cover Intelligence (Arc 1)
# K-means clustering + Random Forest on Sentinel-2 via Planetary Computer.
# No GEE required — uses PC STAC API directly.
# ---------------------------------------------------------------------------

if selected_module == "🌿 Land Cover Intelligence":

    # Initialise session state keys used by land_cover.py
    for _k, _v in [
        ("lc_scene",             None),
        ("lc_kmeans",            None),
        ("lc_rf",                None),
        ("lc_ai",                None),
        ("lc_ai_model",          None),
        ("lc_region",            ""),
        ("lc_pending",           None),
        ("lc_geocoded_place",    ""),
        ("lc_geocoded_bbox",     None),
        ("lc_available_scenes",  None),
        ("lc_search_bbox",       None),
        ("lc_search_region",     ""),
    ]:
        if _k not in st.session_state:
            st.session_state[_k] = _v

    land_cover.render(
        location_name=st.session_state.get("lc_geocoded_place", ""),
        bbox=st.session_state.get("lc_geocoded_bbox", None),
    )

    st.stop()

# ---------------------------------------------------------------------------
# MODULE 8 — Corridor Risk Intelligence (Arc 2)
# Multi-algorithm vegetation encroachment analysis via GEE NDVI time series.
# ---------------------------------------------------------------------------

if selected_module == "⚠️ Corridor Risk Intelligence":

    corridor_risk.render()
    st.stop()

# ---------------------------------------------------------------------------
# MODULE 0 — Welcome panel (default when Welcome is selected)
# ---------------------------------------------------------------------------

st.markdown("## Welcome to the EOIL Portal")
st.markdown(
    "This portal is a working demonstration of AI-native Earth observation analysis. "
    "It pulls real satellite data from ESA, NASA, and Google Earth Engine."
)

st.divider()

st.markdown("### What you can do here")

col1, col2 = st.columns(2)

with col1:
    st.markdown("#### 🔬 Spectral Explorer")
    st.markdown(
        "Fetch real satellite imagery from NASA and ESA archives for any location. "
        "Render the same scene through different spectral band combinations — "
        "vegetation health, water extent, urban heat, burn scars, and more. "
        "Compare all views side by side."
    )
    st.markdown("**Sensor:** Sentinel-2, Landsat 8/9")
    st.markdown("**Data:** Planetary Computer (Microsoft)")
    st.markdown("**Question:** What can satellites see that cameras cannot?")

with col2:
    st.markdown("#### 📈 Time Series Explorer")
    st.markdown(
        "Pull 10+ years of satellite measurements for any location on Earth. "
        "See how vegetation greenness or land surface temperature has shifted "
        "over time. Trend analysis, seasonal cycles, and annual comparison charts."
    )
    st.markdown("**Sensor:** MODIS Terra, Landsat 8")
    st.markdown("**Data:** Google Earth Engine")
    st.markdown("**Question:** How has this location changed over time?")

st.divider()

col3, col4 = st.columns(2)

with col3:
    st.markdown("#### 📡 SAR Explorer")
    st.markdown(
        "Analyse Sentinel-1 radar data for any location and two dates. "
        "SAR works through clouds and at night — it reveals ships, infrastructure, "
        "and surface change that optical sensors miss. VV, VH, false color, "
        "and change maps included."
    )
    st.markdown("**Sensor:** Sentinel-1 GRD (ESA / Copernicus)")
    st.markdown("**Data:** Google Earth Engine")
    st.markdown("**Question:** What can radar see that cameras cannot?")

with col4:
    st.markdown("#### 🔀 Change Detection")
    st.markdown(
        "Select any location and two dates. The module computes the NDVI difference "
        "between the two dates and shows where vegetation increased or decreased. "
        "Interactive change map with toggleable layers and AI interpretation."
    )
    st.markdown("**Sensor:** Sentinel-2 SR, MODIS (fallback)")
    st.markdown("**Data:** Google Earth Engine")
    st.markdown("**Question:** What changed between two dates?")

st.divider()

col5, col6 = st.columns(2)

with col5:
    st.markdown("#### 🔍 AI Imagery Interpreter")
    st.markdown(
        "Pick any location and date. The module fetches a Sentinel-2 true-color chip "
        "and passes it directly to a vision AI. The AI reads the image and describes "
        "land cover, notable features, seasonal state, and decision applications."
    )
    st.markdown("**Sensor:** Sentinel-2 L2A")
    st.markdown("**Data:** Planetary Computer (Microsoft)")
    st.markdown("**Question:** What does this place look like — and what does it mean?")

with col6:
    st.markdown("#### 🌫️ Emissions Explorer")
    st.markdown(
        "Map atmospheric gas concentrations for any region and date. "
        "Visualise methane, nitrogen dioxide, carbon monoxide, and sulphur dioxide "
        "from TROPOMI satellite measurements. Compare two dates to see concentration change."
    )
    st.markdown("**Sensor:** Sentinel-5P TROPOMI (ESA / Copernicus)")
    st.markdown("**Data:** Google Earth Engine")
    st.markdown("**Question:** Where are elevated emissions occurring — and did they change?")

st.divider()

col7, col8 = st.columns(2)

with col7:
    st.markdown("#### 🌿 Land Cover Intelligence")
    st.markdown(
        "Classify land cover for any location using Sentinel-2 multispectral imagery. "
        "Two algorithms run on the same scene: K-means unsupervised clustering and "
        "Random Forest supervised classification. Compare outputs and see where they agree."
    )
    st.markdown("**Sensor:** Sentinel-2 L2A")
    st.markdown("**Data:** Planetary Computer (Microsoft)")
    st.markdown("**Question:** What land cover types are present — and how confident is the classification?")

st.divider()

st.markdown("### How to get started")
st.markdown(
    "Use the **sidebar on the left** to navigate between modules. "
    "Each module has an info panel that explains what it does and how to use it. "
    "All modules pull live data — nothing here is simulated."
)

st.divider()

st.markdown("### About this portal")
col_a, col_b = st.columns(2)

with col_a:
    st.markdown(
        "The EOIL Portal is built as part of the AI-Native Earth Observation "
        "Innovation Lab — a 30-day program to rebuild satellite data analysis skills "
        "and integrate modern AI into every workflow."
    )

with col_b:
    st.markdown("**Technology stack:**")
    st.markdown("- Streamlit — app framework")
    st.markdown("- Google Earth Engine — satellite data processing at scale")
    st.markdown("- Planetary Computer — optical satellite archive")
    st.markdown("- Groq + Gemini — 11-model text chain, 8-model vision chain")
    st.markdown("- Folium / Plotly — interactive maps and charts")

st.divider()
st.caption(
    "EOIL Portal v2.0 — Earth Observation Innovation Lab. Modules: Spectral, Time Series, SAR, Change Detection, AI Imagery Interpreter, Emissions Explorer, Land Cover Intelligence, Corridor Risk Intelligence. "
    "Built with Claude Code. "
    "Login and access controls will be added in a future version."
)
