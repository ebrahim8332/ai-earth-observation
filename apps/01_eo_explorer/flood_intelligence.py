"""
flood_intelligence.py — Flood Detection and Impact Mapping module for the EOIL portal.

Arc 3: Flood Detection and Impact Mapping.

Three-layer architecture:
  Layer 1: Sentinel-1 SAR backscatter change, NDWI from Sentinel-2, SRTM slope mask,
            JRC Global Surface Water permanent water mask
  Layer 2: Combined flood extent classification — flooded km², permanent water km²,
            new flood-only km², confidence level
  Layer 3: Gemini/Groq structured impact brief — five-element flood intelligence output

All algorithm logic is ported from notebook 10_flood_detection.ipynb.
Results are cached in Streamlit session state so switching tabs does not re-run.
"""

import io
import re
import warnings
warnings.filterwarnings('ignore')

import numpy as np
import requests
import streamlit as st
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from PIL import Image
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor

import config
import ai_chain

# ---------------------------------------------------------------------------
# Pre-defined flood events
# ---------------------------------------------------------------------------

FLOOD_EVENTS = {
    "Pakistan 2022 — Indus Floodplain, Sindh": {
        "bbox":         [67.0, 25.5, 69.5, 27.5],
        "center":       [26.5, 68.2],
        "zoom":         7,
        "before_start": "2022-06-01",
        "before_end":   "2022-07-31",
        "after_start":  "2022-08-20",
        "after_end":    "2022-09-30",
        "context": (
            "One-third of Pakistan submerged. 33 million displaced, 1,739 deaths, "
            "2 million homes damaged. Lower Indus floodplain, Sindh province. "
            "Heavy monsoon cloud cover throughout the event."
        ),
        "sar_threshold": -3.0,
        "orbit": "DESCENDING",
    },
    "Bangladesh 2020 — Brahmaputra, Sylhet Division": {
        "bbox":         [90.5, 24.0, 92.5, 25.5],
        "center":       [24.8, 91.5],
        "zoom":         8,
        "before_start": "2020-06-01",
        "before_end":   "2020-06-30",
        "after_start":  "2020-07-01",
        "after_end":    "2020-08-15",
        "context": (
            "Severe monsoon flooding along the Brahmaputra and Meghna river systems. "
            "Sylhet, Mymensingh and Jamalpur divisions heavily affected. "
            "Cloud cover during monsoon season makes SAR the primary detection sensor."
        ),
        "sar_threshold": -3.0,
        "orbit": "DESCENDING",
    },
    "Nigeria 2022 — Niger-Benue Confluence, Lokoja": {
        "bbox":         [6.2, 7.5, 7.8, 9.0],
        "center":       [8.2, 6.9],
        "zoom":         8,
        "before_start": "2022-08-01",
        "before_end":   "2022-09-15",
        "after_start":  "2022-09-16",
        "after_end":    "2022-11-15",
        "context": (
            "Flooding at the confluence of the Niger and Benue rivers. "
            "Over 600 deaths, 1.3 million displaced. Anambra, Kogi, and Delta states affected. "
            "West African rainy season limits optical sensor utility."
        ),
        "sar_threshold": -3.0,
        "orbit": "DESCENDING",
    },
}

# ---------------------------------------------------------------------------
# GEE analysis
# ---------------------------------------------------------------------------

def _run_flood_analysis(event_key: str, sar_threshold: float) -> dict:
    """Run the full four-signal flood detection algorithm for a flood event."""
    import ee

    event = FLOOD_EVENTS[event_key]
    bbox  = event["bbox"]
    aoi   = ee.Geometry.Rectangle(bbox)

    # Layer 1a: Sentinel-1 SAR
    def get_sar(start, end):
        return (
            ee.ImageCollection('COPERNICUS/S1_GRD')
            .filterBounds(aoi)
            .filterDate(start, end)
            .filter(ee.Filter.eq('instrumentMode', 'IW'))
            .filter(ee.Filter.listContains('transmitterReceiverPolarisation', 'VV'))
            .filter(ee.Filter.eq('orbitProperties_pass', event["orbit"]))
            .select('VV')
            .median()
            .clip(aoi)
        )

    sar_before = get_sar(event["before_start"], event["before_end"])
    sar_after  = get_sar(event["after_start"],  event["after_end"])
    sar_change = sar_after.subtract(sar_before).rename('sar_change')

    def count_sar(start, end):
        return (
            ee.ImageCollection('COPERNICUS/S1_GRD')
            .filterBounds(aoi)
            .filterDate(start, end)
            .filter(ee.Filter.eq('instrumentMode', 'IW'))
            .filter(ee.Filter.listContains('transmitterReceiverPolarisation', 'VV'))
            .filter(ee.Filter.eq('orbitProperties_pass', event["orbit"]))
            .size().getInfo()
        )

    before_count = count_sar(event["before_start"], event["before_end"])
    after_count  = count_sar(event["after_start"],  event["after_end"])

    # Layer 1b: Sentinel-2 NDWI
    def compute_ndwi(image):
        return image.normalizedDifference(['B3', 'B8']).rename('NDWI')

    s2_col   = (
        ee.ImageCollection('COPERNICUS/S2_SR_HARMONIZED')
        .filterBounds(aoi)
        .filterDate(event["after_start"], event["after_end"])
        .filter(ee.Filter.lt('CLOUDY_PIXEL_PERCENTAGE', 20))
    )
    s2_count = s2_col.size().getInfo()

    if s2_count > 0:
        ndwi_image = s2_col.map(compute_ndwi).median().clip(aoi)
        ndwi_water = ndwi_image.gt(0.2).rename('ndwi_water')
    else:
        ndwi_image = ee.Image(0).rename('NDWI').clip(aoi)
        ndwi_water = ee.Image(0).rename('ndwi_water').clip(aoi)

    # Layer 1c: SRTM slope mask
    dem        = ee.Image('USGS/SRTMGL1_003').clip(aoi)
    slope      = ee.Terrain.slope(dem)
    slope_mask = slope.lt(5).rename('slope_mask')

    # Layer 1d: JRC permanent water
    gsw             = ee.Image('JRC/GSW1_4/GlobalSurfaceWater').select('occurrence').clip(aoi)
    permanent_water = gsw.gt(75).rename('permanent_water')

    # Layer 2: Combined classification
    sar_candidate   = sar_change.lt(sar_threshold)
    sar_flat        = sar_candidate.And(slope_mask)
    sar_new         = sar_flat.And(permanent_water.Not())
    confirmed_flood = sar_new.And(ndwi_water).multiply(2)
    sar_only_flood  = sar_new.And(ndwi_water.Not()).multiply(1)
    flood_class     = confirmed_flood.add(sar_only_flood).where(permanent_water, -1)
    flood_class     = flood_class.rename('flood_class').clip(aoi)

    # Area statistics
    pixel_area = ee.Image.pixelArea().divide(1e6)

    def class_area(val):
        return (
            flood_class.eq(val).multiply(pixel_area)
            .reduceRegion(reducer=ee.Reducer.sum(), geometry=aoi, scale=500, maxPixels=1e9)
            .getInfo().get('flood_class', 0)
        )

    area_permanent   = class_area(-1)
    area_sar_flood   = class_area(1)
    area_confirmed   = class_area(2)
    area_total_flood = area_sar_flood + area_confirmed

    # Confidence
    if s2_count >= 3 and area_confirmed > 0:
        confidence      = 'HIGH'
        confidence_note = 'SAR change and NDWI in agreement across multiple optical scenes.'
    elif s2_count > 0:
        confidence      = 'MEDIUM'
        confidence_note = 'SAR primary detector. Limited optical confirmation available.'
    else:
        confidence      = 'MEDIUM'
        confidence_note = (
            'SAR only — no cloud-free optical scenes found during the flood period '
            '(cloud cover filter: <20%). Consistent with storm/monsoon conditions '
            'that typically accompany major flood events.'
        )

    # Thumbnail URLs
    region_coords = aoi.getInfo()['coordinates']
    thumb_params  = {'region': region_coords, 'dimensions': '300x300', 'format': 'png'}

    sar_vis    = {'min': -25, 'max': 0,  'palette': ['000000', 'ffffff']}
    change_vis = {'min': -10, 'max': 5,  'palette': ['0000ff', 'ffffff', 'ff0000']}
    flood_vis  = {'min': -1,  'max': 2,  'palette': ['1E90FF', 'f0f0f0', 'FF8C00', 'FF0000']}

    try:
        url_before = sar_before.getThumbURL({**sar_vis,    **thumb_params})
        url_after  = sar_after.getThumbURL( {**sar_vis,    **thumb_params})
        url_change = sar_change.getThumbURL({**change_vis, **thumb_params})
        url_flood  = flood_class.getThumbURL({**flood_vis, **thumb_params})
    except Exception:
        url_before = url_after = url_change = url_flood = None

    # Cloud cover summary for Layer 3 prompt
    avg_cloud = 'unknown'
    s2_any = (
        ee.ImageCollection('COPERNICUS/S2_SR_HARMONIZED')
        .filterBounds(aoi)
        .filterDate(event["after_start"], event["after_end"])
        .size().getInfo()
    )
    if s2_any > 0:
        try:
            cloud_stats = (
                ee.ImageCollection('COPERNICUS/S2_SR_HARMONIZED')
                .filterBounds(aoi)
                .filterDate(event["after_start"], event["after_end"])
                .aggregate_stats('CLOUDY_PIXEL_PERCENTAGE')
                .getInfo()
            )
            mean_cloud = cloud_stats.get('mean')
            if mean_cloud is not None:
                avg_cloud = f'{mean_cloud:.0f}%'
        except Exception:
            pass

    return {
        "event_key":        event_key,
        "event":            event,
        "before_count":     before_count,
        "after_count":      after_count,
        "s2_count":         s2_count,
        "avg_cloud":        avg_cloud,
        "area_permanent":   area_permanent,
        "area_sar_flood":   area_sar_flood,
        "area_confirmed":   area_confirmed,
        "area_total_flood": area_total_flood,
        "confidence":       confidence,
        "confidence_note":  confidence_note,
        "url_before":       url_before,
        "url_after":        url_after,
        "url_change":       url_change,
        "url_flood":        url_flood,
        "sar_threshold":    sar_threshold,
    }


def _fetch_png(url: str):
    """Download a PNG from a GEE thumbnail URL. Returns bytes or None."""
    if not url:
        return None
    try:
        resp = requests.get(url, timeout=30)
        if resp.status_code == 200:
            return resp.content
    except Exception:
        pass
    return None


# ---------------------------------------------------------------------------
# Charts
# ---------------------------------------------------------------------------

def _fig_to_bytes(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _build_sar_row_fig(results: dict) -> bytes:
    """Three SAR panels in one row: before, after, change."""
    panels = [
        (results['url_before'],
         f"SAR Before\n{results['event']['before_start'][:7]} – {results['event']['before_end'][:7]}",
         f"VV median · {results['before_count']} scenes\nDark = smooth surface / water"),
        (results['url_after'],
         f"SAR After\n{results['event']['after_start'][:7]} – {results['event']['after_end'][:7]}",
         f"VV median · {results['after_count']} scenes\nAreas now darker = surface smoothed by flood"),
        (results['url_change'],
         "SAR Backscatter Change",
         "After minus Before (dB)\nBlue = increase · Red = decrease (flood signal)"),
    ]

    fig, axes = plt.subplots(1, 3, figsize=(14, 5))
    fig.suptitle(
        f"Sentinel-1 SAR — {results['event_key']}",
        fontsize=11, fontweight='bold', y=1.01
    )

    for ax, (url, title, subtitle) in zip(axes, panels):
        raw = _fetch_png(url)
        if raw:
            try:
                img = Image.open(BytesIO(raw))
                ax.imshow(img)
            except Exception:
                ax.set_facecolor('#cccccc')
                ax.text(0.5, 0.5, 'Unavailable', ha='center', va='center',
                        transform=ax.transAxes, fontsize=9)
        else:
            ax.set_facecolor('#cccccc')
            ax.text(0.5, 0.5, 'Unavailable', ha='center', va='center',
                    transform=ax.transAxes, fontsize=9)
        ax.set_title(title, fontsize=10, fontweight='bold', pad=4)
        ax.set_xlabel(subtitle, fontsize=8, labelpad=4)
        ax.set_xticks([])
        ax.set_yticks([])

    plt.tight_layout()
    return _fig_to_bytes(fig)


def _build_extent_chart(results: dict) -> bytes:
    """Bar chart: permanent water, SAR-only flood, confirmed flood."""
    r          = results
    categories = ['Permanent\nWater', 'SAR-Only\nFlood', 'Confirmed\nFlood\n(SAR+NDWI)']
    values     = [r['area_permanent'], r['area_sar_flood'], r['area_confirmed']]
    colors     = ['#1E90FF', '#FF8C00', '#CC2200']

    fig, ax = plt.subplots(figsize=(7, 4))
    bars    = ax.bar(categories, values, color=colors, edgecolor='white', linewidth=1.5, width=0.5)

    max_val = max(values) if max(values) > 0 else 1
    for bar, val in zip(bars, values):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + max_val * 0.015,
            f'{val:,.0f} km²',
            ha='center', va='bottom', fontweight='bold', fontsize=10
        )

    ax.set_ylabel('Area (km²)', fontsize=10)
    ax.set_title(
        f"Flood Extent Classification\n"
        f"Total new flood: {r['area_total_flood']:,.0f} km²  |  Confidence: {r['confidence']}",
        fontsize=11, fontweight='bold'
    )
    ax.set_facecolor('#f8f8f8')
    fig.patch.set_facecolor('white')
    ax.grid(axis='y', alpha=0.3)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    return _fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# Layer 3: prompt and fallback
# ---------------------------------------------------------------------------

def _build_impact_prompt(results: dict) -> str:
    e = results["event"]
    r = results
    return f"""You are an Earth observation analyst producing a structured flood impact brief for an emergency coordinator.

EVENT: {r['event_key']}
CONTEXT: {e['context']}

LAYER 2 OUTPUTS (computed from Sentinel-1 SAR and supporting data):
- Permanent water baseline: {r['area_permanent']:.0f} km²
- SAR-detected new flood area: {r['area_sar_flood']:.0f} km²
- Optically confirmed flood area: {r['area_confirmed']:.0f} km²
- Total new flood area: {r['area_total_flood']:.0f} km²
- Confidence level: {r['confidence']}
- Confidence note: {r['confidence_note']}
- SAR scenes used (before period): {r['before_count']}
- SAR scenes used (after period): {r['after_count']}
- Cloud-free optical scenes available: {r['s2_count']}
- Average cloud cover over scene: {r['avg_cloud']}
- SAR backscatter change threshold applied: {r['sar_threshold']:.0f} dB

Write a structured five-element impact brief. Use exactly these five headings.

## 1. Flood Extent
State the total new flood area in km². Distinguish permanent water from new inundation. Use the numbers provided.

## 2. Infrastructure Exposure
Based on the region and flood extent, identify infrastructure categories most likely in the affected zone. Rate exposure HIGH, MEDIUM, or LOW for each category with a one-sentence justification.

## 3. Sensor Justification
Explain why SAR was the appropriate sensor for this event. Reference the actual cloud cover data provided. Explain the C-band wavelength vs cloud droplet physics. Explain the backscatter change principle. Make this educational for a coordinator who does not know SAR.

## 4. Response Recommendation
State what a field team or emergency coordinator should do with this output, in priority order. Include at least one ground-truth verification step the satellite cannot confirm.

## 5. Confidence and Limitations
State clearly what this analysis cannot confirm without ground truth. Address the threshold sensitivity, what may be overcounted, what may be undercounted, and the 500m scale limitation.

End with:
DATA QUALITY: SAR scenes (before/after): {r['before_count']}/{r['after_count']} | Optical scenes: {r['s2_count']} | Cloud cover: {r['avg_cloud']} | Confidence: {r['confidence']}
"""


def _fallback_brief(results: dict) -> str:
    r = results
    return f"""## 1. Flood Extent

Total new flood area detected: **{r['area_total_flood']:.0f} km²**. Permanent water baseline (rivers, canals, lakes): {r['area_permanent']:.0f} km². SAR-detected new inundation: {r['area_sar_flood']:.0f} km². Optically confirmed new water: {r['area_confirmed']:.0f} km².

## 2. Infrastructure Exposure

- **Roads and transport:** HIGH — rural road networks in floodplain regions follow irrigation canal alignments. Inundation of {r['area_total_flood']:.0f} km² implies significant surface access disruption.
- **Agricultural land:** HIGH — floodplain regions in this geography are primarily agricultural. Crop damage and soil erosion are primary near-term impacts.
- **Settlements:** MEDIUM-HIGH — rural settlements in floodplain zones are vulnerable. Urban areas may have partial flood defences.
- **Power distribution:** MEDIUM — distribution lines typically follow road corridors. Substation flooding requires field verification.
- **Water infrastructure:** HIGH — irrigation canals may overflow and silt during recession. Drinking water infrastructure is a priority for assessment.

## 3. Sensor Justification

Sentinel-1 SAR operates at C-band (~5.6 cm wavelength). Cloud water droplets are 10–100 micrometres — more than 500 times smaller than the radar wavelength. The radar pulse passes through cloud without scattering. During this event, average cloud cover was approximately {r['avg_cloud']} over the scene. {r['s2_count']} cloud-free optical scenes were available — {'none, making SAR the sole viable detection instrument' if r['s2_count'] == 0 else 'confirming SAR detections where coverage existed'}.

The flood signal is backscatter change. Dry land returns a strong VV signal. Flooded land returns weak signal — smooth water reflects radar away from the sensor. A threshold of {r['sar_threshold']:.0f} dB was applied, representing a halving of radar return power.

## 4. Response Recommendation

1. Confirm the flood boundary at two to three accessible points along the SAR-detected edge. Satellite boundary accuracy is approximately 500m at analysis scale.
2. Use the permanent water layer to distinguish new water from the pre-existing river and canal system.
3. Prioritise reconnaissance of settlement areas within the flood boundary.
4. Check road corridors at flood boundary crossings before committing vehicles. SAR detects water extent, not depth.
5. Re-run analysis after 7–10 days using the next SAR pass to assess whether extent is growing, stable, or receding.

## 5. Confidence and Limitations

- **Depth:** Not measurable from SAR backscatter change. Analysis detects surface extent only.
- **Urban flooding:** Systematically undercounted. Building corner reflectors dominate SAR backscatter in built-up areas.
- **Vegetated flooding:** Complex backscatter response. Slope and permanent water masks reduce but do not eliminate uncertainty.
- **Threshold sensitivity:** At {r['sar_threshold']:.0f} dB, increasing to −5 dB reduces area and false positives; decreasing to −2 dB increases area and sensitivity to shallow flooding.
- **Scale:** Analysis at 500m. SAR native resolution is ~10m. Mixed pixels at boundaries undercount total extent.

DATA QUALITY: SAR scenes (before/after): {r['before_count']}/{r['after_count']} | Optical scenes: {r['s2_count']} | Cloud cover: {r['avg_cloud']} | Confidence: {r['confidence']}
"""


# ---------------------------------------------------------------------------
# Word document builder
# ---------------------------------------------------------------------------

def _add_inline_bold(paragraph, text):
    """Render **bold** spans correctly inside a Word paragraph."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def _build_word_doc(results: dict, brief_text: str, model_used: str,
                    chart_bytes: bytes, sar_fig_bytes: bytes) -> bytes:
    """Build a formatted Word document with stats, charts, and impact brief."""
    doc = Document()

    # Margins
    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    r = results

    # Title
    title = doc.add_paragraph()
    run   = title.add_run('Flood Intelligence Brief')
    run.bold      = True
    run.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.add_run('Event: ').bold = True
    meta.add_run(r['event_key'])
    doc.add_paragraph(r['event']['context'])
    doc.add_paragraph()

    # Stats table
    p = doc.add_paragraph()
    p.add_run('Layer 2 — Flood Extent Statistics').bold = True
    p.runs[0].font.size = Pt(13)

    tbl = doc.add_table(rows=7, cols=2)
    tbl.style = 'Table Grid'
    rows_data = [
        ('Metric',                        'Value'),
        ('Permanent water baseline',       f"{r['area_permanent']:,.0f} km²"),
        ('SAR-only new flood',             f"{r['area_sar_flood']:,.0f} km²"),
        ('Confirmed flood (SAR + NDWI)',   f"{r['area_confirmed']:,.0f} km²"),
        ('Total new flood area',           f"{r['area_total_flood']:,.0f} km²"),
        ('Confidence level',               r['confidence']),
        ('Confidence note',                r['confidence_note']),
    ]
    for i, (label, value) in enumerate(rows_data):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = value
        if i == 0:
            for cell in tbl.rows[i].cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()

    # Data quality table
    p2 = doc.add_paragraph()
    p2.add_run('Data Quality').bold = True
    p2.runs[0].font.size = Pt(13)

    qtbl = doc.add_table(rows=5, cols=2)
    qtbl.style = 'Table Grid'
    q_rows = [
        ('Parameter',                  'Value'),
        ('SAR scenes — before period', str(r['before_count'])),
        ('SAR scenes — after period',  str(r['after_count'])),
        ('Cloud-free optical scenes',  str(r['s2_count'])),
        ('Average cloud cover',        r['avg_cloud']),
    ]
    for i, (label, value) in enumerate(q_rows):
        qtbl.rows[i].cells[0].text = label
        qtbl.rows[i].cells[1].text = value
        if i == 0:
            for cell in qtbl.rows[i].cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()

    # Extent bar chart
    p3 = doc.add_paragraph()
    p3.add_run('Flood Extent Chart').bold = True
    p3.runs[0].font.size = Pt(13)
    if chart_bytes:
        doc.add_picture(io.BytesIO(chart_bytes), width=Inches(5.0))
    doc.add_paragraph()

    # SAR thumbnails
    p4 = doc.add_paragraph()
    p4.add_run('Layer 1 — SAR Imagery').bold = True
    p4.runs[0].font.size = Pt(13)
    if sar_fig_bytes:
        doc.add_picture(io.BytesIO(sar_fig_bytes), width=Inches(6.0))

    # Spacer + separator before flood classification map
    doc.add_paragraph()
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(4)
    sep.paragraph_format.space_after  = Pt(4)
    sep_run = sep.add_run('─' * 60)
    sep_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    sep_run.font.size      = Pt(9)

    p_map = doc.add_paragraph()
    p_map.add_run('Flood Classification Map').bold = True
    p_map.runs[0].font.size = Pt(11)
    cap = doc.add_paragraph('Blue = permanent water  ·  Orange = SAR-only flood  ·  Red = confirmed flood (SAR + NDWI)')
    cap.runs[0].font.size   = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    flood_raw = _fetch_png(r['url_flood'])
    if flood_raw:
        try:
            doc.add_picture(io.BytesIO(flood_raw), width=Inches(3.5))
        except Exception:
            doc.add_paragraph('[Flood classification map — image unavailable]')
    doc.add_paragraph()

    # Impact brief
    p5 = doc.add_paragraph()
    p5.add_run('Layer 3 — AI Impact Brief').bold = True
    p5.runs[0].font.size = Pt(13)
    if model_used:
        mi = doc.add_paragraph()
        mi.add_run(f'Generated by: {model_used}').italic = True

    lines = brief_text.splitlines()
    i = 0
    while i < len(lines):
        line    = lines[i]
        stripped = line.strip()

        if stripped.startswith('## '):
            p = doc.add_paragraph(stripped[3:])
            p.style = doc.styles['Heading 2']
            i += 1
        elif stripped.startswith('# '):
            p = doc.add_paragraph(stripped[2:])
            p.style = doc.styles['Heading 1']
            i += 1
        elif stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_bold(p, stripped[2:])
            i += 1
        elif stripped.startswith('DATA QUALITY:'):
            p = doc.add_paragraph()
            p.add_run(stripped).bold = True
            i += 1
        elif stripped == '' or stripped.startswith('---'):
            doc.add_paragraph()
            i += 1
        else:
            if stripped:
                p = doc.add_paragraph()
                _add_inline_bold(p, stripped)
            i += 1

    footer = doc.add_paragraph()
    footer.add_run('Generated by EOIL — AI-Native Earth Observation Innovation Lab').italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _build_markdown(results: dict, brief_text: str, model_used: str) -> str:
    r = results
    lines = [
        f"# Flood Intelligence Brief",
        f"",
        f"**Event:** {r['event_key']}",
        f"**Context:** {r['event']['context']}",
        f"",
        f"---",
        f"",
        f"## Layer 2 — Flood Extent Statistics",
        f"",
        f"| Metric | Value |",
        f"|--------|-------|",
        f"| Permanent water baseline | {r['area_permanent']:,.0f} km² |",
        f"| SAR-only new flood | {r['area_sar_flood']:,.0f} km² |",
        f"| Confirmed flood (SAR+NDWI) | {r['area_confirmed']:,.0f} km² |",
        f"| Total new flood area | {r['area_total_flood']:,.0f} km² |",
        f"| Confidence level | {r['confidence']} |",
        f"",
        f"**Confidence note:** {r['confidence_note']}",
        f"",
        f"## Data Quality",
        f"",
        f"| Parameter | Value |",
        f"|-----------|-------|",
        f"| SAR scenes — before period | {r['before_count']} |",
        f"| SAR scenes — after period | {r['after_count']} |",
        f"| Cloud-free optical scenes | {r['s2_count']} |",
        f"| Average cloud cover | {r['avg_cloud']} |",
        f"",
        f"---",
        f"",
        f"## Layer 3 — AI Impact Brief",
        f"",
        f"*Generated by: {model_used or 'built-in fallback'}*",
        f"",
        brief_text,
    ]
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Main render function
# ---------------------------------------------------------------------------

def render():
    """Render the Flood Intelligence module. Called by app.py."""

    st.markdown("## 🌊 Flood Intelligence")
    st.markdown(
        "SAR-based flood detection and impact mapping. "
        "Sentinel-1 radar penetrates cloud cover — the only viable sensor "
        "during a flood event, when cloud cover is heaviest."
    )

    gee_ok = st.session_state.get("gee_available", False)
    if gee_ok:
        st.caption("🟢 GEE connected — live Sentinel-1 and Sentinel-2 data active.")
    else:
        st.caption("🔴 GEE not connected. This module requires Earth Engine credentials.")

    # ------------------------------------------------------------------
    # Controls — all on the main page
    # ------------------------------------------------------------------
    with st.expander("ℹ️ What does this module do?", expanded=False):
        st.markdown("""
**This module answers one question: where did new flooding occur, and what does it mean?**

It combines four data signals to isolate genuinely new inundation:

1. **SAR backscatter change** — Sentinel-1 radar detects surface smoothing caused by flood water, through cloud cover.
2. **NDWI** — Sentinel-2 optical water index confirms standing water where cloud-free scenes are available.
3. **Slope mask** — Flat terrain floods. Steep terrain drains. SRTM elevation data removes slopes above 5 degrees.
4. **Permanent water mask** — JRC Global Surface Water removes rivers, lakes, and canals that existed before the event.

The AI impact brief translates those measurements into a decision-ready output for field teams and coordinators.
        """)

    st.divider()

    event_key = st.selectbox(
        "Select flood event",
        list(FLOOD_EVENTS.keys()),
        key="fi_event_key",
        help="Each event has validated Sentinel-1 coverage and known impact context."
    )
    event = FLOOD_EVENTS[event_key]

    col_b, col_a = st.columns(2)
    with col_b:
        st.info(
            f"**Before period:** {event['before_start']} to {event['before_end']}  \n"
            f"**After period:** {event['after_start']} to {event['after_end']}  \n"
            f"**SAR orbit:** {event['orbit']}"
        )
    with col_a:
        st.info(f"**Context:** {event['context']}")

    sar_threshold = st.slider(
        "SAR backscatter change threshold (dB)",
        min_value=-8.0, max_value=-1.0,
        value=float(event["sar_threshold"]),
        step=0.5,
        key="fi_threshold",
        help="Pixels with backscatter change below this value are classified as flood candidates. −3 dB = backscatter halved."
    )
    st.caption(
        f"Current threshold: **{sar_threshold:.1f} dB**. "
        "More negative = stricter filter, smaller area. "
        "Less negative = more permissive, larger area. "
        "Try −5 dB and −2 dB to see the sensitivity range."
    )

    run_btn = st.button("▶ Run Flood Analysis", type="primary", key="fi_run")

    # ------------------------------------------------------------------
    # Two-step run pattern
    # ------------------------------------------------------------------
    cache_key = f"fi_results_{event_key}_{sar_threshold}"

    if run_btn:
        st.session_state["fi_pending"]     = cache_key
        st.session_state["fi_pending_key"] = event_key
        st.session_state["fi_pending_thr"] = sar_threshold
        st.rerun()

    if st.session_state.get("fi_pending") == cache_key:
        del st.session_state["fi_pending"]

        if not gee_ok:
            st.error("GEE not available. Cannot run flood analysis without Earth Engine credentials.")
            st.stop()

        with st.spinner(
            "Fetching Sentinel-1 SAR composites and running flood classification — allow 60–90 seconds..."
        ):
            try:
                results = _run_flood_analysis(
                    st.session_state["fi_pending_key"],
                    st.session_state["fi_pending_thr"],
                )
                st.session_state[cache_key] = results
            except Exception as ex:
                st.error(f"Flood analysis failed: {ex}")
                st.stop()

    # ------------------------------------------------------------------
    # Results
    # ------------------------------------------------------------------
    if cache_key not in st.session_state:
        st.info("Select an event and click **▶ Run Flood Analysis** to begin.")
        with st.expander("Why SAR for flood detection?", expanded=True):
            st.markdown("""
**The cloud problem.**

Every major flood event is accompanied by heavy cloud cover — the same weather system
that causes the flooding also blocks optical satellites. Sentinel-2 and Landsat detect
reflected sunlight. Cloud blocks that completely. During a flood at its worst, they see nothing.

**Why SAR is different.**

Sentinel-1 transmits its own microwave energy at C-band (~5.6 cm wavelength) and measures
what returns. Cloud water droplets are 10–100 micrometres — more than 500 times smaller
than the radar wavelength. The pulse passes through cloud without scattering.

**The flood signal.**

Dry land returns strong radar. Rough surfaces scatter the pulse back from many angles.
Flooded land returns almost nothing — smooth water reflects the pulse away from the sensor
like a mirror. The change in backscatter between a before and after scene is the flood signal.
A drop of −3 dB means returned power halved.
            """)
        return

    results = st.session_state[cache_key]
    r       = results

    def section_break():
        st.markdown(
            '<hr style="border: none; border-top: 3px solid #d0d0d0; margin: 24px 0 16px 0;">',
            unsafe_allow_html=True,
        )

    # ------------------------------------------------------------------
    # Layer 1: SAR thumbnails — three in one row
    # ------------------------------------------------------------------
    section_break()
    st.subheader("📡 Layer 1 — Signal Processing")
    st.caption(
        "Three SAR panels below show the radar signal before the flood, after the flood, "
        "and the change between them. These are derived from Sentinel-1 median composites — "
        "not single scenes. The median across multiple passes removes noise and transient effects."
    )

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("SAR scenes (before)", r["before_count"])
    with col2:
        st.metric("SAR scenes (after)", r["after_count"])
    with col3:
        st.metric("Optical scenes (<20% cloud)", r["s2_count"])
    with col4:
        st.metric("Avg cloud cover", r["avg_cloud"])

    sar_fig_bytes = _build_sar_row_fig(results)
    st.image(sar_fig_bytes, use_container_width=True)

    with st.expander("ℹ️ How to read the SAR panels", expanded=False):
        st.markdown("""
**SAR Before and SAR After** show radar backscatter in greyscale.
Dark pixels are smooth surfaces — calm water, bare soil, flooded land.
Bright pixels are rough surfaces — vegetation, buildings, bare rock.

In the Before panel, the landscape shows its dry-season texture.
In the After panel, areas that flooded have become darker because flood water is
a smooth mirror that reflects radar away from the sensor.

**SAR Backscatter Change** shows the difference — After minus Before.
Red pixels had lower backscatter after the flood. That is the flood detection signal.
Blue pixels had higher backscatter after — possibly from wind roughening dry soil or crop growth.
The red areas are your flood candidates before the slope and water masks are applied.
        """)

    # ------------------------------------------------------------------
    # Flood classification map — its own section
    # ------------------------------------------------------------------
    section_break()
    st.subheader("🗺️ Flood Classification Map")
    st.markdown(
        "This map is different from the SAR panels above. The SAR panels show raw radar signal. "
        "This map shows the **final decision** after applying all four filters: "
        "SAR change threshold, slope mask, permanent water mask, and NDWI confirmation. "
        "Only pixels that passed all relevant tests appear as flood."
    )

    col_map, col_legend = st.columns([2, 1])
    with col_map:
        flood_raw = _fetch_png(r["url_flood"])
        if flood_raw:
            st.image(flood_raw, width=380,
                     caption="Flood classification map")
        else:
            st.warning("Flood classification map unavailable.")
    with col_legend:
        st.markdown("**Map legend**")
        st.markdown("🟦 Permanent water")
        st.markdown("&nbsp;&nbsp;&nbsp;&nbsp;(rivers, canals, lakes)")
        st.markdown("🟧 SAR-only flood")
        st.markdown("&nbsp;&nbsp;&nbsp;&nbsp;(new water, SAR detected)")
        st.markdown("🟥 Confirmed flood")
        st.markdown("&nbsp;&nbsp;&nbsp;&nbsp;(SAR + NDWI agree)")
        st.markdown("⬜ Not flooded")
        st.markdown("&nbsp;")
        st.markdown(
            f"**Threshold:** {r['sar_threshold']:.1f} dB  \n"
            f"**Slope mask:** <5°  \n"
            f"**Perm. water:** >75% occurrence"
        )

    # ------------------------------------------------------------------
    # Layer 2: Classification statistics
    # ------------------------------------------------------------------
    section_break()
    st.subheader("📊 Layer 2 — Flood Extent Classification")
    st.caption(
        "Area statistics after applying all four filters. "
        "New flood area excludes permanent water — it represents genuinely new inundation only."
    )

    col_a, col_b, col_c, col_d = st.columns(4)
    with col_a:
        st.metric("Permanent water", f"{r['area_permanent']:,.0f} km²")
    with col_b:
        st.metric("SAR-only flood", f"{r['area_sar_flood']:,.0f} km²")
    with col_c:
        st.metric("Confirmed flood", f"{r['area_confirmed']:,.0f} km²")
    with col_d:
        st.metric("Total new flood", f"{r['area_total_flood']:,.0f} km²",
                  delta=r["confidence"], delta_color="off")

    conf_icon = {"HIGH": "🟢", "MEDIUM": "🟡", "LOW": "🔴"}.get(r["confidence"], "⚪")
    st.info(f"{conf_icon} **Confidence: {r['confidence']}** — {r['confidence_note']}")

    chart_bytes = _build_extent_chart(results)
    st.image(chart_bytes, use_column_width=False, width=500)

    with st.expander("ℹ️ What did each algorithm step do?", expanded=False):
        st.markdown(f"""
**SAR backscatter change — threshold: {r['sar_threshold']:.1f} dB**
Pixels where VV backscatter dropped by {abs(r['sar_threshold']):.1f} dB or more are flagged as candidates.
{abs(r['sar_threshold']):.1f} dB represents a {"halving" if r['sar_threshold'] == -3 else "major reduction"} in radar return power.
The before composite used {r['before_count']} scenes; the after composite used {r['after_count']} scenes.

**Slope mask — SRTM DEM, threshold: 5 degrees**
Flood water ponds on flat terrain. Slopes above 5° drain too quickly for surface ponding to persist.
This also removes SAR geometry artefacts: steep slopes produce layover and shadow that mimic water.

**Permanent water mask — JRC Global Surface Water, >75% occurrence**
Pixels flagged as water in >75% of Landsat observations (1984–2021) are classified as permanent water.
The {r['area_permanent']:,.0f} km² baseline includes the main river channel and irrigation network.
Without this mask, the existing water infrastructure would dominate the flood statistics.

**NDWI optical confirmation — threshold: 0.2**
Where cloud-free Sentinel-2 scenes were available ({r['s2_count']} scenes), NDWI > 0.2 confirms water.
{"No cloud-free optical scenes were available — SAR was the sole detection instrument." if r['s2_count'] == 0 else f"Where SAR and NDWI agreed: {r['area_confirmed']:,.0f} km² confirmed. SAR-only: {r['area_sar_flood']:,.0f} km²."}
        """)

    # ------------------------------------------------------------------
    # Layer 3: AI impact brief
    # ------------------------------------------------------------------
    section_break()
    st.subheader("🤖 Layer 3 — AI Impact Brief")
    st.caption(
        "Five-element structured brief for emergency coordinators and field teams. "
        "Every claim is grounded in the Layer 2 measurements above."
    )

    brief_key = f"fi_brief_{cache_key}"
    model_key = f"fi_model_{cache_key}"

    if st.button("Get AI Impact Brief", type="primary", key="fi_ai_btn"):
        with st.spinner("Generating impact brief..."):
            prompt    = _build_impact_prompt(results)
            text, model_used = ai_chain.complete(
                prompt,
                groq_key=config.GROQ_API_KEY,
                gemini_key=config.GEMINI_API_KEY,
            )
            if text:
                st.session_state[brief_key] = text
                st.session_state[model_key] = model_used
            else:
                st.session_state[brief_key] = _fallback_brief(results)
                st.session_state[model_key] = None

    if brief_key in st.session_state:
        brief_text = st.session_state[brief_key]
        model_used = st.session_state.get(model_key)

        with st.expander("📋 AI Impact Brief", expanded=True):
            st.markdown(brief_text)
            if model_used:
                st.caption(f"AI response from {model_used}")
            else:
                st.caption("Built-in analysis — no AI key present")

        # Downloads
        section_break()
        st.markdown("### Downloads")
        dl1, dl2 = st.columns(2)

        with dl1:
            md_content = _build_markdown(results, brief_text, model_used or "built-in fallback")
            safe_name  = event_key[:25].replace(' ', '_').replace(',', '').replace('—', '')
            st.download_button(
                label="⬇ Download brief (.md)",
                data=md_content.encode('utf-8'),
                file_name=f"flood-intelligence-{safe_name}.md",
                mime="text/markdown",
                use_container_width=True,
            )

        with dl2:
            word_bytes = _build_word_doc(
                results, brief_text, model_used or "built-in fallback",
                chart_bytes, sar_fig_bytes
            )
            st.download_button(
                label="⬇ Download brief (.docx)",
                data=word_bytes,
                file_name=f"flood-intelligence-{safe_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
