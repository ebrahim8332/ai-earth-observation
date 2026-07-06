"""
drone_intelligence.py — Drone Intelligence module for the EOIL portal.

Arc 6: Drone Intelligence.

Two sub-modules on one page:

  Sub-module A — Infrastructure Inspection
    Layer 1: Simulate co-registered RGB + thermal rasters over a transmission tower.
             Compute delta-T per component zone against ambient baseline.
    Layer 2: Isolation Forest on thermal pixels (unsupervised anomaly detection).
             SLIC superpixel segmentation as SAM proxy (component boundary delineation).
             Cross-reference: each segment gets mean thermal value and CIGRE severity class.
    Layer 3: Groq structured inspection brief — critical findings, warnings, sensor limits.

  Sub-module B — Vegetation Mapping
    Layer 1: Simulate drone RGB+NIR orthomosaic over Catawba Valley corridor (Arc 5 area).
             Compute NDVI and ExG at 10 cm resolution. Build canopy density map.
    Layer 2: Watershed segmentation for individual crown delineation.
             Random Forest pixel classification (bare soil / low veg / shrub / canopy).
             Compare crown count and density with Arc 5 LiDAR DBSCAN results.
    Layer 3: Groq vegetation management brief — what drone adds, what LiDAR adds, how to combine.

All data is simulated. Modelled on real UAV inspection and corridor survey datasets.
All simulation is labelled clearly in the UI.

Algorithm logic is documented in detail in notebook 10_drone_intelligence.ipynb.
"""

import io
import re
import warnings
warnings.filterwarnings('ignore')

import numpy as np
import streamlit as st
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.colors import ListedColormap
from scipy.ndimage import gaussian_filter, distance_transform_edt, uniform_filter
from scipy.signal import find_peaks
from skimage.segmentation import slic, mark_boundaries, watershed
from skimage.feature import peak_local_max
from skimage.filters import gaussian as sk_gaussian
from skimage.measure import regionprops
from sklearn.ensemble import IsolationForest, RandomForestClassifier
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt

import config
import ai_chain

ASSETS_DIR = Path(__file__).parent / "assets"
AMBIENT_TEMP = 24.0   # degrees Celsius — ambient air baseline for all delta-T calculations

# ---------------------------------------------------------------------------
# Simulation helpers — same logic as notebook, re-runnable in-memory
# ---------------------------------------------------------------------------

def _simulate_tower_scene(seed: int = 42):
    """
    Generate a simulated 200×200 pixel co-registered RGB + thermal raster pair.
    5 cm/pixel resolution. Returns (rgb, thermal) as numpy arrays.
    """
    rng = np.random.default_rng(seed=seed)
    H, W = 200, 200

    # RGB — sky background, tower body, crossarms, insulators, conductors
    rgb = np.ones((H, W, 3), dtype=np.float32)
    rgb[:, :, 0] = 0.55; rgb[:, :, 1] = 0.60; rgb[:, :, 2] = 0.70
    sky_noise = rng.normal(0, 0.03, (H, W))
    for c in range(3):
        rgb[:, :, c] += sky_noise
    rgb[0:180, 95:105, :] = [0.35, 0.35, 0.35]
    rgb[38:46, 40:160, :]  = [0.40, 0.40, 0.40]
    rgb[78:86, 50:150, :]  = [0.40, 0.40, 0.40]
    rgb[36:50, 46:56, :]   = [0.90, 0.88, 0.80]
    rgb[36:50, 144:154, :] = [0.90, 0.88, 0.80]
    rgb[76:90, 56:66, :]   = [0.90, 0.88, 0.80]
    rgb[76:90, 134:144, :] = [0.90, 0.88, 0.80]
    rgb[43, 0:46, :]   = [0.65, 0.65, 0.65]
    rgb[43, 154:200, :] = [0.65, 0.65, 0.65]
    rgb[83, 0:56, :]   = [0.65, 0.65, 0.65]
    rgb[83, 144:200, :] = [0.65, 0.65, 0.65]
    rgb = np.clip(rgb, 0, 1)

    # Thermal — realistic per-component temperatures
    thermal = rng.normal(AMBIENT_TEMP, 1.5, (H, W)).astype(np.float32)
    thermal[0:180, 95:105]  = rng.normal(28.0, 0.8, (180, 10))
    thermal[38:46, 40:160]  = rng.normal(30.0, 1.0, (8, 120))
    thermal[78:86, 50:150]  = rng.normal(30.0, 1.0, (8, 100))
    thermal[36:50, 46:56]   = rng.normal(26.0, 0.5, (14, 10))   # top-left — healthy
    thermal[36:50, 144:154] = rng.normal(58.0, 2.0, (14, 10))   # top-right — CRITICAL
    thermal[76:90, 56:66]   = rng.normal(39.0, 1.0, (14, 10))   # mid-left  — WARNING
    thermal[76:90, 134:144] = rng.normal(27.0, 0.5, (14, 10))   # mid-right — healthy
    thermal[43, 0:46]   = rng.normal(32.0, 1.0, (46,))
    thermal[43, 154:200] = rng.normal(32.0, 1.0, (46,))
    thermal[83, 0:56]   = rng.normal(31.0, 0.8, (56,))
    thermal[83, 144:200] = rng.normal(45.0, 2.0, (56,))   # mid-right conductor — WARM
    thermal = gaussian_filter(thermal, sigma=1.5)
    return rgb, thermal, H, W


ZONES = {
    "Tower Body":           (0,   180, 93,  107),
    "Top Crossarm":         (36,  48,  38,  162),
    "Mid Crossarm":         (76,  88,  48,  152),
    "Insulator Top-Left":   (34,  52,  44,  58),
    "Insulator Top-Right":  (34,  52,  142, 156),
    "Insulator Mid-Left":   (74,  92,  54,  68),
    "Insulator Mid-Right":  (74,  92,  132, 146),
    "Conductor Top-Left":   (41,  45,  0,   44),
    "Conductor Top-Right":  (41,  45,  156, 200),
    "Conductor Mid-Left":   (81,  85,  0,   54),
    "Conductor Mid-Right":  (81,  85,  146, 200),
}


def _classify_delta_t(dt: float) -> str:
    if dt < 5:   return "Normal"
    if dt < 15:  return "Monitor"
    if dt < 30:  return "Warning"
    return "CRITICAL"


def _compute_zone_stats(thermal):
    stats = {}
    for name, (r0, r1, c0, c1) in ZONES.items():
        patch = thermal[r0:r1, c0:c1]
        mean_t = float(patch.mean())
        dt = mean_t - AMBIENT_TEMP
        stats[name] = {"mean_C": round(mean_t, 1), "delta_T": round(dt, 1),
                       "class": _classify_delta_t(dt)}
    return stats


def _simulate_vegetation_scene(seed: int = 42):
    """
    Generate a simulated 200×300 pixel drone vegetation orthomosaic.
    10 cm/pixel. Returns R, G, B, NIR, NDVI, ExG, crown list.
    """
    rng = np.random.default_rng(seed=seed)
    VH, VW = 200, 300
    PIXEL_M = 0.10

    R   = rng.normal(0.22, 0.03, (VH, VW)).astype(np.float32)
    G   = rng.normal(0.18, 0.03, (VH, VW)).astype(np.float32)
    B   = rng.normal(0.14, 0.02, (VH, VW)).astype(np.float32)
    NIR = rng.normal(0.25, 0.04, (VH, VW)).astype(np.float32)

    R[70:130, :]   += rng.normal(-0.02, 0.01, (60, VW))
    G[70:130, :]   += rng.normal(+0.03, 0.01, (60, VW))
    NIR[70:130, :] += rng.normal(+0.05, 0.02, (60, VW))

    crown_params = [
        (20,25,12),(15,70,10),(25,120,14),(18,175,11),(22,230,13),(12,270,9),
        (30,50,8),(25,155,10),(40,100,12),(35,200,11),(45,260,9),
        (62,40,11),(65,140,10),(60,220,12),
        (150,30,13),(160,80,10),(155,130,12),(165,190,11),(158,245,10),
        (170,60,9),(175,160,13),(145,110,11),(185,220,10),(180,270,8),
        (138,55,10),(135,175,11),(140,255,9),
    ]
    rows_grid, cols_grid = np.mgrid[0:VH, 0:VW]
    for (cy, cx, cr) in crown_params:
        dist   = np.sqrt((rows_grid - cy)**2 + (cols_grid - cx)**2)
        weight = np.clip(np.exp(-0.5 * (dist / (cr * 0.7))**2), 0, 1)
        R   += weight * rng.normal(-0.08, 0.01, (VH, VW))
        G   += weight * rng.normal(+0.10, 0.02, (VH, VW))
        B   += weight * rng.normal(-0.02, 0.01, (VH, VW))
        NIR += weight * rng.normal(+0.35, 0.04, (VH, VW))

    R   = np.clip(R,   0, 1)
    G   = np.clip(G,   0, 1)
    B   = np.clip(B,   0, 1)
    NIR = np.clip(NIR, 0, 1)
    eps = 1e-8
    NDVI = (NIR - R) / (NIR + R + eps)
    ExG  = 2.0 * G - R - B
    return R, G, B, NIR, NDVI, ExG, crown_params, VH, VW, PIXEL_M


# ---------------------------------------------------------------------------
# Figure builders — return PNG bytes for display and Word doc
# ---------------------------------------------------------------------------

def _fig_to_bytes(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=130, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _build_rgb_thermal_fig(rgb, thermal) -> bytes:
    fig, axes = plt.subplots(1, 2, figsize=(13, 5))
    axes[0].imshow(rgb)
    axes[0].set_title("RGB Orthomosaic — 5 cm/pixel", fontweight='bold')
    axes[0].text(2, 195, "[SIMULATED DATA]", color='yellow', fontsize=7,
                 bbox=dict(boxstyle='round', fc='black', alpha=0.6))
    axes[0].annotate('Tower body', xy=(100,90), xytext=(125,90),
                     arrowprops=dict(arrowstyle='->', color='white'), color='white', fontsize=8)
    axes[0].annotate('Faulty insulator', xy=(149,43), xytext=(158,18),
                     arrowprops=dict(arrowstyle='->', color='yellow'), color='yellow', fontsize=8)
    im = axes[1].imshow(thermal, cmap='inferno', vmin=20, vmax=65)
    axes[1].set_title("Thermal Raster — Surface Temperature (°C)", fontweight='bold')
    axes[1].text(2, 195, "[SIMULATED DATA]", color='yellow', fontsize=7,
                 bbox=dict(boxstyle='round', fc='black', alpha=0.6))
    axes[1].annotate('CRITICAL 58°C', xy=(149,43), xytext=(158,20),
                     arrowprops=dict(arrowstyle='->', color='cyan'), color='cyan', fontsize=8)
    axes[1].annotate('WARNING 39°C', xy=(61,83), xytext=(8,100),
                     arrowprops=dict(arrowstyle='->', color='orange'), color='orange', fontsize=8)
    plt.colorbar(im, ax=axes[1], label='Temperature (°C)', fraction=0.046, pad=0.04)
    plt.suptitle("Layer 1 — Signal Processing: Co-registered RGB and Thermal Rasters",
                 fontsize=11, fontweight='bold')
    plt.tight_layout()
    return _fig_to_bytes(fig)


def _build_isolation_forest_fig(thermal, H, W) -> tuple[bytes, np.ndarray, dict]:
    """Run Isolation Forest and return (figure bytes, anomaly_map, per-zone anomaly stats)."""
    X = thermal.reshape(-1, 1)
    iso = IsolationForest(n_estimators=100, contamination=0.03, random_state=42)
    iso.fit(X)
    predictions = iso.predict(X)
    scores      = iso.decision_function(X)
    anomaly_map = (predictions == -1).reshape(H, W).astype(np.uint8)
    score_map   = scores.reshape(H, W)

    zone_anomaly = {}
    for name, (r0, r1, c0, c1) in ZONES.items():
        patch = anomaly_map[r0:r1, c0:c1]
        frac  = float(patch.sum()) / patch.size if patch.size > 0 else 0.0
        zone_anomaly[name] = round(frac * 100, 1)

    fig, axes = plt.subplots(1, 3, figsize=(16, 5))
    im1 = axes[0].imshow(thermal, cmap='inferno', vmin=20, vmax=65)
    axes[0].set_title("Raw Thermal", fontweight='bold')
    plt.colorbar(im1, ax=axes[0], label='°C', fraction=0.046)

    axes[1].imshow(-score_map, cmap='hot')
    axes[1].set_title("Anomaly Score\n(bright = more anomalous)", fontweight='bold')

    axes[2].imshow(thermal, cmap='inferno', vmin=20, vmax=65)
    anom_overlay = np.ma.masked_where(anomaly_map == 0, anomaly_map)
    axes[2].imshow(anom_overlay, cmap='cool', alpha=0.6, vmin=0, vmax=1)
    axes[2].set_title("Flags on Thermal\n(cyan = anomalous pixel)", fontweight='bold')

    for ax in axes:
        ax.text(2, 196, "[SIMULATED]", color='white', fontsize=7,
                bbox=dict(boxstyle='round', fc='black', alpha=0.5))

    plt.suptitle("Layer 2 — Isolation Forest: Unsupervised Thermal Anomaly Detection",
                 fontsize=11, fontweight='bold')
    plt.tight_layout()
    return _fig_to_bytes(fig), anomaly_map, zone_anomaly


def _build_segmentation_fig(rgb, thermal, H, W) -> tuple[bytes, list]:
    """Run SLIC segmentation, classify each segment, return (figure bytes, flagged list)."""
    segments = slic(rgb, n_segments=80, compactness=20, sigma=1, start_label=1)
    n_seg    = segments.max()

    class_int = {"Normal": 0, "Monitor": 1, "Warning": 2, "CRITICAL": 3}
    severity_img = np.zeros((H, W), dtype=np.int32)
    flagged = []

    for seg_id in range(1, n_seg + 1):
        mask   = (segments == seg_id)
        if not mask.any():
            continue
        mean_t = float(thermal[mask].mean())
        dt     = mean_t - AMBIENT_TEMP
        cls    = _classify_delta_t(dt)
        severity_img[mask] = class_int[cls]
        rows, cols = np.where(mask)
        cy, cx = rows.mean(), cols.mean()
        if cls in ("Warning", "CRITICAL"):
            flagged.append({"class": cls, "mean_C": round(mean_t, 1),
                            "delta_T": round(dt, 1), "cx": cx, "cy": cy})

    seg_cmap = ListedColormap(['#2ecc71', '#f1c40f', '#e67e22', '#e74c3c'])
    fig, axes = plt.subplots(1, 3, figsize=(16, 5))

    axes[0].imshow(mark_boundaries(rgb, segments, color=(0, 1, 1), mode='thick'))
    axes[0].set_title("Segment Boundaries on RGB\n(SLIC superpixel — SAM proxy)",
                      fontweight='bold')

    im2 = axes[1].imshow(severity_img, cmap=seg_cmap, vmin=-0.5, vmax=3.5)
    axes[1].set_title("Segment Thermal Classification\nNormal / Monitor / Warning / Critical",
                      fontweight='bold')
    cbar = plt.colorbar(im2, ax=axes[1], fraction=0.046, ticks=[0, 1, 2, 3])
    cbar.ax.set_yticklabels(['Normal', 'Monitor', 'Warning', 'CRITICAL'])

    axes[2].imshow(thermal, cmap='inferno', vmin=20, vmax=65)
    for rec in flagged:
        colour = 'red' if rec['class'] == 'CRITICAL' else 'orange'
        axes[2].plot(rec['cx'], rec['cy'], 's', color=colour,
                     markersize=9, markeredgecolor='white')
        axes[2].text(rec['cx'] + 3, rec['cy'], rec['class'],
                     color=colour, fontsize=7, fontweight='bold')
    axes[2].set_title("Flagged Segments on Thermal\n(red = Critical, orange = Warning)",
                      fontweight='bold')

    for ax in axes:
        ax.text(2, 196, "[SIMULATED]", color='white', fontsize=7,
                bbox=dict(boxstyle='round', fc='black', alpha=0.5))

    plt.suptitle("Layer 2 — Component Segmentation + Thermal Classification",
                 fontsize=11, fontweight='bold')
    plt.tight_layout()
    return _fig_to_bytes(fig), flagged


def _build_ndvi_exg_fig(R, G, B, NIR, NDVI, ExG, VH, VW, PIXEL_M) -> tuple:
    """Returns (figure_bytes, mean_canopy_density_pct).

    mean_canopy_density_pct is the corridor-wide mean of the per-cell canopy
    density grid shown in the bottom-right panel — computed here so it can
    also be passed to the Layer 3 prompt instead of being visible only in
    this chart.
    """
    NDVI_THRESH = 0.45
    GRID_PX = 10
    n_rows_grid = VH // GRID_PX
    n_cols_grid = VW // GRID_PX
    canopy_density = np.zeros((n_rows_grid, n_cols_grid), dtype=np.float32)
    for gr in range(n_rows_grid):
        for gc in range(n_cols_grid):
            cell = NDVI[gr*GRID_PX:(gr+1)*GRID_PX, gc*GRID_PX:(gc+1)*GRID_PX]
            canopy_density[gr, gc] = float((cell > NDVI_THRESH).mean())

    mean_canopy_density_pct = round(float(canopy_density.mean()) * 100, 1)

    rgb_drone = np.stack([R, G, B], axis=-1)
    fig, axes = plt.subplots(2, 2, figsize=(15, 9))

    axes[0,0].imshow(rgb_drone)
    axes[0,0].set_title("RGB Orthomosaic — 10 cm/pixel", fontweight='bold')
    axes[0,0].axhline(70,  color='white', linestyle='--', linewidth=1, alpha=0.7)
    axes[0,0].axhline(130, color='white', linestyle='--', linewidth=1, alpha=0.7)
    axes[0,0].text(230, 100, 'Clear strip', color='white', fontsize=9)
    axes[0,0].text(3, 193, "[SIMULATED DATA]", color='yellow', fontsize=7,
                   bbox=dict(boxstyle='round', fc='black', alpha=0.6))

    im_ndvi = axes[0,1].imshow(NDVI, cmap='RdYlGn', vmin=-0.2, vmax=0.95)
    axes[0,1].set_title("NDVI = (NIR − R) / (NIR + R)", fontweight='bold')
    plt.colorbar(im_ndvi, ax=axes[0,1], label='NDVI', fraction=0.046)
    axes[0,1].axhline(70,  color='white', linestyle='--', linewidth=1, alpha=0.7)
    axes[0,1].axhline(130, color='white', linestyle='--', linewidth=1, alpha=0.7)

    im_exg = axes[1,0].imshow(ExG, cmap='Greens', vmin=-0.15, vmax=0.55)
    axes[1,0].set_title("ExG = 2G − R − B  (RGB only — no NIR needed)", fontweight='bold')
    plt.colorbar(im_exg, ax=axes[1,0], label='ExG', fraction=0.046)
    axes[1,0].axhline(70,  color='black', linestyle='--', linewidth=1, alpha=0.6)
    axes[1,0].axhline(130, color='black', linestyle='--', linewidth=1, alpha=0.6)

    im_cd = axes[1,1].imshow(canopy_density, cmap='YlGn', vmin=0, vmax=1,
                              extent=[0, VW, VH, 0])
    axes[1,1].set_title(f"Canopy Density — fraction of 1 m² cells above NDVI {NDVI_THRESH}",
                        fontweight='bold')
    plt.colorbar(im_cd, ax=axes[1,1], label='Fraction (0–1)', fraction=0.046)
    axes[1,1].axhline(70,  color='white', linestyle='--', linewidth=1, alpha=0.7)
    axes[1,1].axhline(130, color='white', linestyle='--', linewidth=1, alpha=0.7)

    plt.suptitle("Layer 1 — Signal Processing: Drone Vegetation Orthomosaic\n"
                 "Catawba Valley Corridor — Simulated 10 cm/pixel",
                 fontsize=11, fontweight='bold')
    plt.tight_layout()
    return _fig_to_bytes(fig), mean_canopy_density_pct


def _build_watershed_fig(NDVI, VH, VW, PIXEL_M) -> tuple[bytes, list, list, int]:
    """Watershed segmentation. Returns (figure bytes, crown_props, encroaching, n_crowns)."""
    NDVI_THRESH = 0.45
    ndvi_smooth  = sk_gaussian(NDVI, sigma=3)
    canopy_mask  = (ndvi_smooth > NDVI_THRESH)
    distance     = distance_transform_edt(canopy_mask)
    local_max    = peak_local_max(distance, min_distance=12, labels=canopy_mask)

    markers = np.zeros((VH, VW), dtype=np.int32)
    for i, (r, c) in enumerate(local_max, start=1):
        markers[r, c] = i

    ws_labels   = watershed(-distance, markers, mask=canopy_mask)
    crown_props = regionprops(ws_labels, intensity_image=NDVI)
    n_crowns    = len(crown_props)

    encroaching = [p for p in crown_props
                   if (60 <= p.centroid[0] <= 80) or (120 <= p.centroid[0] <= 140)]

    rgb_drone = np.stack([NDVI, NDVI, NDVI], axis=-1)  # grey background for overlay
    fig, axes = plt.subplots(1, 3, figsize=(18, 6))

    axes[0].imshow(ndvi_smooth, cmap='RdYlGn', vmin=0, vmax=0.9)
    axes[0].plot(local_max[:, 1], local_max[:, 0], 'b.', markersize=5,
                 label=f'Crown seeds ({len(local_max)})')
    axes[0].axhline(70,  color='white', linestyle='--', linewidth=1)
    axes[0].axhline(130, color='white', linestyle='--', linewidth=1)
    axes[0].legend(loc='lower right', fontsize=8)
    axes[0].set_title(f"Smoothed NDVI + Crown Seeds\n({len(local_max)} local maxima detected)",
                      fontweight='bold')

    ws_display = ws_labels.copy().astype(float)
    ws_display[ws_labels == 0] = np.nan
    axes[1].imshow(NDVI, cmap='RdYlGn', vmin=-0.1, vmax=0.95)
    axes[1].imshow(ws_display, cmap='Set3', alpha=0.5)
    for p in encroaching:
        cy2, cx2 = p.centroid
        r2 = np.sqrt(p.area / np.pi)
        circ = plt.Circle((cx2, cy2), r2, color='red', fill=False, linewidth=2)
        axes[1].add_patch(circ)
    axes[1].axhline(70,  color='white', linestyle='--', linewidth=1.5)
    axes[1].axhline(130, color='white', linestyle='--', linewidth=1.5)
    axes[1].set_title(f"Watershed Crown Delineation\n{n_crowns} crowns | {len(encroaching)} encroaching (red circles)",
                      fontweight='bold')

    centroid_rows = [p.centroid[0] for p in crown_props]
    dist_from_strip = [
        abs(r - 70) * PIXEL_M if r < 100 else abs(r - 130) * PIXEL_M
        for r in centroid_rows
    ]
    axes[2].hist(dist_from_strip, bins=10, color='forestgreen', edgecolor='white', alpha=0.85)
    axes[2].axvline(1.5, color='red', linestyle='--', linewidth=1.5, label='1.5 m buffer')
    axes[2].set_xlabel("Distance from clear strip edge (m)")
    axes[2].set_ylabel("Number of crowns")
    axes[2].set_title("Crown Distribution by Distance\nfrom Clear Strip Edge", fontweight='bold')
    axes[2].legend()

    for ax in axes[:2]:
        ax.text(5, 192, "[SIMULATED]", color='white', fontsize=7,
                bbox=dict(boxstyle='round', fc='black', alpha=0.5))

    plt.suptitle("Layer 2 — Watershed Segmentation: Individual Crown Delineation",
                 fontsize=11, fontweight='bold')
    plt.tight_layout()
    return _fig_to_bytes(fig), crown_props, encroaching, n_crowns


def _build_rf_fig(R, G, B, NIR, NDVI, ExG, VH, VW, PIXEL_M, crown_params) -> tuple[bytes, list, float]:
    """Random Forest pixel classification. Returns (figure bytes, pred_counts, accuracy)."""
    def local_std(arr, size=5):
        lm = uniform_filter(arr.astype(np.float32), size=size)
        lq = uniform_filter(arr.astype(np.float32)**2, size=size)
        return np.sqrt(np.clip(lq - lm**2, 0, None))

    ndvi_tex = local_std(NDVI)
    nir_tex  = local_std(NIR)
    features = np.stack([R.ravel(), G.ravel(), B.ravel(), NIR.ravel(),
                         NDVI.ravel(), ExG.ravel(),
                         ndvi_tex.ravel(), nir_tex.ravel()], axis=1).astype(np.float32)
    feat_names = ['R', 'G', 'B', 'NIR', 'NDVI', 'ExG', 'NDVI_texture', 'NIR_texture']

    rows_g, cols_g = np.mgrid[0:VH, 0:VW]
    true_labels = np.zeros((VH, VW), dtype=np.int32)
    true_labels[70:130, :] = 1
    for (cy, cx, cr) in crown_params:
        d = np.sqrt((rows_g - cy)**2 + (cols_g - cx)**2)
        true_labels[d < cr * 0.5] = 3
        outer = (d >= cr * 0.5) & (d < cr * 0.9)
        true_labels[outer] = np.maximum(true_labels[outer], 2)

    true_flat = true_labels.ravel()
    rng2 = np.random.default_rng(seed=99)
    train_idx = []
    for c in range(4):
        ci = np.where(true_flat == c)[0]
        n  = max(50, int(len(ci) * 0.02))
        train_idx.extend(rng2.choice(ci, size=n, replace=False).tolist())
    train_idx = np.array(train_idx)

    clf = RandomForestClassifier(n_estimators=100, max_depth=12,
                                 min_samples_leaf=5, random_state=42, n_jobs=-1)
    clf.fit(features[train_idx], true_flat[train_idx])
    pred = clf.predict(features)
    classified = pred.reshape(VH, VW)
    accuracy = float((pred == true_flat).mean())
    pred_counts = [(pred == c).sum() for c in range(4)]
    importances = clf.feature_importances_

    class_names   = ['Bare soil', 'Low vegetation', 'Woody shrub', 'Tree canopy']
    class_colours = np.array([[0.76,0.60,0.42,1],[0.75,0.90,0.35,1],
                               [0.40,0.65,0.20,1],[0.10,0.40,0.10,1]])
    class_cmap    = ListedColormap(class_colours)

    fig, axes = plt.subplots(1, 3, figsize=(18, 6))

    axes[0].imshow(true_labels, cmap=class_cmap, vmin=-0.5, vmax=3.5)
    axes[0].set_title("Ground Truth Labels\n(from simulation parameters)", fontweight='bold')
    axes[0].axhline(70,  color='white', linestyle='--', linewidth=1.2)
    axes[0].axhline(130, color='white', linestyle='--', linewidth=1.2)

    axes[1].imshow(classified, cmap=class_cmap, vmin=-0.5, vmax=3.5)
    axes[1].set_title(f"Random Forest Classification\n({accuracy*100:.1f}% pixel accuracy)",
                      fontweight='bold')
    axes[1].axhline(70,  color='white', linestyle='--', linewidth=1.2)
    axes[1].axhline(130, color='white', linestyle='--', linewidth=1.2)
    patches = [mpatches.Patch(color=class_colours[i], label=class_names[i]) for i in range(4)]
    axes[1].legend(handles=patches, loc='lower right', fontsize=8)

    sorted_idx = np.argsort(importances)[::-1]
    axes[2].barh([feat_names[i] for i in sorted_idx], importances[sorted_idx],
                 color='steelblue', edgecolor='white')
    axes[2].set_xlabel("Importance")
    axes[2].set_title("Feature Importances\n(Random Forest Gini)", fontweight='bold')
    axes[2].invert_yaxis()

    for ax in axes[:2]:
        ax.text(5, 192, "[SIMULATED]", color='white', fontsize=7,
                bbox=dict(boxstyle='round', fc='black', alpha=0.5))

    plt.suptitle("Layer 2 — Random Forest Pixel Classification", fontsize=11, fontweight='bold')
    plt.tight_layout()
    return _fig_to_bytes(fig), pred_counts, accuracy


def _build_comparison_fig(n_drone_crowns, n_encroaching, PIXEL_M, VH, VW) -> tuple[bytes, dict]:
    """Bar chart comparing drone watershed vs Arc 5 LiDAR. Returns (figure bytes, lidar_stats)."""
    lidar_path = ASSETS_DIR / "lidar_catawba_nc.npz"
    if not lidar_path.exists():
        return None, {}

    lidar_data    = np.load(str(lidar_path), allow_pickle=False)
    lidar_n_crowns   = int(lidar_data['n_trees'])
    lidar_n_viol     = int(lidar_data['n_violating'])
    lidar_mean_h     = float(lidar_data['mean_height_m'])
    lidar_viol_pct   = float(lidar_data['violation_pct'])
    lidar_area_m2    = 300 * 150
    drone_area_m2    = VH * VW * PIXEL_M**2
    lidar_density    = lidar_n_crowns / lidar_area_m2 * 1000
    drone_density    = n_drone_crowns  / drone_area_m2 * 1000

    lidar_stats = {
        "n_crowns": lidar_n_crowns, "n_violating": lidar_n_viol,
        "mean_h": lidar_mean_h, "viol_pct": lidar_viol_pct,
        "area_m2": lidar_area_m2, "density": lidar_density,
        "drone_area_m2": drone_area_m2, "drone_density": drone_density,
    }

    metrics    = ['Crown density\n(per 1,000 m²)', 'Encroaching\ncrowns (count)']
    lidar_vals = [lidar_density, lidar_n_viol]
    drone_vals = [drone_density, n_encroaching]
    x     = np.arange(len(metrics))
    width = 0.35

    fig, ax = plt.subplots(figsize=(8, 5))
    b1 = ax.bar(x - width/2, lidar_vals, width, label='Arc 5 LiDAR DBSCAN',
                color='steelblue', edgecolor='white')
    b2 = ax.bar(x + width/2, drone_vals, width, label='Arc 6 Drone Watershed',
                color='forestgreen', edgecolor='white')
    ax.set_xticks(x)
    ax.set_xticklabels(metrics, fontsize=10)
    ax.legend(fontsize=9)
    ax.set_title("Arc 5 LiDAR DBSCAN vs Arc 6 Drone Watershed\n"
                 "Catawba Valley Corridor — Crown Detection Comparison",
                 fontweight='bold')
    ax.set_facecolor('#f8f8f8')
    fig.patch.set_facecolor('white')
    ax.grid(axis='y', alpha=0.3)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    for bar in b1:
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05,
                f'{bar.get_height():.1f}', ha='center', va='bottom', fontsize=9)
    for bar in b2:
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05,
                f'{bar.get_height():.1f}', ha='center', va='bottom', fontsize=9)
    plt.tight_layout()
    return _fig_to_bytes(fig), lidar_stats


# ---------------------------------------------------------------------------
# AI brief builders
# ---------------------------------------------------------------------------

def _inspection_prompt(zone_stats: dict, flagged_segs: list, zone_anomaly: dict = None) -> str:
    zone_anomaly = zone_anomaly or {}
    lines = [
        f"{n}: {s['mean_C']}°C (delta-T {s['delta_T']}°C) — {s['class']} — "
        f"Isolation Forest anomalous pixels: {zone_anomaly.get(n, 'not computed')}%"
        for n, s in zone_stats.items()
    ]
    return f"""You are an Earth observation and thermal inspection analyst producing a structured UAV inspection brief for a utility engineering team.

TOWER: 115 kV transmission tower, Catawba Valley NC corridor. Ambient air temperature: {AMBIENT_TEMP}°C.
[Data is simulated — modelled on real UAV thermal inspection datasets]

LAYER 2 OUTPUTS (computed from co-registered RGB and thermal rasters):
{chr(10).join(lines)}

Flagged segments (Warning or Critical): {len(flagged_segs)}

Write a structured five-element inspection brief. Use exactly these five headings.
Each section must be a minimum of 60 words (minimum 300 words total across all five sections),
except sections that legitimately have nothing to report (e.g. zero Critical findings) —
in that case state so explicitly in 1-2 sentences rather than padding.

## 1. Critical Findings
List every component classified Critical (delta-T > 30°C). State the component name, mean temperature, delta-T, and the specific failure mode most consistent with that thermal signature (e.g. cracked ceramic disc, contaminated surface, loose compression joint). State the required action and urgency for each.

## 2. Warnings
List every component classified Warning (delta-T 15–30°C). State the component name, delta-T, and the most likely cause. State the recommended inspection timeframe.

## 3. Normal Components
Briefly confirm which components are within normal range and why no action is required.

## 4. Sensor Capability and Methodology
Explain what the co-registered thermal + RGB approach provides that a single-sensor inspection cannot. Cover: why delta-T is used instead of absolute temperature, how Isolation Forest detects anomalies without labelled training data, and how SLIC segmentation assigns anomalous pixels to specific components. Reference the actual per-component anomalous-pixel percentages listed above — do not describe Isolation Forest only in the abstract.

## 5. Field Verification Requirements
State clearly what this thermal analysis cannot confirm without field inspection. Cover: internal component condition vs surface temperature, conductor sag under load, and thermal camera resolution limits relative to the RGB camera.

End with:
DATA QUALITY: Component zones: {len(zone_stats)} | Flagged segments: {len(flagged_segs)} | Source: simulated data modelled on real UAV inspection datasets"""


def _inspection_fallback(zone_stats: dict, flagged_segs: list) -> str:
    critical = [n for n, s in zone_stats.items() if s['class'] == 'CRITICAL']
    warnings = [n for n, s in zone_stats.items() if s['class'] == 'Warning']
    critical_str = ", ".join(critical) if critical else "None"
    warning_str  = ", ".join(warnings) if warnings else "None"
    return f"""## Critical Findings
**{critical_str}** — delta-T exceeds 30°C above ambient. Consistent with cracked ceramic disc or \
partial flashover development. Remove from service within 24–48 hours. Schedule live-line insulator replacement.

## Warnings
**{warning_str}** — delta-T in the 15–30°C warning band. Likely early-stage surface contamination \
reducing insulation resistance. Schedule field inspection within 7 days.

## Normal Components
All remaining components within normal range (delta-T under 10°C). No action required.

## Sensor Note
Thermal camera resolution is 3–5× lower than the co-registered RGB. Anomaly boundaries \
may extend 1–2 pixels beyond the actual fault zone. Field verification required before work order issuance.

DATA QUALITY: Component zones: {len(zone_stats)} | Flagged segments: {len(flagged_segs)} | \
Source: simulated data modelled on real UAV inspection datasets"""


def _fmt_lidar_stat(lidar_stats: dict, key: str, suffix: str = "", decimals: int = 1) -> str:
    """Format a numeric value from lidar_stats safely.

    lidar_stats is {} whenever the Arc 5 LiDAR comparison asset file is
    missing (see _build_comparison_fig). Piping a missing value straight into
    a numeric format spec, e.g. f"{lidar_stats.get('viol_pct', '—'):.1f}",
    crashes with ValueError the instant the placeholder string '—' hits
    ':.1f'. This returns plain text instead, so a missing asset degrades
    gracefully rather than crashing the whole brief.
    """
    value = lidar_stats.get(key)
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        return f"{value:.{decimals}f}{suffix}"
    return "not available"


def _vegetation_prompt(n_crowns, n_encroaching, pred_counts, PIXEL_M, VH, VW,
                       lidar_stats: dict, canopy_density_pct: float = None) -> str:
    class_names = ['Bare soil', 'Low vegetation', 'Woody shrub', 'Tree canopy']
    n_total = VH * VW
    class_lines = "\n".join(
        f"  {class_names[c]}: {pred_counts[c]:,} px = {pred_counts[c]*PIXEL_M**2:.0f} m²"
        for c in range(4)
    )
    canopy_density_line = (
        f"  Mean canopy density (Layer 1, 1 m² grid): {canopy_density_pct}% of cells above the NDVI threshold"
        if canopy_density_pct is not None
        else "  Mean canopy density: not computed for this run"
    )
    return f"""You are an Earth observation and vegetation management analyst producing a structured drone survey brief for a utility corridor management team.

CORRIDOR: Catawba Valley 115kV transmission corridor — same area as Arc 5 LiDAR survey.
[Data is simulated — modelled on real UAV corridor survey datasets]

LAYER 2 OUTPUTS — DRONE (watershed segmentation + Random Forest, 10 cm/pixel):
  Survey area: {VH*VW*PIXEL_M**2:.0f} m²
  Watershed crowns detected: {n_crowns}
  Encroaching crowns (within 1 m of clear strip boundary): {n_encroaching}
{canopy_density_line}
  Random Forest pixel classification:
{class_lines}

LAYER 2 OUTPUTS — ARC 5 LIDAR (DBSCAN, same corridor, full extent):
  Coverage: {lidar_stats.get('area_m2', 45000):,} m²
  DBSCAN crowns detected: {lidar_stats.get('n_crowns', 102)}
  Violating crowns: {lidar_stats.get('n_violating', 'not available')}
  Violation rate: {_fmt_lidar_stat(lidar_stats, 'viol_pct', '%')}
  Mean tree height: {_fmt_lidar_stat(lidar_stats, 'mean_h', ' m')}
  LiDAR crown density: {_fmt_lidar_stat(lidar_stats, 'density', ' per 1,000 m²')}
  Drone crown density: {_fmt_lidar_stat(lidar_stats, 'drone_density', ' per 1,000 m²')}

Write a structured five-element vegetation management brief. Use exactly these five headings.
Each section must be a minimum of 60 words (minimum 300 words total across all five sections).

## 1. Encroachment Status
State the number of encroaching crowns and their proximity to the clear strip. Assess the immediate vegetation management risk. Reference the drone crown count and the LiDAR violation rate from the same corridor.

## 2. Land Cover Breakdown
Interpret the Random Forest classification results together with the mean canopy density figure given above. State what each class means for corridor management. Identify which land cover types present a risk and which do not.

## 3. Drone vs LiDAR Comparison
Explain where the two methods agree and where they diverge. Specifically address why the drone finds higher crown density than LiDAR DBSCAN (different algorithms, not different trees). State what each method measures that the other cannot — height from LiDAR, spectral class and crown texture from drone.

## 4. Algorithm Capability
Explain what watershed segmentation and Random Forest provide that manual photo interpretation cannot. Cover: how watershed finds crown boundaries without knowing the number of crowns in advance, and how texture features in Random Forest separate tree canopy from smooth grass at 10 cm resolution.

## 5. Recommended Operational Approach
State how to combine drone and LiDAR in a production vegetation management workflow. Specify which method is appropriate for corridor-wide screening vs targeted follow-up inspection. Include at least one field verification step neither method can replace.

End with:
DATA QUALITY: Drone crowns: {n_crowns} | Encroaching: {n_encroaching} | LiDAR crowns: {lidar_stats.get('n_crowns', 102)} | Source: simulated data modelled on real corridor survey datasets"""


def _vegetation_fallback(n_crowns, n_encroaching, lidar_stats: dict) -> str:
    return f"""## Drone Findings
Watershed segmentation identified **{n_crowns} individual crowns** in the 30 m × 20 m survey patch. \
**{n_encroaching} crowns** were detected within 1 m of the clear strip boundary — these represent \
active encroachment risk. The Random Forest classifier confirms tree canopy and woody shrub dominate \
both sides of the clear strip, with bare soil confined to the cleared zone.

## Comparison with Arc 5 LiDAR
LiDAR DBSCAN found {lidar_stats.get('n_crowns', 102)} crowns in the full \
{lidar_stats.get('area_m2', 45000):,} m² corridor \
(density: {_fmt_lidar_stat(lidar_stats, 'density', ' per 1,000 m²')}). Drone watershed found \
{_fmt_lidar_stat(lidar_stats, 'drone_density', ' per 1,000 m²')}. The higher drone density is expected — \
watershed separates touching crowns at 10 cm resolution that DBSCAN merges in 3D point density space.

## Where They Agree
Both methods confirm the same pattern: dense canopy on both sides of the clear strip with active \
encroachment. The LiDAR violation rate ({_fmt_lidar_stat(lidar_stats, 'viol_pct', '%')}) is consistent with \
the drone encroachment count.

## Where They Diverge
Drone imagery cannot measure tree height. LiDAR gives height to within 0.2 m — essential for \
NERC FAC-003 clearance compliance. For violation determination you need LiDAR. For precise crown \
mapping and species identification you need drone imagery.

## Operational Approach
Use LiDAR for quarterly corridor-wide screening. Use drone for targeted follow-up on priority zones \
LiDAR flags. The two are complementary, not competing.

DATA QUALITY: Drone crowns: {n_crowns} | LiDAR crowns: {lidar_stats.get('n_crowns','—')} | \
Source: simulated data modelled on real corridor survey datasets"""


# ---------------------------------------------------------------------------
# Word document builders
# ---------------------------------------------------------------------------

def _add_bold_runs(paragraph, text: str):
    for part in re.split(r'(\*\*.*?\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def _build_word_inspection(zone_stats, brief_text, model_used,
                            rgb_thermal_bytes, iso_bytes, seg_bytes) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(1)
    sec.left_margin = sec.right_margin = Inches(1.1)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    t = doc.add_paragraph()
    t.add_run('Drone Infrastructure Inspection Brief').bold = True
    t.runs[0].font.size = Pt(16)
    doc.add_paragraph("Corridor: Catawba Valley NC — 115kV transmission tower")
    doc.add_paragraph("[SIMULATED DATA — modelled on real UAV inspection datasets]").italic = True
    doc.add_paragraph()

    # Layer 1 figure
    p = doc.add_paragraph()
    p.add_run('Layer 1 — Co-registered RGB and Thermal Rasters').bold = True
    p.runs[0].font.size = Pt(13)
    if rgb_thermal_bytes:
        doc.add_picture(io.BytesIO(rgb_thermal_bytes), width=Inches(6.0))
        c = doc.add_paragraph()
        c.add_run('Left: RGB orthomosaic at 5 cm/pixel. Right: thermal raster (°C). '
                  'Both images are pixel-for-pixel aligned (co-registered).').italic = True
        c.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # Layer 2 delta-T table
    p2 = doc.add_paragraph()
    p2.add_run('Layer 2 — Component Delta-T and CIGRE Classification').bold = True
    p2.runs[0].font.size = Pt(13)
    tbl = doc.add_table(rows=len(zone_stats) + 1, cols=4)
    tbl.style = 'Table Grid'
    for j, h in enumerate(['Component', 'Mean Temp (°C)', 'Delta-T (°C)', 'CIGRE Class']):
        tbl.rows[0].cells[j].text = h
        tbl.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, (name, s) in enumerate(zone_stats.items(), start=1):
        tbl.rows[i].cells[0].text = name
        tbl.rows[i].cells[1].text = str(s['mean_C'])
        tbl.rows[i].cells[2].text = str(s['delta_T'])
        tbl.rows[i].cells[3].text = s['class']
    doc.add_paragraph()

    # Isolation Forest figure
    p3 = doc.add_paragraph()
    p3.add_run('Layer 2 — Isolation Forest: Unsupervised Thermal Anomaly Detection').bold = True
    p3.runs[0].font.size = Pt(13)
    doc.add_paragraph(
        'Isolation Forest scores each pixel by how few random cuts are needed to isolate it. '
        'Anomalous pixels (rare, extreme temperatures) are isolated quickly. '
        'Normal pixels, which are densely clustered near ambient, require many cuts. '
        'No labelled training data is needed — the algorithm learns "normal" from the data itself.'
    )
    if iso_bytes:
        doc.add_picture(io.BytesIO(iso_bytes), width=Inches(6.0))
        c2 = doc.add_paragraph()
        c2.add_run('Left: raw thermal. Centre: continuous anomaly score (brighter = more anomalous). '
                   'Right: binary anomaly flags overlaid on thermal (cyan = flagged pixel).').italic = True
        c2.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # Segmentation figure
    p4 = doc.add_paragraph()
    p4.add_run('Layer 2 — Component Segmentation + Thermal Classification').bold = True
    p4.runs[0].font.size = Pt(13)
    doc.add_paragraph(
        'SLIC superpixel segmentation (SAM proxy) groups pixels into compact regions based on '
        'colour similarity and spatial proximity. Each segment is then assigned the mean thermal '
        'value from the co-registered thermal raster. This connects the two sensors: RGB defines '
        'the component boundaries; thermal defines the temperature of each component. '
        'CIGRE delta-T thresholds are then applied to classify each segment.'
    )
    if seg_bytes:
        doc.add_picture(io.BytesIO(seg_bytes), width=Inches(6.0))
        c3 = doc.add_paragraph()
        c3.add_run('Left: segment boundaries on RGB. Centre: severity colour map. '
                   'Right: flagged segments on thermal (red = Critical, orange = Warning).').italic = True
        c3.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # Algorithm connection note
    p5 = doc.add_paragraph()
    p5.add_run('How the algorithms connect').bold = True
    doc.add_paragraph(
        'Isolation Forest and segmentation are independent algorithms applied to the same thermal raster. '
        'Isolation Forest works at the pixel level — it flags individual pixels that are statistically '
        'anomalous. Segmentation works at the region level — it groups pixels into component-sized regions '
        'and classifies each region as a whole. Together they provide two complementary views: '
        'Isolation Forest finds WHERE the hot pixels are; segmentation tells you WHICH component they belong to.'
    )
    doc.add_paragraph()

    # AI brief
    p6 = doc.add_paragraph()
    p6.add_run('Layer 3 — AI Inspection Brief').bold = True
    p6.runs[0].font.size = Pt(13)
    if model_used:
        doc.add_paragraph(f'Generated by: {model_used}').italic = True

    for line in brief_text.splitlines():
        s = line.strip()
        if s.startswith('## '):
            doc.add_paragraph(s[3:], style='Heading 2')
        elif s.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            _add_bold_runs(p, s[2:])
        elif s.startswith('DATA QUALITY:'):
            doc.add_paragraph().add_run(s).bold = True
        elif s == '' or s.startswith('---'):
            doc.add_paragraph()
        elif s:
            p = doc.add_paragraph()
            _add_bold_runs(p, s)

    doc.add_paragraph()
    doc.add_paragraph('Generated by EOIL — AI-Native Earth Observation Innovation Lab').italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _build_word_vegetation(pred_counts, n_crowns, n_encroaching, accuracy,
                            brief_text, model_used, PIXEL_M, VH, VW, lidar_stats,
                            ndvi_bytes, watershed_bytes, rf_bytes, comparison_bytes) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(1)
    sec.left_margin = sec.right_margin = Inches(1.1)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    t = doc.add_paragraph()
    t.add_run('Drone Vegetation Mapping Brief').bold = True
    t.runs[0].font.size = Pt(16)
    doc.add_paragraph("Corridor: Catawba Valley NC — 115kV transmission corridor (Arc 5 area)")
    doc.add_paragraph("[SIMULATED DATA — modelled on real corridor survey datasets]").italic = True
    doc.add_paragraph()

    # Layer 1 figure
    p = doc.add_paragraph()
    p.add_run('Layer 1 — Signal Processing: NDVI, ExG, Canopy Density').bold = True
    p.runs[0].font.size = Pt(13)
    doc.add_paragraph(
        'NDVI (Normalized Difference Vegetation Index) = (NIR − R) / (NIR + R). '
        'Requires a multispectral camera. Range −1 to +1; healthy canopy typically 0.7–0.9. '
        'ExG (Excess Green) = 2G − R − B. Works with a standard RGB camera, no NIR needed. '
        'Both indices are computed at 10 cm/pixel — 100× finer than the Sentinel-2 satellite data '
        'used in Arc 1 and Arc 2.'
    )
    if ndvi_bytes:
        doc.add_picture(io.BytesIO(ndvi_bytes), width=Inches(6.0))
        c = doc.add_paragraph()
        c.add_run('Top left: RGB. Top right: NDVI. Bottom left: ExG. Bottom right: canopy density grid (1 m² cells). '
                  'Dashed white lines mark the clear strip boundaries.').italic = True
        c.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # Watershed figure
    p2 = doc.add_paragraph()
    p2.add_run('Layer 2 — Watershed Segmentation: Individual Crown Delineation').bold = True
    p2.runs[0].font.size = Pt(13)
    doc.add_paragraph(
        'Watershed treats the NDVI surface as a topographic landscape. '
        'Crown centres are local NDVI peaks (high ground). Gaps between crowns are valleys. '
        'The algorithm floods outward from each peak, with flood boundaries becoming crown edges. '
        'A distance transform is applied first to prevent over-segmentation in touching crowns. '
        'The result is one labelled region per individual plant crown, regardless of whether crowns touch.'
    )
    if watershed_bytes:
        doc.add_picture(io.BytesIO(watershed_bytes), width=Inches(6.0))
        c2 = doc.add_paragraph()
        c2.add_run(f'Left: smoothed NDVI with crown seeds (blue dots). '
                   f'Centre: {n_crowns} watershed crowns overlaid on NDVI — red circles mark encroaching crowns. '
                   f'Right: crown count by distance from clear strip edge.').italic = True
        c2.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # RF figure
    p3 = doc.add_paragraph()
    p3.add_run('Layer 2 — Random Forest Pixel Classification').bold = True
    p3.runs[0].font.size = Pt(13)
    doc.add_paragraph(
        'The same Random Forest algorithm used in Arc 5 LiDAR is applied here, '
        'but the features are different: spectral bands (R, G, B, NIR, NDVI, ExG) '
        'and texture (local standard deviation in a 5×5 window). '
        'Texture captures the visual roughness of tree canopy vs smooth grass — '
        'a feature that spectral values alone cannot distinguish at this resolution. '
        f'Classification accuracy vs simulation ground truth: {accuracy*100:.1f}%.'
    )
    if rf_bytes:
        doc.add_picture(io.BytesIO(rf_bytes), width=Inches(6.0))
        c3 = doc.add_paragraph()
        c3.add_run('Left: ground truth from simulation. Centre: Random Forest prediction. '
                   'Right: feature importances — NDVI and NIR dominate.').italic = True
        c3.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # How algorithms connect
    p_conn = doc.add_paragraph()
    p_conn.add_run('How the algorithms connect').bold = True
    doc.add_paragraph(
        'Watershed and Random Forest are applied to the same image but answer different questions. '
        'Watershed asks: where are the crown boundaries? It produces spatial regions. '
        'Random Forest asks: what land cover type is each pixel? It produces a class label. '
        'Together: watershed tells you how many crowns there are and where they are; '
        'Random Forest tells you whether those crowns are tree canopy, shrub, or something else. '
        'In production, you would apply Random Forest first to filter out non-canopy pixels, '
        'then run watershed only on confirmed canopy pixels — reducing false crown detections.'
    )
    doc.add_paragraph()

    # Stats table
    p4 = doc.add_paragraph()
    p4.add_run('Layer 2 — Classification Summary').bold = True
    p4.runs[0].font.size = Pt(13)
    class_names = ['Bare soil', 'Low vegetation', 'Woody shrub', 'Tree canopy']
    tbl = doc.add_table(rows=len(class_names) + 1, cols=3)
    tbl.style = 'Table Grid'
    for j, h in enumerate(['Class', 'Pixels', 'Area (m²)']):
        tbl.rows[0].cells[j].text = h
        tbl.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, (name, cnt) in enumerate(zip(class_names, pred_counts), start=1):
        tbl.rows[i].cells[0].text = name
        tbl.rows[i].cells[1].text = f"{cnt:,}"
        tbl.rows[i].cells[2].text = f"{cnt * PIXEL_M**2:.1f}"
    doc.add_paragraph()

    # Comparison figure
    p5 = doc.add_paragraph()
    p5.add_run('Arc 5 LiDAR DBSCAN vs Arc 6 Drone Watershed — Comparison').bold = True
    p5.runs[0].font.size = Pt(13)
    doc.add_paragraph(
        'DBSCAN (used in Arc 5) clusters 3D point density. Crowns that touch in 3D space merge '
        'into one cluster. Watershed (used here) finds 2D NDVI local maxima. Adjacent crowns '
        'produce separate peaks and are delineated individually. This explains the higher '
        'crown density from the drone — the algorithm sees more crowns, not more trees. '
        'Neither result is wrong. They answer different questions at different scales.'
    )
    if comparison_bytes:
        doc.add_picture(io.BytesIO(comparison_bytes), width=Inches(5.5))
    doc.add_paragraph()

    # AI brief
    p6 = doc.add_paragraph()
    p6.add_run('Layer 3 — AI Vegetation Management Brief').bold = True
    p6.runs[0].font.size = Pt(13)
    if model_used:
        doc.add_paragraph(f'Generated by: {model_used}').italic = True

    for line in brief_text.splitlines():
        s = line.strip()
        if s.startswith('## '):
            doc.add_paragraph(s[3:], style='Heading 2')
        elif s.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            _add_bold_runs(p, s[2:])
        elif s.startswith('DATA QUALITY:'):
            doc.add_paragraph().add_run(s).bold = True
        elif s == '' or s.startswith('---'):
            doc.add_paragraph()
        elif s:
            p = doc.add_paragraph()
            _add_bold_runs(p, s)

    doc.add_paragraph()
    doc.add_paragraph('Generated by EOIL — AI-Native Earth Observation Innovation Lab').italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Main render function
# ---------------------------------------------------------------------------

def render():
    st.header("🚁 Drone Intelligence")
    st.markdown(
        "Two sub-modules. One portal page. "
        "Drone sensors operate at centimetre resolution — a scale no satellite can match. "
        "Sub-module A uses a drone-mounted thermal camera to detect electrical faults on a transmission tower before failure. "
        "Sub-module B uses a multispectral drone to map individual tree crowns at 10 cm resolution over the same corridor as Arc 5 LiDAR, "
        "then compares the two methods directly."
    )
    st.info("**All data is simulated.** Modelled on real UAV inspection and corridor survey datasets. "
            "Labelled throughout.", icon="ℹ️")

    tab_a, tab_b = st.tabs(["🔥 Sub-module A — Infrastructure Inspection",
                             "🌿 Sub-module B — Vegetation Mapping"])

    # ==========================================================================
    # SUB-MODULE A — INFRASTRUCTURE INSPECTION
    # ==========================================================================
    with tab_a:
        st.subheader("Infrastructure Inspection — Thermal + RGB Drone Survey")
        st.markdown(
            "A drone flies a 115 kV transmission tower with two sensors mounted simultaneously: "
            "an RGB camera and a thermal infrared camera. Both capture the same ground at 5 cm/pixel. "
            "The thermal image reveals component temperatures — electrical faults generate heat "
            "before they cause failure. Finding a hot insulator costs a repair crew a few hours. "
            "Missing it can cause a flashover and outage affecting thousands of customers."
        )

        # --- Run button ---
        if st.button("▶ Run Inspection Analysis", type="primary", key="run_inspection"):
            st.session_state["inspection_pending"] = True

        if st.session_state.get("inspection_pending"):
            st.session_state["inspection_pending"] = False

            with st.spinner("Simulating tower scene and running algorithms..."):
                rgb, thermal, H, W = _simulate_tower_scene()
                zone_stats = _compute_zone_stats(thermal)
                rgb_thermal_bytes = _build_rgb_thermal_fig(rgb, thermal)
                iso_bytes, anomaly_map, zone_anomaly = _build_isolation_forest_fig(thermal, H, W)
                seg_bytes, flagged_segs = _build_segmentation_fig(rgb, thermal, H, W)

            st.session_state["insp_zone_stats"]       = zone_stats
            st.session_state["insp_rgb_thermal_bytes"] = rgb_thermal_bytes
            st.session_state["insp_iso_bytes"]         = iso_bytes
            st.session_state["insp_zone_anomaly"]      = zone_anomaly
            st.session_state["insp_seg_bytes"]         = seg_bytes
            st.session_state["insp_flagged_segs"]      = flagged_segs
            st.session_state["insp_done"]              = True
            st.session_state.pop("insp_brief", None)

        if st.session_state.get("insp_done"):
            zone_stats    = st.session_state["insp_zone_stats"]
            rgb_th_bytes  = st.session_state["insp_rgb_thermal_bytes"]
            iso_bytes     = st.session_state["insp_iso_bytes"]
            zone_anomaly  = st.session_state["insp_zone_anomaly"]
            seg_bytes     = st.session_state["insp_seg_bytes"]
            flagged_segs  = st.session_state["insp_flagged_segs"]

            # ── Layer 1 ──────────────────────────────────────────────────────
            st.divider()
            st.markdown("### Layer 1 — Signal Processing")
            with st.expander("What is Layer 1 doing?", expanded=False):
                st.markdown("""
**Layer 1 is engineering, not AI.** It prepares the raw sensor data for the algorithms.

Two rasters are generated at 5 cm/pixel over a 10 m × 10 m scene:
- **RGB orthomosaic** — visible-light imagery. Shows what the tower looks like. Used for component segmentation.
- **Thermal raster** — surface temperature in °C. Used for anomaly detection and delta-T calculation.

**Co-registration** means both images are pixel-for-pixel aligned. Pixel (row=50, col=100) in the RGB
shows the exact same 5×5 cm patch as pixel (50, 100) in the thermal. Without co-registration, you cannot
assign a temperature to a specific component — the images would be misaligned.

**Delta-T** = component temperature − ambient baseline (24°C). Used in CIGRE thresholds below.
A component at 58°C in 24°C air has delta-T = 34°C — that is the fault indicator, not the absolute temperature.
""")

            st.image(rgb_th_bytes, caption="Layer 1 — Co-registered RGB and Thermal", use_container_width=True)

            # Delta-T table
            st.markdown("**Component delta-T table (CIGRE thresholds):**")
            import pandas as pd
            df_zones = pd.DataFrame([
                {"Component": n, "Mean Temp (°C)": s["mean_C"],
                 "Delta-T (°C)": s["delta_T"], "CIGRE Class": s["class"]}
                for n, s in zone_stats.items()
            ])
            def _colour_class(val):
                c = {"Normal": "background-color:#d4edda",
                     "Monitor": "background-color:#fff3cd",
                     "Warning": "background-color:#ffe5b4",
                     "CRITICAL": "background-color:#f8d7da; font-weight:bold"}.get(val, "")
                return c
            st.dataframe(
                df_zones.style.map(_colour_class, subset=["CIGRE Class"]),
                use_container_width=True, hide_index=True
            )

            with st.expander("CIGRE delta-T thresholds — what do they mean?", expanded=False):
                st.markdown("""
CIGRE is the international power engineering body. Their guidance on thermal inspection thresholds:

| Delta-T | Class | Required action |
|---------|-------|----------------|
| < 5°C | Normal | No action |
| 5–15°C | Monitor | Inspect within 30 days |
| 15–30°C | Warning | Inspect within 7 days |
| > 30°C | **Critical** | Remove from service / inspect immediately |

Delta-T is used instead of absolute temperature because ambient conditions change.
A conductor at 42°C on a cold day (ambient 10°C, delta-T 32°C) is more alarming than
the same conductor on a hot day (ambient 38°C, delta-T 4°C).
""")

            # ── Layer 2: Isolation Forest ─────────────────────────────────────
            st.divider()
            st.markdown("### Layer 2 — Algorithm 1: Isolation Forest")
            with st.expander("What is Isolation Forest and why was it chosen?", expanded=False):
                st.markdown("""
**Isolation Forest** is an *unsupervised* anomaly detection algorithm.
Unsupervised means it does not need labelled training data — you do not tell it which pixels are faulty.
It learns what "normal" looks like from the data itself.

**How it works:** Build 100 random decision trees. For each data point, count how many
random cuts are needed to isolate it into its own branch.
- **Normal pixels** are dense (many similar temperatures nearby). Many cuts needed to isolate them.
- **Anomalous pixels** are rare and different. Very few cuts needed — they are already alone.

The algorithm returns an anomaly score per pixel. Pixels with the lowest scores (easiest to isolate)
are flagged as anomalous.

**Why not simple thresholding?** A threshold like "flag anything above 40°C" requires knowing
the threshold in advance. On a hot day, normal components run warmer. Isolation Forest adapts
to the data — it always flags the statistically unusual pixels relative to the full scene.

**The contamination parameter** = 0.03 means we expect 3% of pixels to be anomalous.
In a real inspection, most of a tower is healthy — faults are rare. Setting this too high
creates too many false alarms; too low misses real faults.

**Connection to segmentation (Algorithm 2):** Isolation Forest finds WHERE anomalous pixels are.
It does not know what component they belong to. Segmentation (next step) defines the component boundaries
so you can answer: *this anomalous cluster is inside the top-right insulator.*
""")

            st.image(iso_bytes, caption="Layer 2 — Isolation Forest anomaly detection", use_container_width=True)

            # Per-zone anomaly coverage
            st.markdown("**Anomaly flag coverage per component zone:**")
            df_anom = pd.DataFrame([
                {"Component": n, "Anomalous pixels (%)": v,
                 "CIGRE Class": zone_stats[n]["class"]}
                for n, v in zone_anomaly.items()
            ])
            st.dataframe(df_anom.style.map(_colour_class, subset=["CIGRE Class"]),
                         use_container_width=True, hide_index=True)

            # ── Layer 2: Segmentation ─────────────────────────────────────────
            st.divider()
            st.markdown("### Layer 2 — Algorithm 2: Component Segmentation (SAM proxy)")
            with st.expander("What is segmentation and how does it connect to Isolation Forest?", expanded=False):
                st.markdown("""
**Segmentation** divides the image into coherent regions — in this case, tower components.

**SAM (Segment Anything Model)** from Meta AI (2023) can segment any image into precise
component-level regions without domain-specific training. It would produce precise pixel-level
outlines for each insulator, crossarm section, and conductor span. Here we use **SLIC superpixels**
as a proxy — a simpler 2012 algorithm that groups pixels by colour and spatial proximity.
The result is less precise than SAM but sufficient to demonstrate the concept.

**The connection between the two algorithms:**
1. **Isolation Forest** (pixel-level) → flags individual anomalous pixels in the thermal raster
2. **Segmentation** (region-level) → defines which pixels belong to each component
3. **Cross-reference** → for each segment, compute mean thermal value and delta-T → apply CIGRE threshold

This is how you go from "pixel at row 43, column 149 is anomalous" to
"the top-right insulator is Critical." The two algorithms together answer the full question.

**Layer 3 receives:** a structured list of components with their mean temperature, delta-T,
and severity class — not raw pixels. Generative AI interprets structure, not images directly.
""")

            st.image(seg_bytes, caption="Layer 2 — Component segmentation and thermal classification",
                     use_container_width=True)

            if flagged_segs:
                st.markdown("**Flagged segments (Warning or Critical):**")
                df_flag = pd.DataFrame(flagged_segs)[["class", "mean_C", "delta_T"]]
                df_flag.columns = ["Class", "Mean Temp (°C)", "Delta-T (°C)"]
                st.dataframe(df_flag.style.map(_colour_class, subset=["Class"]),
                             use_container_width=True, hide_index=True)

            # ── Layer 3: AI brief ─────────────────────────────────────────────
            st.divider()
            st.markdown("### Layer 3 — Generative AI: Inspection Brief")
            with st.expander("What does Layer 3 do and why is it not a substitute for Layers 1 and 2?", expanded=False):
                st.markdown("""
**Layer 3 is interpretation, not analysis.** It receives structured output from Layers 1 and 2
and converts numbers into narrative.

Layer 2 produces: *"Insulator Top-Right: delta-T 34°C, CRITICAL."*
Layer 3 produces: *"The top-right insulator signature is consistent with a cracked ceramic disc
allowing partial discharge. Remove from service within 24–48 hours."*

**Why it is not a substitute for Layers 1 and 2:**
A generative AI model cannot analyse thermal pixel values. It cannot run Isolation Forest.
It cannot segment an image. If you ask it to "look at the thermal image and find faults,"
it would either hallucinate results or describe what the image looks like without finding
the statistical anomaly. The algorithm (Layer 2) finds; the AI (Layer 3) explains.

**The three layers working together:**
- Layer 1 produces calibrated numbers (temperatures, delta-T)
- Layer 2 finds the statistical pattern (which pixels are anomalous, which components are flagged)
- Layer 3 translates the finding into actionable language a field crew can act on
""")

            if "insp_brief" not in st.session_state:
                if st.button("Generate AI Inspection Brief", key="gen_insp_brief"):
                    with st.spinner("Calling AI..."):
                        prompt  = _inspection_prompt(zone_stats, flagged_segs, zone_anomaly)
                        brief_text, model_used = ai_chain.complete(
                            prompt,
                            groq_key=config.GROQ_API_KEY,
                            gemini_key=config.GEMINI_API_KEY,
                        )
                        st.session_state["insp_brief"]       = brief_text or ""
                        st.session_state["insp_brief_model"]  = model_used or ""
                        if not st.session_state["insp_brief"]:
                            st.session_state["insp_brief"] = _inspection_fallback(zone_stats, flagged_segs)
                        st.rerun()

            if "insp_brief" in st.session_state:
                brief_text  = st.session_state["insp_brief"]
                model_used  = st.session_state.get("insp_brief_model", "")
                with st.expander("📋 AI Inspection Brief", expanded=True):
                    st.markdown(brief_text)
                    if model_used:
                        st.caption(f"AI response from {model_used}")
                    else:
                        st.caption("Built-in analysis — no AI key present")

                # Downloads
                st.divider()
                col1, col2 = st.columns(2)
                with col1:
                    md_text = f"# Drone Infrastructure Inspection Brief\n\n" \
                              f"**Corridor:** Catawba Valley NC — 115kV  \n" \
                              f"**Simulated data** — modelled on real UAV inspection datasets\n\n" \
                              f"---\n\n{brief_text}"
                    st.download_button(
                        "⬇ Download brief (.md)",
                        data=md_text,
                        file_name="drone_inspection_brief.md",
                        mime="text/markdown",
                        use_container_width=True,
                    )
                with col2:
                    word_bytes = _build_word_inspection(
                        zone_stats, brief_text, model_used,
                        rgb_th_bytes, iso_bytes, seg_bytes
                    )
                    st.download_button(
                        "⬇ Download brief (.docx)",
                        data=word_bytes,
                        file_name="drone_inspection_brief.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )

    # ==========================================================================
    # SUB-MODULE B — VEGETATION MAPPING
    # ==========================================================================
    with tab_b:
        st.subheader("Vegetation Mapping — High-Resolution Crown Detection")
        st.markdown(
            "A multispectral drone surveys a 30 m × 20 m section of the Catawba Valley transmission corridor — "
            "the same corridor analysed in Arc 5 with airborne LiDAR. "
            "At 10 cm/pixel the drone can delineate individual plant crowns. "
            "The result is compared directly with Arc 5 LiDAR DBSCAN crown counts to show "
            "what each method sees and what each cannot."
        )

        if st.button("▶ Run Vegetation Analysis", type="primary", key="run_veg"):
            st.session_state["veg_pending"] = True

        if st.session_state.get("veg_pending"):
            st.session_state["veg_pending"] = False

            with st.spinner("Simulating vegetation scene and running algorithms..."):
                R, G, B, NIR, NDVI, ExG, crown_params, VH, VW, PIXEL_M = _simulate_vegetation_scene()
                ndvi_bytes, canopy_density_pct = _build_ndvi_exg_fig(R, G, B, NIR, NDVI, ExG, VH, VW, PIXEL_M)
                ws_bytes, crown_props, encroaching, n_crowns = _build_watershed_fig(NDVI, VH, VW, PIXEL_M)
                rf_bytes, pred_counts, accuracy = _build_rf_fig(R, G, B, NIR, NDVI, ExG, VH, VW, PIXEL_M, crown_params)
                comp_bytes, lidar_stats = _build_comparison_fig(n_crowns, len(encroaching), PIXEL_M, VH, VW)

            st.session_state["veg_ndvi_bytes"]    = ndvi_bytes
            st.session_state["veg_canopy_density_pct"] = canopy_density_pct
            st.session_state["veg_ws_bytes"]      = ws_bytes
            st.session_state["veg_rf_bytes"]      = rf_bytes
            st.session_state["veg_comp_bytes"]    = comp_bytes
            st.session_state["veg_n_crowns"]      = n_crowns
            st.session_state["veg_n_encroaching"] = len(encroaching)
            st.session_state["veg_pred_counts"]   = pred_counts
            st.session_state["veg_accuracy"]      = accuracy
            st.session_state["veg_lidar_stats"]   = lidar_stats
            st.session_state["veg_PIXEL_M"]       = PIXEL_M
            st.session_state["veg_VH"]            = VH
            st.session_state["veg_VW"]            = VW
            st.session_state["veg_done"]          = True
            st.session_state.pop("veg_brief", None)

        if st.session_state.get("veg_done"):
            ndvi_bytes   = st.session_state["veg_ndvi_bytes"]
            canopy_density_pct = st.session_state["veg_canopy_density_pct"]
            ws_bytes     = st.session_state["veg_ws_bytes"]
            rf_bytes     = st.session_state["veg_rf_bytes"]
            comp_bytes   = st.session_state["veg_comp_bytes"]
            n_crowns     = st.session_state["veg_n_crowns"]
            n_encroach   = st.session_state["veg_n_encroaching"]
            pred_counts  = st.session_state["veg_pred_counts"]
            accuracy     = st.session_state["veg_accuracy"]
            lidar_stats  = st.session_state["veg_lidar_stats"]
            PIXEL_M      = st.session_state["veg_PIXEL_M"]
            VH           = st.session_state["veg_VH"]
            VW           = st.session_state["veg_VW"]

            # ── Layer 1 ──────────────────────────────────────────────────────
            st.divider()
            st.markdown("### Layer 1 — Signal Processing")
            with st.expander("What is Layer 1 doing here?", expanded=False):
                st.markdown("""
**Layer 1** prepares the drone imagery for the algorithms.

**NDVI = (NIR − R) / (NIR + R)**
Requires a multispectral camera with a near-infrared band. NIR is invisible to the human eye
but healthy vegetation reflects it strongly (~50%). Bare soil reflects much less.
NDVI ranges from −1 (no vegetation) to +1 (dense healthy canopy).
At 10 cm/pixel, individual leaves are partially resolved — NDVI variation within a single crown is visible.

**ExG = 2G − R − B (Excess Green)**
Requires only a standard RGB camera — no NIR needed. Much cheaper equipment.
Exaggerates the green channel relative to red and blue. Vegetation is green,
so ExG is high where plants are. Less sensitive than NDVI but widely used in
drone-based precision agriculture because it works with any consumer drone.

**Why both?** NDVI is more precise but requires multispectral hardware. ExG is less precise
but available on any RGB drone. Showing both demonstrates the hardware cost vs quality tradeoff.

**Canopy density map** = fraction of 10×10 pixel cells (1 m²) where NDVI > 0.45.
Provides a coarser view of vegetation coverage at 1 m resolution — useful for comparing
with satellite data which cannot resolve below ~10 m.
""")

            st.image(ndvi_bytes, caption="Layer 1 — NDVI, ExG, canopy density", use_container_width=True)

            # ── Layer 2: Watershed ────────────────────────────────────────────
            st.divider()
            st.markdown("### Layer 2 — Algorithm 1: Watershed Segmentation")
            with st.expander("What is watershed segmentation and how does it work?", expanded=False):
                st.markdown("""
**Watershed** is one of the oldest image segmentation algorithms (1979). The name is topographic.

Imagine inverting the NDVI image so high NDVI (crown centres) become valleys,
and low NDVI (gaps between crowns) become ridges. Then flood the landscape from
every local low point simultaneously. Where two floods would meet, a boundary forms.
Each flooded basin = one tree crown.

**The steps:**
1. **Smooth NDVI** with a Gaussian filter (radius = 3 pixels = 30 cm). Removes noise that would
   create thousands of micro-basins.
2. **Distance transform** on the binary canopy mask. Each canopy pixel gets a value equal to
   its distance from the nearest non-canopy pixel. This creates a smooth surface that peaks
   at crown centres.
3. **Find local maxima** in the distance transform. These are the crown centres — one seed per crown.
4. **Flood outward** from each seed. Boundaries form where floods meet.

**Why watershed instead of thresholding?**
Thresholding would give you one large connected blob wherever crowns touch.
Watershed separates individual crowns even when they are physically touching —
which is exactly what you need for crown counting and encroachment analysis.

**Connection to Random Forest (next step):**
Watershed tells you WHERE each crown is and how many there are.
Random Forest tells you WHAT each pixel is. Together they answer:
*how many tree canopy crowns are in the clear strip buffer zone?*
""")

            c1, c2, c3 = st.columns(3)
            c1.metric("Crowns detected", n_crowns)
            c2.metric("Encroaching crowns", n_encroach)
            c3.metric("Scene area", f"{VH * VW * PIXEL_M**2:.0f} m²")

            st.image(ws_bytes, caption="Layer 2 — Watershed crown delineation", use_container_width=True)

            # ── Layer 2: Random Forest ────────────────────────────────────────
            st.divider()
            st.markdown("### Layer 2 — Algorithm 2: Random Forest Pixel Classification")
            with st.expander("How does Random Forest work here, and how is it different from Arc 5?", expanded=False):
                st.markdown("""
**Random Forest** is the same algorithm used in Arc 5 LiDAR — but applied at 100× finer spatial resolution
with completely different input features.

**Arc 5 features (per LiDAR point):** X, Y, Z coordinates, height above ground, return number, intensity.
**Arc 6 features (per drone pixel):** R, G, B, NIR band values + NDVI + ExG + texture.

**Texture features** measure how variable pixel values are within a 5×5 pixel window (50 cm × 50 cm).
Tree canopy has high texture at 10 cm resolution — individual leaves, gaps, and shadows create a noisy pattern.
Smooth grass has low texture. Adding texture alongside spectral values makes the classifier significantly more accurate
because two land cover types can have similar mean NDVI but very different texture patterns.

**The four classes:** bare soil / low vegetation / woody shrub / tree canopy.

**Training without field data:** Real projects use GPS field samples to provide ground truth.
Here, training labels are derived from the simulation parameters — we know which pixels we intended
to be trees vs soil. In production, this is the step that requires the most fieldwork.

**How it connects to watershed:**
Watershed segments the image by NDVI topology.
Random Forest classifies each pixel by spectral and texture features.
In production: run RF first to identify confirmed canopy pixels, then run watershed
only within the canopy mask — this eliminates false crown detections in shadow or soil areas.
""")

            class_names = ['Bare soil', 'Low vegetation', 'Woody shrub', 'Tree canopy']
            import pandas as pd
            df_rf = pd.DataFrame({
                "Class": class_names,
                "Pixels": pred_counts,
                "Area (m²)": [f"{c * PIXEL_M**2:.1f}" for c in pred_counts],
                "Coverage (%)": [f"{c/(VH*VW)*100:.1f}" for c in pred_counts],
            })
            st.dataframe(df_rf, use_container_width=True, hide_index=True)
            st.image(rf_bytes, caption="Layer 2 — Random Forest classification",
                     use_container_width=True)

            # ── LiDAR comparison ──────────────────────────────────────────────
            st.divider()
            st.markdown("### Arc 5 LiDAR vs Arc 6 Drone — Direct Comparison")
            with st.expander("Why do the two methods find different crown counts?", expanded=False):
                st.markdown("""
**Neither is wrong.** They see the forest differently.

| Attribute | LiDAR + DBSCAN (Arc 5) | Drone + Watershed (Arc 6) |
|-----------|----------------------|--------------------------|
| Sensor | 3D point cloud (X, Y, Z) | 2D image (R, G, B, NIR) |
| Crown separation | 3D point density clusters | 2D NDVI local maxima |
| Resolution | ~10 cm point spacing | 10 cm pixel |
| Coverage | Full 300 m corridor | 30 m × 20 m patch |
| Tree height | Measured precisely | Not available |
| Understory | Partially visible (multi-return) | Hidden under canopy |
| Cost | Expensive (manned aircraft) | Low (consumer drone) |

**Why drone finds higher crown density:**
DBSCAN clusters 3D point density. Two touching crowns create one dense point cloud blob → one cluster.
Watershed finds 2D NDVI local maxima. Two touching crowns produce two separate NDVI peaks → two crowns.
More detection ≠ more trees. It means finer separation of individual plants.

**Production workflow:**
LiDAR first — for corridor-wide quarterly screening and height-based FAC-003 compliance.
Drone follow-up — for targeted inspection of priority zones, species identification,
and precise crown mapping where LiDAR flagged a problem.
""")

            if lidar_stats:
                col1, col2, col3, col4 = st.columns(4)
                col1.metric("LiDAR crowns (full corridor)", lidar_stats['n_crowns'])
                col2.metric("Drone crowns (30m × 20m patch)", n_crowns)
                col3.metric("LiDAR density (per 1,000 m²)", f"{lidar_stats['density']:.1f}")
                col4.metric("Drone density (per 1,000 m²)", f"{lidar_stats['drone_density']:.1f}")

                if comp_bytes:
                    st.image(comp_bytes, caption="Crown density and encroachment comparison",
                             use_container_width=False)

            # ── Layer 3: AI brief ─────────────────────────────────────────────
            st.divider()
            st.markdown("### Layer 3 — Generative AI: Vegetation Management Brief")
            with st.expander("What does Layer 3 receive and what does it produce?", expanded=False):
                st.markdown("""
Layer 3 receives structured data from Layers 1 and 2:
- Drone crown count and density
- Number of encroaching crowns
- RF pixel classification breakdown (area per class in m²)
- Arc 5 LiDAR statistics for comparison

It produces a plain-English vegetation management brief: where the two methods agree,
where they diverge and why, and how to use them together operationally.

**It does not invent the analysis.** Every claim in the brief is traceable to a Layer 2 number.
The algorithm finds; the AI explains.
""")

            if "veg_brief" not in st.session_state:
                if st.button("Generate AI Vegetation Brief", key="gen_veg_brief"):
                    with st.spinner("Calling AI..."):
                        prompt = _vegetation_prompt(
                            n_crowns, n_encroach, pred_counts,
                            PIXEL_M, VH, VW, lidar_stats,
                            canopy_density_pct=canopy_density_pct,
                        )
                        brief_text, model_used = ai_chain.complete(
                            prompt,
                            groq_key=config.GROQ_API_KEY,
                            gemini_key=config.GEMINI_API_KEY,
                        )
                        st.session_state["veg_brief"]       = brief_text or ""
                        st.session_state["veg_brief_model"] = model_used or ""
                        if not st.session_state["veg_brief"]:
                            st.session_state["veg_brief"] = _vegetation_fallback(
                                n_crowns, n_encroach, lidar_stats
                            )
                        st.rerun()

            if "veg_brief" in st.session_state:
                brief_text = st.session_state["veg_brief"]
                model_used = st.session_state.get("veg_brief_model", "")
                with st.expander("📋 AI Vegetation Management Brief", expanded=True):
                    st.markdown(brief_text)
                    if model_used:
                        st.caption(f"AI response from {model_used}")
                    else:
                        st.caption("Built-in analysis — no AI key present")

                st.divider()
                col1, col2 = st.columns(2)
                with col1:
                    md_text = f"# Drone Vegetation Mapping Brief\n\n" \
                              f"**Corridor:** Catawba Valley NC — 115kV  \n" \
                              f"**Simulated data** — modelled on real corridor survey datasets\n\n" \
                              f"---\n\n{brief_text}"
                    st.download_button(
                        "⬇ Download brief (.md)",
                        data=md_text,
                        file_name="drone_vegetation_brief.md",
                        mime="text/markdown",
                        use_container_width=True,
                    )
                with col2:
                    word_bytes = _build_word_vegetation(
                        pred_counts, n_crowns, n_encroach, accuracy,
                        brief_text, model_used, PIXEL_M, VH, VW, lidar_stats,
                        ndvi_bytes, ws_bytes, rf_bytes, comp_bytes
                    )
                    st.download_button(
                        "⬇ Download brief (.docx)",
                        data=word_bytes,
                        file_name="drone_vegetation_brief.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
